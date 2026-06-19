"""Generate the Gmail Google Approval Guidelines Word document.

Renders GMAIL_GOOGLE_APPROVAL_GUIDELINES.md into a branded .docx, reusing the
shared markdown helpers in generate_docs.py and adding fenced-code-block support
(the base renderer does not handle ``` fences).
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from generate_docs import (
    BODY_FONT_SIZE,
    BULLET_PATTERN,
    HEADING_PATTERN,
    ORDERED_LIST_PATTERN,
    add_inline_markdown_runs,
    add_marked_list_paragraph,
    add_markdown_table,
    add_styled_paragraph,
    add_text_run,
    extract_feedback_header,
    set_cell_shading,
    split_document_title,
    strip_inline_markdown,
)

BASE_DIR = Path(__file__).resolve().parent
SOURCE_FILE = "GMAIL_GOOGLE_APPROVAL_GUIDELINES.md"
OUTPUT_FILE = "GMAIL_GOOGLE_APPROVAL_GUIDELINES.docx"
CODE_FILL = "F2F2F2"


def add_code_block(doc, code_lines):
    """Render a fenced code block as a single shaded, monospace paragraph."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.4)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)

    for index, line in enumerate(code_lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # Light shading on the paragraph for a code-block feel.
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): CODE_FILL,
    })
    p_pr.append(shd)
    return paragraph


def render_body(doc, lines):
    """Markdown body renderer with fenced-code-block support."""
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        stripped = raw_line.strip()

        if not stripped:
            index += 1
            continue

        # Fenced code block ``` ... ```
        if stripped.startswith("```"):
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index].rstrip("\n"))
                index += 1
            index += 1  # consume closing fence
            add_code_block(doc, code_lines)
            continue

        if stripped == "---":
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)
            index += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            add_markdown_table(doc, table_lines)
            continue

        heading_match = HEADING_PATTERN.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = strip_inline_markdown(heading_match.group(2))
            if level <= 2:
                add_styled_paragraph(doc, heading_text, "Heading 1",
                                     space_before=Pt(18), space_after=Pt(6))
            elif level == 3:
                add_styled_paragraph(doc, heading_text, "Heading 2",
                                     space_before=Pt(12), space_after=Pt(4))
            else:
                add_styled_paragraph(doc, heading_text, "Heading 3",
                                     space_before=Pt(8), space_after=Pt(4))
            index += 1
            continue

        # Blockquote (used for required disclosures / reviewer answers)
        if stripped.startswith(">"):
            quote_text = stripped.lstrip(">").strip()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.75)
            paragraph.paragraph_format.space_after = Pt(6)
            add_inline_markdown_runs(paragraph, quote_text)
            for run in paragraph.runs:
                run.italic = True
            index += 1
            continue

        bullet_match = BULLET_PATTERN.match(raw_line)
        if bullet_match:
            level = len(bullet_match.group(1).expandtabs(2)) // 2
            add_marked_list_paragraph(doc, "-", bullet_match.group(2).strip(), level=level)
            index += 1
            continue

        ordered_match = ORDERED_LIST_PATTERN.match(raw_line)
        if ordered_match:
            level = len(ordered_match.group(1).expandtabs(2)) // 2
            add_marked_list_paragraph(
                doc, f"{ordered_match.group(2)}.", ordered_match.group(3).strip(), level=level
            )
            index += 1
            continue

        # Paragraph (accumulate wrapped lines)
        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            next_raw = lines[index]
            next_stripped = next_raw.strip()
            if not next_stripped:
                break
            if (next_stripped in ("---",) or next_stripped.startswith("|")
                    or next_stripped.startswith("```") or next_stripped.startswith(">")
                    or HEADING_PATTERN.match(next_stripped)
                    or BULLET_PATTERN.match(next_raw)
                    or ORDERED_LIST_PATTERN.match(next_raw)):
                break
            paragraph_lines.append(next_stripped)
            index += 1

        paragraph = doc.add_paragraph()
        add_inline_markdown_runs(paragraph, " ".join(paragraph_lines))
        paragraph.paragraph_format.space_after = Pt(4)


def main():
    lines = (BASE_DIR / SOURCE_FILE).read_text(encoding="utf-8-sig").splitlines()
    title, metadata, body_lines = extract_feedback_header(lines)
    main_title, subtitle = split_document_title(title)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # --- TITLE PAGE ---
    doc.add_paragraph()
    doc.add_paragraph()
    add_styled_paragraph(doc, "VELVET ELVES", "Title", bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(28))
    add_styled_paragraph(doc, "AI-First Transaction Management Platform", "Subtitle",
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(18))
    doc.add_paragraph()
    add_styled_paragraph(doc, main_title or "GMAIL INTEGRATION GOOGLE APPROVAL GUIDELINES",
                         "Heading 1", bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(20))
    if subtitle:
        add_styled_paragraph(doc, subtitle, "Heading 2",
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14))
    doc.add_paragraph()

    if metadata:
        table = doc.add_table(rows=len(metadata), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Light Grid Accent 1"
        for row_index, (label, value) in enumerate(metadata):
            table.rows[row_index].cells[0].text = label
            cell_value = table.rows[row_index].cells[1]
            cell_value.text = ""
            add_inline_markdown_runs(cell_value.paragraphs[0], value)
            for column_index, cell in enumerate(table.rows[row_index].cells):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = BODY_FONT_SIZE
                        if column_index == 0:
                            run.bold = True
            set_cell_shading(table.rows[row_index].cells[0], "EEF2F8")

    doc.add_page_break()

    # --- CONTENT ---
    render_body(doc, body_lines)

    output_path = BASE_DIR / OUTPUT_FILE
    doc.save(output_path)
    print(f"Created {output_path}")


if __name__ == "__main__":
    main()
