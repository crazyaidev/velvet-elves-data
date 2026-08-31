"""Render AIME_AND_YOUR_TRANSACTION_FILE_A.md to a branded .docx.

The markdown is a Word export (escaped checkboxes, bold headings, a base64
screenshot) plus Jan's 18 Aug 2026 replies under each of Jake's answers.
"""
from __future__ import annotations

import base64
import re
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from generate_docs import (
    BODY_FONT_SIZE,
    BULLET_PATTERN,
    HEADING_PATTERN,
    ORDERED_LIST_PATTERN,
    TABLE_HEADER_FILL,
    add_inline_markdown_runs,
    add_markdown_table,
    add_marked_list_paragraph,
    add_styled_paragraph,
    add_text_run,
    set_cell_shading,
    split_markdown_table_row,
    strip_inline_markdown,
)

BASE_DIR = Path(__file__).resolve().parent
SOURCE_FILE = "AIME_AND_YOUR_TRANSACTION_FILE_A.md"
OUTPUT_FILE = "AIME_AND_YOUR_TRANSACTION_FILE_A.docx"

JAN_FILL = "DEEBF7"
JAN_HEADING = RGBColor(0x1F, 0x4E, 0x79)
CHECKBOX_RE = re.compile(r"^-\s+\[([ xX]*)\]\s*(.*)$")
IMAGE_REF_RE = re.compile(r"^!\[\]\[image1\]\s*$")
IMAGE_DATA_RE = re.compile(
    r"\[image1\]:\s*<data:image/png;base64,(.+)>", re.DOTALL
)
UNESCAPE_RE = re.compile(r"\\([-*\[\]_#])")
JAN_HEADING_RE = re.compile(r"jan.?s response", re.IGNORECASE)


def unescape_word_md(text: str) -> str:
    """Undo Word-export backslash escapes: \\- \\[ \\] \\* \\_."""
    return UNESCAPE_RE.sub(r"\1", text)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): fill,
        },
    )
    p_pr.append(shd)


def extract_png_bytes(raw_text: str) -> bytes | None:
    match = IMAGE_DATA_RE.search(raw_text)
    if not match:
        return None
    payload = match.group(1).strip()
    return base64.b64decode(payload)


def parse_source() -> tuple[list[tuple[str, str]], list[str], bytes | None]:
    raw = (BASE_DIR / SOURCE_FILE).read_text(encoding="utf-8-sig")
    png_bytes = extract_png_bytes(raw)
    lines = [unescape_word_md(line) for line in raw.splitlines()]

    metadata: list[tuple[str, str]] = []
    body_start = 0
    index = 0
    while index < len(lines):
        stripped = lines[index].strip()
        if stripped.startswith("|") and "Prepared for" in stripped:
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            parsed = [split_markdown_table_row(row) for row in table_lines]
            for row in parsed:
                if len(row) < 2 or not row[0] or set(row[0]) <= {":", "-"}:
                    continue
                metadata.append(
                    (strip_inline_markdown(row[0]), strip_inline_markdown(row[1]))
                )
            continue
        heading = HEADING_PATTERN.match(stripped)
        if heading and "Before you start" in strip_inline_markdown(heading.group(2)):
            body_start = index
            break
        index += 1

    body_lines = []
    for line in lines[body_start:]:
        if line.strip().startswith("[image1]:"):
            continue
        body_lines.append(line)
    return metadata, body_lines, png_bytes


def add_checkbox(doc, checked: bool, text: str, jan: bool) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.63)
    paragraph.paragraph_format.first_line_indent = Cm(-0.45)
    paragraph.paragraph_format.space_after = Pt(2)
    marker = "☑ " if checked else "☐ "
    add_text_run(paragraph, marker, size=BODY_FONT_SIZE)
    add_inline_markdown_runs(paragraph, text)
    if jan:
        shade_paragraph(paragraph, JAN_FILL)


def add_screenshot(doc, png_bytes: bytes) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run()
    run.add_picture(BytesIO(png_bytes), width=Inches(6.3))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(10)
    cap_run = caption.add_run(
        "Settings → AI & Automation — How it runs (Jake’s screenshot)"
    )
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = RGBColor(0x5B, 0x5B, 0x5B)


def add_jan_heading(doc, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(strip_inline_markdown(text).rstrip(":"))
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = JAN_HEADING
    shade_paragraph(paragraph, JAN_FILL)


def add_body_paragraph(doc, text: str, jan: bool) -> None:
    paragraph = doc.add_paragraph()
    add_inline_markdown_runs(paragraph, text)
    paragraph.paragraph_format.space_after = Pt(4)
    if jan:
        shade_paragraph(paragraph, JAN_FILL)
        for run in paragraph.runs:
            if run.font.color.rgb is None:
                run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def shade_last_table(doc) -> None:
    if not doc.tables:
        return
    table = doc.tables[-1]
    for row_index, row in enumerate(table.rows):
        fill = "BDD7EE" if row_index == 0 else JAN_FILL
        for cell in row.cells:
            set_cell_shading(cell, fill)


def render_body(doc, lines: list[str], png_bytes: bytes | None) -> None:
    index = 0
    jan = False
    while index < len(lines):
        raw_line = lines[index]
        stripped = raw_line.strip()

        if not stripped:
            index += 1
            continue

        if IMAGE_REF_RE.match(stripped):
            if png_bytes:
                add_screenshot(doc, png_bytes)
            index += 1
            continue

        heading_match = HEADING_PATTERN.match(stripped)
        if heading_match:
            jan = False
            level = len(heading_match.group(1))
            heading_text = strip_inline_markdown(heading_match.group(2))
            style = "Heading 1" if level <= 1 else "Heading 2"
            before = Pt(18) if level <= 1 else Pt(14)
            add_styled_paragraph(
                doc,
                heading_text,
                style,
                space_before=before,
                space_after=Pt(6),
            )
            index += 1
            continue

        if JAN_HEADING_RE.search(stripped) and stripped.startswith("**"):
            jan = True
            add_jan_heading(doc, stripped)
            index += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            add_markdown_table(doc, table_lines)
            if jan:
                shade_last_table(doc)
            continue

        checkbox_match = CHECKBOX_RE.match(stripped)
        if checkbox_match:
            checked = "x" in checkbox_match.group(1).lower()
            add_checkbox(doc, checked, checkbox_match.group(2), jan)
            index += 1
            continue

        bullet_match = BULLET_PATTERN.match(raw_line)
        if bullet_match:
            level = len(bullet_match.group(1).expandtabs(2)) // 2
            paragraph = add_marked_list_paragraph(
                doc, "–", bullet_match.group(2).strip(), level=level
            )
            if jan:
                shade_paragraph(paragraph, JAN_FILL)
            index += 1
            continue

        ordered_match = ORDERED_LIST_PATTERN.match(raw_line)
        if ordered_match:
            level = len(ordered_match.group(1).expandtabs(2)) // 2
            paragraph = add_marked_list_paragraph(
                doc,
                f"{ordered_match.group(2)}.",
                ordered_match.group(3).strip(),
                level=level,
            )
            if jan:
                shade_paragraph(paragraph, JAN_FILL)
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            next_raw = lines[index]
            next_stripped = next_raw.strip()
            if not next_stripped:
                break
            if (
                next_stripped.startswith("|")
                or IMAGE_REF_RE.match(next_stripped)
                or HEADING_PATTERN.match(next_stripped)
                or CHECKBOX_RE.match(next_stripped)
                or BULLET_PATTERN.match(next_raw)
                or ORDERED_LIST_PATTERN.match(next_raw)
                or (JAN_HEADING_RE.search(next_stripped) and next_stripped.startswith("**"))
            ):
                break
            paragraph_lines.append(next_stripped)
            index += 1

        add_body_paragraph(doc, " ".join(paragraph_lines), jan)


def build_title_page(doc, metadata: list[tuple[str, str]]) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    doc.add_paragraph()
    doc.add_paragraph()
    add_styled_paragraph(
        doc,
        "VELVET ELVES",
        "Title",
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        size=Pt(28),
    )
    add_styled_paragraph(
        doc,
        "AI-First Transaction Management Platform",
        "Subtitle",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        size=Pt(18),
    )
    doc.add_paragraph()
    add_styled_paragraph(
        doc,
        "AIME AND YOUR TRANSACTION FILE",
        "Heading 1",
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        size=Pt(20),
    )
    add_styled_paragraph(
        doc,
        "Jake’s answers and Jan’s replies",
        "Heading 2",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        size=Pt(14),
    )
    doc.add_paragraph()

    if not metadata:
        return

    table = doc.add_table(rows=len(metadata), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    for row_index, (label, value) in enumerate(metadata):
        table.rows[row_index].cells[0].text = label
        table.rows[row_index].cells[1].text = value
        for column_index, cell in enumerate(table.rows[row_index].cells):
            set_cell_shading(cell, TABLE_HEADER_FILL if column_index == 0 else "FFFFFF")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = BODY_FONT_SIZE
                    if column_index == 0:
                        run.bold = True


def main() -> None:
    metadata, body_lines, png_bytes = parse_source()
    doc = Document()
    build_title_page(doc, metadata)
    doc.add_page_break()
    render_body(doc, body_lines, png_bytes)
    output_path = BASE_DIR / OUTPUT_FILE
    doc.save(output_path)
    print(f"Created {output_path}")


if __name__ == "__main__":
    main()
