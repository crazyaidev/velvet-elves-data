"""Render AIME_AI_AUTOMATION_TESTING_GUIDELINES_2026-08-18.md to a branded .docx."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from generate_docs import (
    add_styled_paragraph,
    read_source_lines,
    render_feedback_body,
    save_document,
    split_document_title,
)

SOURCE_FILE = "AIME_AI_AUTOMATION_TESTING_GUIDELINES_2026-08-18.md"
OUTPUT_FILE = "AIME_AI_AUTOMATION_TESTING_GUIDELINES_2026-08-18.docx"


def create_aime_ai_automation_testing_guidelines_doc() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    lines = read_source_lines(SOURCE_FILE)
    title = "Aime and AI Automation Testing"
    body_start = 0
    for index, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith("# "):
            title = stripped[2:].strip()
            body_start = index + 1
            break

    main_title, subtitle = split_document_title(title)

    doc.add_paragraph()
    doc.add_paragraph()
    add_styled_paragraph(
        doc,
        "VELVET ELVES",
        "Title",
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        size=Pt(28),
    )
    add_styled_paragraph(
        doc,
        "AI-First Transaction Management Platform",
        "Subtitle",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        size=Pt(18),
    )
    doc.add_paragraph()
    add_styled_paragraph(
        doc,
        main_title,
        "Heading 1",
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        size=Pt(20),
    )
    if subtitle:
        add_styled_paragraph(
            doc,
            subtitle,
            "Heading 2",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            size=Pt(14),
        )
    doc.add_paragraph()
    doc.add_page_break()
    render_feedback_body(doc, lines[body_start:])
    save_document(doc, OUTPUT_FILE)


if __name__ == "__main__":
    create_aime_ai_automation_testing_guidelines_doc()
    print(f"Wrote {Path(__file__).resolve().parent / OUTPUT_FILE}")
