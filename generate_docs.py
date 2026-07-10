"""Generate Word documents from project source files."""
import argparse
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


BASE_DIR = Path(__file__).resolve().parent
BODY_FONT_SIZE = Pt(10)
TABLE_HEADER_FILL = "D9E2F3"
INLINE_MARKDOWN_PATTERN = re.compile(r"(\*\*.+?\*\*|`.+?`)")
HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+)$")
BULLET_PATTERN = re.compile(r"^(\s*)-\s+(.+)$")
ORDERED_LIST_PATTERN = re.compile(r"^(\s*)(\d+)\.\s+(.+)$")
METADATA_PATTERN = re.compile(r"^\*\*(.+?):\*\*\s*(.+)$")
TABLE_DIVIDER_PATTERN = re.compile(r"^:?-{3,}:?$")


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shd)


def add_styled_paragraph(doc, text, style_name, bold=False, color=None, size=None, alignment=None, space_after=None, space_before=None):
    """Add a paragraph with optional styling."""
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    return p


def read_source_lines(filename):
    """Read a source file from the repo root."""
    return (BASE_DIR / filename).read_text(encoding="utf-8-sig").splitlines()


def save_document(doc, filename):
    """Save a document into the repo root."""
    output_path = BASE_DIR / filename
    doc.save(output_path)
    print(f"Created {output_path.name}")


def add_text_run(paragraph, text, bold=False, code=False, size=BODY_FONT_SIZE):
    """Add a run with consistent sizing and optional inline-code styling."""
    if not text:
        return None

    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = size
    if code:
        run.font.name = "Consolas"
    return run


def add_inline_markdown_runs(paragraph, text, size=BODY_FONT_SIZE, force_bold=False):
    """Render simple markdown inline styles inside a paragraph."""
    cursor = 0

    for match in INLINE_MARKDOWN_PATTERN.finditer(text):
        if match.start() > cursor:
            add_text_run(paragraph, text[cursor:match.start()], bold=force_bold, size=size)

        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            add_text_run(paragraph, token[2:-2], bold=True, size=size)
        else:
            add_text_run(paragraph, token[1:-1], bold=force_bold, code=True, size=size)

        cursor = match.end()

    if cursor < len(text):
        add_text_run(paragraph, text[cursor:], bold=force_bold, size=size)


def strip_inline_markdown(text):
    """Remove inline markdown markers for headings and titles."""
    return text.replace("**", "").replace("`", "")


def split_markdown_table_row(line):
    """Split a markdown table row into cells."""
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_markdown_table_divider(cells):
    """Detect the markdown divider row between table header and body."""
    return bool(cells) and all(TABLE_DIVIDER_PATTERN.fullmatch(cell.replace(" ", "")) for cell in cells)


def set_markdown_cell_text(cell, text, force_bold=False):
    """Write markdown-aware text into a table cell."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    add_inline_markdown_runs(paragraph, text, force_bold=force_bold)
    paragraph.paragraph_format.space_after = Pt(0)


def add_markdown_table(doc, table_lines):
    """Render a basic markdown table into the Word document."""
    parsed_rows = [split_markdown_table_row(line) for line in table_lines]
    parsed_rows = [row for row in parsed_rows if any(cell for cell in row)]
    if not parsed_rows:
        return

    header = parsed_rows[0]
    data_rows = parsed_rows[1:]
    if data_rows and is_markdown_table_divider(data_rows[0]):
        data_rows = data_rows[1:]

    all_rows = [header] + data_rows
    column_count = max(len(row) for row in all_rows)

    table = doc.add_table(rows=len(all_rows), cols=column_count)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for column_index in range(column_count):
        header_text = header[column_index] if column_index < len(header) else ""
        cell = table.rows[0].cells[column_index]
        set_markdown_cell_text(cell, header_text, force_bold=True)
        set_cell_shading(cell, TABLE_HEADER_FILL)

    for row_index, row_data in enumerate(data_rows, start=1):
        for column_index in range(column_count):
            value = row_data[column_index] if column_index < len(row_data) else ""
            set_markdown_cell_text(table.rows[row_index].cells[column_index], value)

    doc.add_paragraph()


def extract_feedback_header(lines):
    """Extract the title and metadata block from design-feedback.md."""
    title = "Design Feedback"
    metadata = []
    body_start = len(lines)

    for index, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue

        heading_match = re.match(r"^#\s+(.+)$", stripped)
        if heading_match and title == "Design Feedback":
            title = heading_match.group(1).strip()
            continue

        metadata_match = METADATA_PATTERN.match(stripped)
        if metadata_match:
            metadata.append((metadata_match.group(1).strip(), metadata_match.group(2).strip()))
            continue

        if stripped == "---":
            body_start = index + 1
            break

        body_start = index
        break

    return title, metadata, lines[body_start:]


def split_document_title(title):
    """Split a markdown title into primary and secondary title lines."""
    parts = re.split(r"\s+[\u2014-]\s+", title, maxsplit=1)
    main_title = parts[0].strip().upper() if parts else "DESIGN FEEDBACK"
    subtitle = parts[1].strip() if len(parts) > 1 else ""
    return main_title, subtitle


def add_marked_list_paragraph(doc, marker, text, level=0, size=BODY_FONT_SIZE):
    """Add an indented paragraph that uses a visible list marker."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.63 * (level + 1))
    paragraph.paragraph_format.first_line_indent = Cm(-0.45)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text_run(paragraph, f"{marker} ", size=size)
    add_inline_markdown_runs(paragraph, text, size=size)
    return paragraph


def render_feedback_body(doc, lines):
    """Render the main markdown body of design-feedback.md."""
    index = 0

    while index < len(lines):
        raw_line = lines[index]
        stripped = raw_line.strip()

        if not stripped:
            index += 1
            continue

        if stripped == "---":
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)
            index += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            add_markdown_table(doc, table_lines)
            continue

        heading_match = HEADING_PATTERN.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = strip_inline_markdown(heading_match.group(2))

            if level <= 2:
                add_styled_paragraph(doc, heading_text, 'Heading 1',
                                     space_before=Pt(18), space_after=Pt(6))
            elif level == 3:
                add_styled_paragraph(doc, heading_text, 'Heading 2',
                                     space_before=Pt(12), space_after=Pt(4))
            else:
                add_styled_paragraph(doc, heading_text, 'Heading 3',
                                     space_before=Pt(8), space_after=Pt(4))

            index += 1
            continue

        bullet_match = BULLET_PATTERN.match(raw_line)
        if bullet_match:
            level = len(bullet_match.group(1).expandtabs(2)) // 2
            add_marked_list_paragraph(doc, "-", bullet_match.group(2).strip(), level=level)
            index += 1
            continue

        ordered_match = ORDERED_LIST_PATTERN.match(raw_line)
        if ordered_match:
            level = len(ordered_match.group(1).expandtabs(2)) // 2
            add_marked_list_paragraph(doc, f"{ordered_match.group(2)}.", ordered_match.group(3).strip(), level=level)
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1

        while index < len(lines):
            next_raw = lines[index]
            next_stripped = next_raw.strip()
            if not next_stripped:
                break
            if next_stripped == "---" or next_stripped.startswith("|") or \
               HEADING_PATTERN.match(next_stripped) or BULLET_PATTERN.match(next_raw) or \
               ORDERED_LIST_PATTERN.match(next_raw):
                break
            paragraph_lines.append(next_stripped)
            index += 1

        paragraph = doc.add_paragraph()
        add_inline_markdown_runs(paragraph, " ".join(paragraph_lines))
        paragraph.paragraph_format.space_after = Pt(4)


def create_requirements_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # --- TITLE PAGE ---
    doc.add_paragraph()
    doc.add_paragraph()
    add_styled_paragraph(doc, "VELVET ELVES", 'Title', bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(28))
    add_styled_paragraph(doc, "AI-First Transaction Management Platform", 'Subtitle',
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(18))
    doc.add_paragraph()
    add_styled_paragraph(doc, "ESTABLISHED PROJECT REQUIREMENTS", 'Heading 1',
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # Info table
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    info = [
        ("Date Established", "March 4, 2026"),
        ("Timeline", "22 Weeks (MVP Delivery)"),
        ("Budget", "Under USD $40,000 (Fixed-Price)"),
        ("Tech Stack", "React (Frontend) | FastAPI (Backend) | OpenAI GPT (AI) | Supabase (Database & Auth) | AWS EC2 (Hosting)"),
        ("Project Type", "Complex SaaS Rebuild"),
    ]
    for i, (label, value) in enumerate(info):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if cell == table.rows[i].cells[0]:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_page_break()

    # --- TABLE OF CONTENTS placeholder ---
    add_styled_paragraph(doc, "Table of Contents", 'Heading 1', space_after=Pt(12))
    toc_items = [
        "1. User Management & Authentication",
        "2. Transaction Management",
        "3. The Wizard (AI-Driven Transaction Intake Engine)",
        "4. Task Engine (Most Critical Component)",
        "5. Document Management",
        "6. Communication Engine",
        "7. Integrations",
        "8. AI & Automation",
        "9. User Interface & Experience",
        "10. Data Migration & Security",
        "11. Advertising & Monetization",
        "12. Future Features (Non-MVP)",
        "13. Deployment & Infrastructure",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # --- CONTENT ---
    # Read the requirements file
    lines = read_source_lines("requirements.txt")

    # Parse and format
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip the header block (already on title page)
        if stripped.startswith("====") or stripped == "" or stripped.startswith("Date Established") or \
           stripped.startswith("Timeline:") or stripped.startswith("Budget:") or \
           stripped.startswith("Tech Stack:") or stripped == "VELVET ELVES - AI-FIRST TRANSACTION MANAGEMENT PLATFORM" or \
           stripped == "ESTABLISHED PROJECT REQUIREMENTS" or stripped == "END OF REQUIREMENTS":
            i += 1
            continue

        # Major section headers (e.g., "1. USER MANAGEMENT & AUTHENTICATION")
        major_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if major_match and stripped == stripped.upper():
            section_num = major_match.group(1)
            section_title = major_match.group(2).title()
            add_styled_paragraph(doc, f"{section_num}. {section_title}", 'Heading 1',
                                 space_before=Pt(18), space_after=Pt(8))
            i += 1
            continue

        # Sub-section headers (e.g., "1.1 Account Registration & Login")
        sub_match = re.match(r'^(\d+\.\d+)\s+(.+)$', stripped)
        if sub_match:
            sub_num = sub_match.group(1)
            sub_title = sub_match.group(2)
            add_styled_paragraph(doc, f"{sub_num} {sub_title}", 'Heading 2',
                                 space_before=Pt(12), space_after=Pt(6))
            i += 1
            continue

        # Role headers (e.g., "a) Agent")
        role_match = re.match(r'^([a-f])\)\s+(.+)$', stripped)
        if role_match:
            add_styled_paragraph(doc, stripped, 'Heading 3',
                                 space_before=Pt(8), space_after=Pt(4))
            i += 1
            continue

        # Sub-headers like "Roles:", "Deliverables:", etc.
        if stripped.endswith(':') and len(stripped) < 60 and not stripped.startswith('-') and not stripped.startswith('['):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Bullet items (lines starting with -)
        if stripped.startswith('- '):
            text = stripped[2:]
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Sub-bullet items (indented deeper)
        if stripped.startswith('CANNOT') or stripped.startswith('Cannot'):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            run.font.size = Pt(10)
            i += 1
            continue

        # Numbered items within text (e.g., "1. Buyer - Financing")
        num_item = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if num_item and stripped != stripped.upper():
            p = doc.add_paragraph(stripped, style='List Number')
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Regular text
        if stripped:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.size = Pt(10)

        i += 1

    save_document(doc, "requirements.docx")


def create_milestones_doc():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # --- TITLE PAGE ---
    doc.add_paragraph()
    doc.add_paragraph()
    add_styled_paragraph(doc, "VELVET ELVES", 'Title', bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(28))
    add_styled_paragraph(doc, "AI-First Transaction Management Platform", 'Subtitle',
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(18))
    doc.add_paragraph()
    add_styled_paragraph(doc, "PROJECT MILESTONES", 'Heading 1',
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "22-Week Timeline", 'Heading 2',
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # Info table
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    info = [
        ("Date Established", "March 4, 2026"),
        ("Start Date", "Week 1 (March 9, 2026)"),
        ("End Date", "Week 22 (August 2, 2026)"),
        ("Budget", "Under USD $40,000 (Fixed-Price)"),
        ("Tech Stack", "React | FastAPI | OpenAI GPT | Supabase | AWS EC2"),
    ]
    for i, (label, value) in enumerate(info):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if cell == table.rows[i].cells[0]:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_page_break()

    # --- MILESTONE SUMMARY TABLE ---
    add_styled_paragraph(doc, "Milestone Summary", 'Heading 1', space_after=Pt(12))

    summary_table = doc.add_table(rows=8, cols=4)
    summary_table.style = 'Medium Shading 1 Accent 1'
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Phase", "Weeks", "Focus Area", "Duration"]
    for j, h in enumerate(headers):
        summary_table.rows[0].cells[j].text = h

    summary_data = [
        ("1", "1-3", "Discovery, Architecture & Foundation", "3 weeks"),
        ("2", "4-8", "Core Transaction & Task Engine", "5 weeks"),
        ("3", "9-13", "AI Wizard & Document Management", "5 weeks"),
        ("4", "14-16", "Communication Engine & AI Email", "3 weeks"),
        ("5", "17-19", "Dashboards, Payments & Profiles", "3 weeks"),
        ("6", "20-21", "White-Label, Integrations & Ad Hooks", "2 weeks"),
        ("7", "22", "Testing, Deployment & Launch", "1 week"),
    ]
    for i, row_data in enumerate(summary_data):
        for j, val in enumerate(row_data):
            summary_table.rows[i+1].cells[j].text = val

    doc.add_page_break()

    # --- DETAILED MILESTONES ---
    lines = read_source_lines("milestones.txt")

    i = 0
    in_summary_section = False
    in_dependencies_section = False
    in_roadmap_section = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip header block and separator lines
        if stripped.startswith("====") or stripped.startswith("------") or stripped.startswith("------|"):
            i += 1
            continue

        if stripped in ["", "VELVET ELVES - AI-FIRST TRANSACTION MANAGEMENT PLATFORM",
                        "PROJECT MILESTONES (22-WEEK TIMELINE)",
                        "END OF MILESTONES"]:
            i += 1
            continue

        if stripped.startswith("Date Established:") or stripped.startswith("Start Date:") or \
           stripped.startswith("End Date:") or stripped.startswith("Budget:") or \
           stripped.startswith("Tech Stack:"):
            i += 1
            continue

        # Skip the text summary table (we already made a nice one)
        if stripped == "MILESTONE SUMMARY":
            in_summary_section = True
            i += 1
            continue
        if in_summary_section:
            if stripped.startswith("KEY DEPENDENCIES"):
                in_summary_section = False
            elif stripped.startswith("Phase") or stripped.startswith("TOTAL") or re.match(r'^\d+\s+\|', stripped):
                i += 1
                continue
            else:
                i += 1
                continue

        # Phase headers (e.g., "PHASE 1: DISCOVERY...")
        phase_match = re.match(r'^PHASE\s+(\d+):\s+(.+)$', stripped)
        if phase_match:
            phase_num = phase_match.group(1)
            phase_title = phase_match.group(2).title()
            add_styled_paragraph(doc, f"Phase {phase_num}: {phase_title}", 'Heading 1',
                                 space_before=Pt(24), space_after=Pt(4))
            # Next line should be the week range
            i += 1
            if i < len(lines):
                week_line = lines[i].strip()
                if week_line.startswith("Weeks") or week_line.startswith("Week"):
                    p = doc.add_paragraph()
                    run = p.add_run(week_line)
                    run.italic = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                    p.paragraph_format.space_after = Pt(8)
            i += 1
            continue

        # Milestone headers (e.g., "MILESTONE 1.1 - Project Setup...")
        milestone_match = re.match(r'^MILESTONE\s+(\d+\.\d+)\s+-\s+(.+?)(?:\s+\((.+?)\))?$', stripped)
        if milestone_match:
            m_num = milestone_match.group(1)
            m_title = milestone_match.group(2)
            m_weeks = milestone_match.group(3) or ""
            heading_text = f"Milestone {m_num} - {m_title}"
            if m_weeks:
                heading_text += f" ({m_weeks})"
            add_styled_paragraph(doc, heading_text, 'Heading 2',
                                 space_before=Pt(14), space_after=Pt(6))
            i += 1
            continue

        # Section labels
        if stripped == "Deliverables:":
            p = doc.add_paragraph()
            run = p.add_run("Deliverables:")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if stripped.startswith("Success Criteria:"):
            p = doc.add_paragraph()
            run = p.add_run("Success Criteria: ")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            criteria_text = stripped.replace("Success Criteria: ", "").replace("Success Criteria:", "")
            run2 = p.add_run(criteria_text)
            run2.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            # Continuation lines
            while i < len(lines):
                next_stripped = lines[i].strip()
                if next_stripped and not next_stripped.startswith("[") and not next_stripped.startswith("MILESTONE") and \
                   not next_stripped.startswith("PHASE") and not next_stripped.startswith("====") and \
                   not next_stripped.startswith("KEY") and not next_stripped.startswith("POST-MVP"):
                    run3 = p.add_run(" " + next_stripped)
                    run3.font.size = Pt(10)
                    i += 1
                else:
                    break
            continue

        # Checkbox items
        if stripped.startswith("[ ]"):
            text = stripped[4:]
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run("[ ] " + text)
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(1.27)
            i += 1
            # Check for continuation lines (indented sub-items starting with -)
            while i < len(lines):
                next_line = lines[i]
                next_stripped = next_line.strip()
                if next_stripped.startswith("- ") and next_line.startswith("      "):
                    sub_text = next_stripped[2:]
                    sp = doc.add_paragraph(style='List Bullet 2')
                    sr = sp.add_run(sub_text)
                    sr.font.size = Pt(9)
                    sp.paragraph_format.space_after = Pt(1)
                    sp.paragraph_format.left_indent = Cm(2.54)
                    i += 1
                elif next_stripped and not next_stripped.startswith("[") and not next_stripped.startswith("Success") and \
                     not next_stripped.startswith("MILESTONE") and not next_stripped.startswith("PHASE") and \
                     not next_stripped.startswith("====") and not next_stripped.startswith("KEY") and \
                     not next_stripped.startswith("POST-MVP") and not next_stripped.startswith("Deliverables") and \
                     next_line.startswith("      "):
                    # Continuation of the checkbox item
                    run2 = p.add_run(" " + next_stripped)
                    run2.font.size = Pt(10)
                    i += 1
                else:
                    break
            continue

        # Key Dependencies section
        if stripped == "KEY DEPENDENCIES & RISKS":
            in_dependencies_section = True
            doc.add_page_break()
            add_styled_paragraph(doc, "Key Dependencies & Risks", 'Heading 1',
                                 space_before=Pt(18), space_after=Pt(8))
            i += 1
            continue

        # Post-MVP section
        if stripped == "POST-MVP ROADMAP (Beyond 22 Weeks)":
            in_roadmap_section = True
            in_dependencies_section = False
            add_styled_paragraph(doc, "Post-MVP Roadmap (Beyond 22 Weeks)", 'Heading 1',
                                 space_before=Pt(18), space_after=Pt(8))
            i += 1
            continue

        # Numbered dependency items
        dep_match = re.match(r'^(\d+)\.\s+(.+?):\s+(.+)$', stripped)
        if dep_match and in_dependencies_section:
            num = dep_match.group(1)
            title = dep_match.group(2)
            desc = dep_match.group(3)
            p = doc.add_paragraph()
            run_num = p.add_run(f"{num}. {title}: ")
            run_num.bold = True
            run_num.font.size = Pt(10)
            run_desc = p.add_run(desc)
            run_desc.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Roadmap items (starting with -)
        if stripped.startswith("- ") and in_roadmap_section:
            text = stripped[2:]
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Default: regular paragraph
        if stripped:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.size = Pt(10)

        i += 1

    save_document(doc, "milestones.docx")


def create_design_feedback_doc():
    """Generate a .docx export of design-feedback.md."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    lines = read_source_lines("design-feedback.md")
    title, metadata, body_lines = extract_feedback_header(lines)
    main_title, subtitle = split_document_title(title)

    # --- TITLE PAGE ---
    doc.add_paragraph()
    doc.add_paragraph()
    add_styled_paragraph(doc, "VELVET ELVES", 'Title', bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(28))
    add_styled_paragraph(doc, "AI-First Transaction Management Platform", 'Subtitle',
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(18))
    doc.add_paragraph()
    add_styled_paragraph(doc, main_title, 'Heading 1', bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(20))
    if subtitle:
        add_styled_paragraph(doc, subtitle, 'Heading 2',
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14))
    doc.add_paragraph()

    if metadata:
        table = doc.add_table(rows=len(metadata), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Light Grid Accent 1'
        for row_index, (label, value) in enumerate(metadata):
            table.rows[row_index].cells[0].text = label
            table.rows[row_index].cells[1].text = value
            for column_index, cell in enumerate(table.rows[row_index].cells):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = BODY_FONT_SIZE
                        if column_index == 0:
                            run.bold = True

    doc.add_page_break()

    # --- CONTENT ---
    render_feedback_body(doc, body_lines)

    save_document(doc, "design-feedback.docx")


def create_transaction_system_doc():
    """Generate a .docx export of TRANSACTION_SYSTEM_GUIDE.md (real-estate audience).

    Reuses the generic markdown renderer (extract_feedback_header +
    render_feedback_body) so the guide gets the same title page, metadata table,
    headings, bullets, and tables as the other generated documents.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    lines = read_source_lines("TRANSACTION_SYSTEM_GUIDE.md")
    title, metadata, body_lines = extract_feedback_header(lines)
    main_title, subtitle = split_document_title(title)

    # --- TITLE PAGE ---
    doc.add_paragraph()
    doc.add_paragraph()
    add_styled_paragraph(doc, "VELVET ELVES", 'Title', bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(28))
    add_styled_paragraph(doc, "AI-First Transaction Management Platform", 'Subtitle',
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(18))
    doc.add_paragraph()
    add_styled_paragraph(doc, main_title, 'Heading 1', bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(20))
    if subtitle:
        add_styled_paragraph(doc, subtitle, 'Heading 2',
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14))
    doc.add_paragraph()

    if metadata:
        table = doc.add_table(rows=len(metadata), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Light Grid Accent 1'
        for row_index, (label, value) in enumerate(metadata):
            table.rows[row_index].cells[0].text = label
            table.rows[row_index].cells[1].text = value
            for column_index, cell in enumerate(table.rows[row_index].cells):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = BODY_FONT_SIZE
                        if column_index == 0:
                            run.bold = True

    doc.add_page_break()

    # --- CONTENT ---
    render_feedback_body(doc, body_lines)

    save_document(doc, "TRANSACTION_SYSTEM_GUIDE.docx")


def create_payment_system_doc():
    """Generate a .docx export of PAYMENT_SYSTEM_GUIDE.md (stakeholder audience).

    Mirrors create_transaction_system_doc: reuses the generic markdown renderer
    (extract_feedback_header + render_feedback_body) so the guide gets the same
    title page, metadata table, headings, bullets, and tables as the other
    generated documents.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    lines = read_source_lines("PAYMENT_SYSTEM_GUIDE.md")
    title, metadata, body_lines = extract_feedback_header(lines)
    main_title, subtitle = split_document_title(title)

    # --- TITLE PAGE ---
    doc.add_paragraph()
    doc.add_paragraph()
    add_styled_paragraph(doc, "VELVET ELVES", 'Title', bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(28))
    add_styled_paragraph(doc, "AI-First Transaction Management Platform", 'Subtitle',
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(18))
    doc.add_paragraph()
    add_styled_paragraph(doc, main_title, 'Heading 1', bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(20))
    if subtitle:
        add_styled_paragraph(doc, subtitle, 'Heading 2',
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14))
    doc.add_paragraph()

    if metadata:
        table = doc.add_table(rows=len(metadata), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Light Grid Accent 1'
        for row_index, (label, value) in enumerate(metadata):
            table.rows[row_index].cells[0].text = label
            table.rows[row_index].cells[1].text = value
            for column_index, cell in enumerate(table.rows[row_index].cells):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = BODY_FONT_SIZE
                        if column_index == 0:
                            run.bold = True

    doc.add_page_break()

    # --- CONTENT ---
    render_feedback_body(doc, body_lines)

    save_document(doc, "PAYMENT_SYSTEM_GUIDE.docx")


def create_godaddy_route53_config_doc():
    """Generate a .docx export of the GoDaddy Route 53 configuration guide."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    lines = read_source_lines("GODADDY_ROUTE53_CONFIGURATION_GUIDE.md")
    title, metadata, body_lines = extract_feedback_header(lines)
    main_title, subtitle = split_document_title(title)

    # --- TITLE PAGE ---
    doc.add_paragraph()
    doc.add_paragraph()
    add_styled_paragraph(doc, "VELVET ELVES", 'Title', bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(28))
    add_styled_paragraph(doc, "GoDaddy DNS Configuration Guide", 'Subtitle',
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(18))
    doc.add_paragraph()
    add_styled_paragraph(doc, main_title, 'Heading 1', bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(20))
    if subtitle:
        add_styled_paragraph(doc, subtitle, 'Heading 2',
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14))
    doc.add_paragraph()

    if metadata:
        table = doc.add_table(rows=len(metadata), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Light Grid Accent 1'
        for row_index, (label, value) in enumerate(metadata):
            table.rows[row_index].cells[0].text = label
            table.rows[row_index].cells[1].text = value
            for column_index, cell in enumerate(table.rows[row_index].cells):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = BODY_FONT_SIZE
                        if column_index == 0:
                            run.bold = True

    doc.add_page_break()

    # --- CONTENT ---
    render_feedback_body(doc, body_lines)

    save_document(doc, "GODADDY_ROUTE53_CONFIGURATION_GUIDE.docx")


def create_transaction_processing_method_doc():
    """Generate a .docx export of TRANSACTION_PROCESSING_METHOD.md (client audience).

    Mirrors create_transaction_system_doc: reuses the generic markdown renderer
    (extract_feedback_header + render_feedback_body) so the guide gets the same
    title page, metadata table, headings, bullets, and tables as the other
    generated documents.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    lines = read_source_lines("TRANSACTION_PROCESSING_METHOD.md")
    title, metadata, body_lines = extract_feedback_header(lines)
    main_title, subtitle = split_document_title(title)

    # --- TITLE PAGE ---
    doc.add_paragraph()
    doc.add_paragraph()
    add_styled_paragraph(doc, "VELVET ELVES", 'Title', bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(28))
    add_styled_paragraph(doc, "AI-First Transaction Management Platform", 'Subtitle',
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(18))
    doc.add_paragraph()
    add_styled_paragraph(doc, main_title, 'Heading 1', bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(20))
    if subtitle:
        add_styled_paragraph(doc, subtitle, 'Heading 2',
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14))
    doc.add_paragraph()

    if metadata:
        table = doc.add_table(rows=len(metadata), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Light Grid Accent 1'
        for row_index, (label, value) in enumerate(metadata):
            table.rows[row_index].cells[0].text = label
            table.rows[row_index].cells[1].text = value
            for column_index, cell in enumerate(table.rows[row_index].cells):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = BODY_FONT_SIZE
                        if column_index == 0:
                            run.bold = True

    doc.add_page_break()

    # --- CONTENT ---
    render_feedback_body(doc, body_lines)

    save_document(doc, "TRANSACTION_PROCESSING_METHOD.docx")


def create_transaction_automation_guide_doc():
    """Generate a .docx export of TRANSACTION_PROCESSING_AND_AUTOMATION_GUIDE.md
    (client audience: Jake + the testing team).

    Mirrors create_transaction_processing_method_doc: reuses the generic
    markdown renderer (extract_feedback_header + render_feedback_body) so the
    guide gets the same title page, metadata table, headings, bullets, and
    tables as the other generated documents.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    lines = read_source_lines("TRANSACTION_PROCESSING_AND_AUTOMATION_GUIDE.md")
    title, metadata, body_lines = extract_feedback_header(lines)
    main_title, subtitle = split_document_title(title)

    # --- TITLE PAGE ---
    doc.add_paragraph()
    doc.add_paragraph()
    add_styled_paragraph(doc, "VELVET ELVES", 'Title', bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(28))
    add_styled_paragraph(doc, "AI-First Transaction Management Platform", 'Subtitle',
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(18))
    doc.add_paragraph()
    add_styled_paragraph(doc, main_title, 'Heading 1', bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(20))
    if subtitle:
        add_styled_paragraph(doc, subtitle, 'Heading 2',
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14))
    doc.add_paragraph()

    if metadata:
        table = doc.add_table(rows=len(metadata), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Light Grid Accent 1'
        for row_index, (label, value) in enumerate(metadata):
            table.rows[row_index].cells[0].text = label
            table.rows[row_index].cells[1].text = value
            for column_index, cell in enumerate(table.rows[row_index].cells):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = BODY_FONT_SIZE
                        if column_index == 0:
                            run.bold = True

    doc.add_page_break()

    # --- CONTENT ---
    render_feedback_body(doc, body_lines)

    save_document(doc, "TRANSACTION_PROCESSING_AND_AUTOMATION_GUIDE.docx")


def create_gmail_approval_plan_doc():
    """Generate a .docx export of GMAIL_GOOGLE_APPROVAL_PLAN.md (client audience).

    Mirrors create_transaction_processing_method_doc: reuses the generic
    markdown renderer (extract_feedback_header + render_feedback_body) so the
    plan gets the same title page, metadata table, headings, bullets, and
    tables as the other generated documents.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    lines = read_source_lines("GMAIL_GOOGLE_APPROVAL_PLAN.md")
    title, metadata, body_lines = extract_feedback_header(lines)
    main_title, subtitle = split_document_title(title)

    # --- TITLE PAGE ---
    doc.add_paragraph()
    doc.add_paragraph()
    add_styled_paragraph(doc, "VELVET ELVES", 'Title', bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(28))
    add_styled_paragraph(doc, "AI-First Transaction Management Platform", 'Subtitle',
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(18))
    doc.add_paragraph()
    add_styled_paragraph(doc, main_title, 'Heading 1', bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(20))
    if subtitle:
        add_styled_paragraph(doc, subtitle, 'Heading 2',
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14))
    doc.add_paragraph()

    if metadata:
        table = doc.add_table(rows=len(metadata), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Light Grid Accent 1'
        for row_index, (label, value) in enumerate(metadata):
            table.rows[row_index].cells[0].text = label
            table.rows[row_index].cells[1].text = value
            for column_index, cell in enumerate(table.rows[row_index].cells):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = BODY_FONT_SIZE
                        if column_index == 0:
                            run.bold = True

    doc.add_page_break()

    # --- CONTENT ---
    render_feedback_body(doc, body_lines)

    save_document(doc, "GMAIL_GOOGLE_APPROVAL_PLAN.docx")


TESTING_REVIEW_FEATURES = None  # Populated at bottom of file


def _add_plain_paragraph(doc, text, size=Pt(10), bold=False, space_after=Pt(4), space_before=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = size
    run.bold = bold
    p.paragraph_format.space_after = space_after
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    return p


def _add_dash_bullet(doc, text, size=Pt(10)):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.75)
    paragraph.paragraph_format.first_line_indent = Cm(-0.45)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text_run(paragraph, "-  ", size=size)
    add_text_run(paragraph, text, size=size)
    return paragraph


def _add_dot_bullet(doc, text, size=Pt(10)):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(1.6)
    paragraph.paragraph_format.first_line_indent = Cm(-0.45)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text_run(paragraph, "\u2022  ", size=size)
    add_text_run(paragraph, text, size=size)
    return paragraph


def _add_section_label(doc, text, color=RGBColor(0x1F, 0x3A, 0x5F)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    return p


def _render_bullet_group(doc, items):
    """Render a list of (main_text, [sub_bullets]) tuples."""
    for entry in items:
        if isinstance(entry, str):
            _add_dash_bullet(doc, entry)
        else:
            main, subs = entry
            _add_dash_bullet(doc, main)
            for sub in subs:
                _add_dot_bullet(doc, sub)


def _add_feedback_block(doc, lines=8):
    """Add a blank, visually-delimited block where the client can write feedback."""
    _add_section_label(doc, "Feedback", color=RGBColor(0xC0, 0x50, 0x4D))

    hint = doc.add_paragraph()
    hint_run = hint.add_run(
        "Please note below: Status (Pass / Fail / Needs Work), any comments or issues you hit, "
        "and your priority for the improvement ideas above (High / Medium / Low / Skip)."
    )
    hint_run.italic = True
    hint_run.font.size = Pt(9)
    hint_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    hint.paragraph_format.space_after = Pt(4)

    for _ in range(lines):
        blank = doc.add_paragraph()
        underline_run = blank.add_run("_" * 90)
        underline_run.font.size = Pt(10)
        underline_run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
        blank.paragraph_format.space_after = Pt(6)


def _write_testing_review_md(output_filename="FRONTEND_CLIENT_TESTING_REVIEW.md"):
    """Generate a client-friendly Markdown version of the testing review."""
    lines = []
    lines.append("# Velvet Elves — Frontend Client Testing Review")
    lines.append("")
    lines.append("## Features Currently Complete — Client Feedback Requested")
    lines.append("")
    lines.append("**Last Updated:** June 30, 2026  ")
    lines.append("**Test Environment:** http://dev.velvetelves.com/  ")
    lines.append("**Recommended Browsers:** Chrome or Edge (please allow pop-ups and downloads)  ")
    lines.append("**Reviewer:** Client — please fill in the Feedback block under each feature")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## How To Use This Document")
    lines.append("")
    lines.append("### What is in this document")
    lines.append("")
    lines.append("- This document lists every frontend feature that is currently complete and needs your review.")
    lines.append("- Each feature includes the page address, the exact steps to test, the expected result, our ideas for future improvements, and a blank Feedback area for your notes.")
    lines.append("- Features that are still being built (for example placeholder 'Coming Soon' pages) are intentionally left out of this review. They are listed separately in the companion `todo_list.md` so you know what is still on the way.")
    lines.append("")
    lines.append("### How to fill in the Feedback area")
    lines.append("")
    lines.append("- **Status** — write Pass, Fail, or Needs Work after you try the feature.")
    lines.append("- **Comments** — anything you noticed: confusing text, slow actions, wrong results, missing fields, visual issues.")
    lines.append("- **Improvement priority** — for the ideas listed under 'Future Improvement Suggestions', please mark each as High, Medium, Low, or Skip.")
    lines.append("")
    lines.append("### Accounts you will need")
    lines.append("")
    lines.append("- **Agent or Elf** — covers the main day-to-day workflow.")
    lines.append("- **Team Lead or Admin** — needed to see the Delete button on transactions, the Task Templates and Users & Invites pages in Settings, the Deletion Queue on the Documents page, and the Team and Teams pages.")
    lines.append("- **Workspace Owner** — the very first person who registered the brokerage. Required for the Transfer ownership flow and Settings → Delete Organization (schedule deletion).")
    lines.append("- **Invited member** — sign up by clicking an invite-email link (instead of /register). Required for the invite-accept flow and the invitee branch of the onboarding wizard.")
    lines.append("- **Attorney** — loads the attorney workspace (Matters, Releases Queue, State Rules, Recording Calendar).")
    lines.append("- **Client** — a buyer or seller invited to a transaction; loads the 'closing concierge' client workspace at /client/home.")
    lines.append("- **FSBO Customer** — loads the for-sale-by-owner seller workspace at /fsbo.")
    lines.append("- **Vendor** — loads the vendor document portal at /portal/vendor.")
    lines.append("- **Platform admin** (internal Velvet Elves staff only) — required for the /platform/tenants, /platform/advertising, /platform/billing, and /platform/ai-usage pages.")
    lines.append("")
    lines.append("### Suggested order of testing")
    lines.append("")
    lines.append("1. Public pages and sign-in / sign-up (including the new Organization field on /register)")
    lines.append("2. Invite-accept flow (open an invite link as a brand-new user)")
    lines.append("3. Onboarding wizard (test both founder and invitee branches) and the product tour overlay")
    lines.append("4. Standard Agent or Elf workflow (dashboard, new transaction, transactions list, opening a single deal in the AI deal workspace, My Task Queue, Closing Calendar, Clients, Contacts, All Documents)")
    lines.append("5. The Settings hub (avatar menu → Settings): your Profile, Notifications, Email & E-signature connections (needed before AI Email Review can send), Email Templates, My Playbook, Help & Tour, and — for Admins / owners — the Workspace cards (Company, Branding, Users & Invites, Task Templates, Document Templates, Vendor Templates, Team Playbook, Integrations, AI & Automation, Payment Access, Advertising)")
    lines.append("6. Intelligence — AI Email Review at /ai-emails, AI Suggestions, Analytics, and Vendors")
    lines.append("7. Payments — invoices and commission payouts")
    lines.append("8. Team Lead or Admin extras — the Team Overview and Teams pages (sidebar Team group), invite / ownership / deactivate a member, the Oversight group (Communication Audit, Audit Log), and the Admin configuration cards now in the Settings hub (Integrations, AI & Automation, Payment Access, Advertising)")
    lines.append("9. Attorney workspace (Matters, Matter detail, Releases Queue, State Rules, Recording Calendar)")
    lines.append("10. Client, FSBO, and Vendor portals")
    lines.append("11. Public links (milestone viewer, invoice payment, advertise storefront) — no sign-in")
    lines.append("12. Platform admin pages (internal Velvet Elves staff only)")
    lines.append("13. Direct links and error pages")
    lines.append("")
    lines.append("---")
    lines.append("")

    def _render_md_bullets(items, indent=""):
        out = []
        for entry in items:
            if isinstance(entry, str):
                out.append(f"{indent}- {entry}")
            else:
                main, subs = entry
                out.append(f"{indent}- {main}")
                for sub in subs:
                    out.append(f"{indent}  - {sub}")
        return out

    category_order = []
    categories = {}
    for f in TESTING_REVIEW_FEATURES:
        cat = f["category"]
        if cat not in categories:
            categories[cat] = []
            category_order.append(cat)
        categories[cat].append(f)

    for section_index, cat in enumerate(category_order, start=1):
        lines.append(f"## Section {section_index} — {cat}")
        lines.append("")
        for feature in categories[cat]:
            lines.append(f"### {feature['no']}. {feature['feature']}")
            lines.append("")

            lines.append("**Route / Location**")
            lines.append("")
            lines.append(feature["route"])
            lines.append("")

            lines.append("**How To Test**")
            lines.append("")
            lines.extend(_render_md_bullets(feature["how_to_test"]))
            lines.append("")

            lines.append("**Expected Result**")
            lines.append("")
            lines.extend(_render_md_bullets(feature["expected_result"]))
            lines.append("")

            lines.append("**Future Improvement Suggestions**")
            lines.append("")
            lines.extend(_render_md_bullets(feature["future_ideas"]))
            lines.append("")

            lines.append("**Feedback**")
            lines.append("")
            lines.append("_Please note: Status (Pass / Fail / Needs Work), any comments or issues you hit, and your priority for the improvement ideas above (High / Medium / Low / Skip)._")
            lines.append("")
            lines.append("> _Status:_ ")
            lines.append("> ")
            lines.append("> _Comments:_ ")
            lines.append("> ")
            lines.append("> _Improvement priority:_ ")
            lines.append("")
            lines.append("---")
            lines.append("")

    lines.append("## Overall Feedback")
    lines.append("")
    overall_labels = [
        "Biggest usability wins you noticed",
        "Biggest friction points you noticed",
        "Features you would like prioritized next",
        "Additional requests or general notes",
    ]
    for label in overall_labels:
        lines.append(f"### {label}")
        lines.append("")
        lines.append("> ")
        lines.append("> ")
        lines.append("> ")
        lines.append("")

    output_path = BASE_DIR / output_filename
    output_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"Created {output_path.name}")


def create_testing_review_doc():
    """Generate both .docx and .md for client testing review, with per-feature feedback space."""
    _write_testing_review_md()

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # --- TITLE PAGE ---
    doc.add_paragraph()
    doc.add_paragraph()
    add_styled_paragraph(doc, "VELVET ELVES", 'Title', bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(28))
    add_styled_paragraph(doc, "AI-First Transaction Management Platform", 'Subtitle',
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(18))
    doc.add_paragraph()
    add_styled_paragraph(doc, "FRONTEND CLIENT TESTING REVIEW", 'Heading 1', bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(20))
    add_styled_paragraph(doc, "Features Currently Complete — Client Feedback Requested", 'Heading 2',
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(13))
    doc.add_paragraph()

    _add_plain_paragraph(doc, "Last Updated: June 30, 2026",
                         bold=True, space_after=Pt(3))
    _add_plain_paragraph(doc, "Test Environment: http://dev.velvetelves.com/",
                         space_after=Pt(3))
    _add_plain_paragraph(doc, "Recommended Browsers: Chrome or Edge (please allow pop-ups and downloads)",
                         space_after=Pt(3))
    _add_plain_paragraph(doc, "Reviewer: Client — please fill in the Feedback block under each feature",
                         space_after=Pt(6))

    doc.add_page_break()

    # --- HOW TO USE ---
    add_styled_paragraph(doc, "How To Use This Document", 'Heading 1',
                         space_before=Pt(6), space_after=Pt(8))

    _add_section_label(doc, "What is in this document")
    _render_bullet_group(doc, [
        "This document lists every frontend feature that is currently complete and needs your review.",
        "Each feature includes the page address, the exact steps to test, the expected result, our ideas for future improvements, and a blank Feedback area for your notes.",
        "Features that are still being built (for example placeholder 'Coming Soon' pages) are intentionally left out of this review. They are listed separately in the companion todo_list.md so you know what is still on the way.",
    ])

    _add_section_label(doc, "How to fill in the Feedback area")
    _render_bullet_group(doc, [
        "Status — write Pass, Fail, or Needs Work after you try the feature.",
        "Comments — anything you noticed: confusing text, slow actions, wrong results, missing fields, visual issues.",
        "Improvement priority — for the ideas listed under 'Future Improvement Suggestions', please mark each as High, Medium, Low, or Skip.",
    ])

    _add_section_label(doc, "Accounts you will need")
    _render_bullet_group(doc, [
        ("Agent or Elf — covers the main day-to-day workflow.", []),
        ("Team Lead or Admin — needed to see the Delete button on transactions, the Task Templates and Users & Invites pages in Settings, the Deletion Queue on the Documents page, and the Team and Teams pages.", []),
        ("Workspace Owner — the very first person who registered the brokerage. Required for the Transfer ownership flow and Settings → Delete Organization (schedule deletion).", []),
        ("Invited member — sign up by clicking an invite-email link (instead of /register). Required for the invite-accept flow and the invitee branch of the onboarding wizard.", []),
        ("Attorney — loads the attorney workspace (Matters, Releases Queue, State Rules, Recording Calendar).", []),
        ("Client — a buyer or seller invited to a transaction; loads the 'closing concierge' client workspace at /client/home.", []),
        ("FSBO Customer — loads the for-sale-by-owner seller workspace at /fsbo.", []),
        ("Vendor — loads the vendor document portal at /portal/vendor.", []),
        ("Platform admin (internal Velvet Elves staff only) — required for the /platform/tenants, /platform/advertising, /platform/billing, and /platform/ai-usage pages.", []),
    ])

    _add_section_label(doc, "Suggested order of testing")
    _render_bullet_group(doc, [
        "1. Public pages and sign-in / sign-up (including the new Organization field on /register)",
        "2. Invite-accept flow (open an invite link as a brand-new user)",
        "3. Onboarding wizard (test both founder and invitee branches) and the product tour overlay",
        "4. Standard Agent or Elf workflow (dashboard, new transaction, transactions list, opening a single deal in the AI deal workspace, My Task Queue, Closing Calendar, Clients, Contacts, All Documents)",
        "5. The Settings hub (avatar menu → Settings): your Profile, Notifications, Email & E-signature connections (needed before AI Email Review can send), Email Templates, My Playbook, Help & Tour, and — for Admins / owners — the Workspace cards (Company, Branding, Users & Invites, Task Templates, Document Templates, Vendor Templates, Team Playbook, Integrations, AI & Automation, Payment Access, Advertising)",
        "6. Intelligence — AI Email Review at /ai-emails, AI Suggestions, Analytics, and Vendors",
        "7. Payments — invoices and commission payouts",
        "8. Team Lead or Admin extras — the Team Overview and Teams pages (sidebar Team group), invite / ownership / deactivate a member, the Oversight group (Communication Audit, Audit Log), and the Admin configuration cards now in the Settings hub (Integrations, AI & Automation, Payment Access, Advertising)",
        "9. Attorney workspace (Matters, Matter detail, Releases Queue, State Rules, Recording Calendar)",
        "10. Client, FSBO, and Vendor portals",
        "11. Public links (milestone viewer, invoice payment, advertise storefront) — no sign-in",
        "12. Platform admin pages (internal Velvet Elves staff only)",
        "13. Direct links and error pages",
    ])

    doc.add_page_break()

    # --- FEATURES ---
    features = TESTING_REVIEW_FEATURES
    category_order = []
    categories = {}
    for feature in features:
        cat = feature["category"]
        if cat not in categories:
            categories[cat] = []
            category_order.append(cat)
        categories[cat].append(feature)

    for section_index, cat in enumerate(category_order, start=1):
        add_styled_paragraph(doc, f"Section {section_index} — {cat}", 'Heading 1',
                             space_before=Pt(18), space_after=Pt(8))

        for feature in categories[cat]:
            add_styled_paragraph(doc, f"{feature['no']}. {feature['feature']}", 'Heading 2',
                                 space_before=Pt(14), space_after=Pt(4))

            _add_section_label(doc, "Route / Location")
            _add_plain_paragraph(doc, feature["route"], size=Pt(10), space_after=Pt(2))

            _add_section_label(doc, "How To Test")
            _render_bullet_group(doc, feature["how_to_test"])

            _add_section_label(doc, "Expected Result")
            _render_bullet_group(doc, feature["expected_result"])

            _add_section_label(doc, "Future Improvement Suggestions")
            _render_bullet_group(doc, feature["future_ideas"])

            _add_feedback_block(doc, lines=6)

            doc.add_paragraph()

        doc.add_page_break()

    # --- OVERALL FEEDBACK ---
    add_styled_paragraph(doc, "Overall Feedback", 'Heading 1',
                         space_before=Pt(6), space_after=Pt(8))

    overall_labels = [
        "Biggest usability wins you noticed",
        "Biggest friction points you noticed",
        "Features you would like prioritized next",
        "Additional requests or general notes",
    ]
    for label in overall_labels:
        _add_section_label(doc, label)
        for _ in range(6):
            blank = doc.add_paragraph()
            underline_run = blank.add_run("_" * 90)
            underline_run.font.size = Pt(10)
            underline_run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            blank.paragraph_format.space_after = Pt(6)

    save_document(doc, "FRONTEND_CLIENT_TESTING_REVIEW.docx")


TESTING_REVIEW_FEATURES = [
    # ============================================================
    # SECTION 1 — PUBLIC & SIGN-IN PAGES
    # ============================================================
    {
        "no": 1,
        "category": "Public & Sign-In Pages",
        "feature": "Terms of Service page",
        "route": "/terms",
        "how_to_test": [
            ("Open the page in two different ways.", [
                "Open /terms while you are signed out.",
                "Open /terms again while signed in as any role.",
            ]),
            ("Scroll the page from top to bottom and read the legal text.", [
                "Check headings and paragraphs render without missing text.",
                "Confirm any links inside the document open correctly.",
            ]),
        ],
        "expected_result": [
            "The page loads without asking you to sign in.",
            "The page title at the top clearly says Terms of Service.",
            "The full legal content appears in a clean, readable layout.",
        ],
        "future_ideas": [
            "Show a 'Last updated' date at the top so readers know which version they are looking at.",
            "Add a clickable Table of Contents so long sections are easy to jump to.",
            "Offer a 'Save as PDF' button for clients who want to keep a copy.",
        ],
    },
    {
        "no": 2,
        "category": "Public & Sign-In Pages",
        "feature": "Privacy Policy page",
        "route": "/privacy",
        "how_to_test": [
            ("Open /privacy while signed out, then open it again while signed in.", []),
            ("Scroll the whole page and read through the policy sections.", [
                "Make sure nothing is cut off on the right side on a wide screen.",
                "Confirm any in-page links work.",
            ]),
        ],
        "expected_result": [
            "The page loads without any sign-in prompt.",
            "The page title clearly says Privacy Policy.",
            "The content reads cleanly in both signed-out and signed-in states.",
        ],
        "future_ideas": [
            "Add jump links to common sections such as data storage and user rights.",
            "Provide a print-friendly layout so clients can save a paper copy.",
            "Summarize key points at the top in plain language.",
        ],
    },
    {
        "no": 3,
        "category": "Public & Sign-In Pages",
        "feature": "Protected page redirect",
        "route": "/dashboard (while signed out)",
        "how_to_test": [
            ("Sign out, then paste /dashboard directly into the browser address bar and press Enter.", []),
        ],
        "expected_result": [
            "The app should send you straight to /login instead of showing the dashboard.",
            "No protected content (such as transactions) should appear on screen.",
        ],
        "future_ideas": [
            "Remember the page you originally tried to open, and take you there after you sign in.",
            "Show a short explanation on the login screen that says 'Please sign in to continue'.",
        ],
    },
    {
        "no": 4,
        "category": "Public & Sign-In Pages",
        "feature": "Register (sign-up) page",
        "route": "/register",
        "how_to_test": [
            "Open /register and confirm the fields: Full name, Organization (optional), Email, Password, Confirm Password, Phone (optional), and a Terms / Privacy checkbox. A Google sign-up button is also available.",
            "Note: there is no Role dropdown. The first person to register for an organization is automatically the Admin and Owner of a brand-new workspace. To join an existing brokerage you must use an invite link — typing the brokerage name in Organization does not join you to it.",
            "Type an email already in use and confirm the page blocks Create Account and offers a 'Sign in instead' link.",
            "Type a weak password (under 8 characters, no number, no symbol, etc.) and confirm the page shows which rules still need to be met.",
            "Type mismatching passwords and confirm Create Account is blocked until they match.",
            "Leave the Terms / Privacy box unchecked and confirm Create Account is blocked.",
            "Submit a valid registration using an email you can check.",
        ],
        "expected_result": [
            "Each invalid case shows a clear message next to the field, and Create Account stays disabled.",
            "After a successful submission you are either signed in and taken to /onboarding, or taken to /login with a 'please confirm your email' message.",
        ],
        "future_ideas": [
            "Add an eye icon that shows / hides the password while typing.",
            "Show a one-line preview of how the Organization name will appear on outbound emails before you submit.",
            "Add Microsoft / Outlook and Apple sign-up buttons next to Google.",
            "Auto-suggest an Organization name from the email domain (for example 'acme-realty.com' suggests 'Acme Realty').",
        ],
    },
    {
        "no": 5,
        "category": "Public & Sign-In Pages",
        "feature": "Login (sign-in) page",
        "route": "/login",
        "how_to_test": [
            ("Check the page content.", [
                "A Google sign-in button is visible at the top.",
                "Email and password fields are visible.",
                "A 'Forgot password?' link is visible.",
                "A link to the Register page is visible.",
            ]),
            ("Try bad inputs.", [
                "Invalid email format — an error should appear.",
                "Wrong password — a clear error banner should appear.",
            ]),
            ("Sign in with a valid account, then refresh the browser tab.", []),
        ],
        "expected_result": [
            "Users who have not finished onboarding go to /onboarding.",
            "Users who have finished onboarding go to /dashboard.",
            "After a page refresh, you are still signed in.",
        ],
        "future_ideas": [
            "Add a 'Remember me' checkbox for longer sessions.",
            "Show a clear lockout message after several failed attempts.",
            "Add optional two-step verification for extra safety.",
        ],
    },
    {
        "no": 6,
        "category": "Public & Sign-In Pages",
        "feature": "Log out",
        "route": "User menu (top-right avatar)",
        "how_to_test": [
            ("Sign in, then open the avatar menu in the top-right corner.", []),
            ("Click Log Out.", []),
            ("After signing out, try to re-open /dashboard directly in the browser.", []),
        ],
        "expected_result": [
            "You are returned to the /login page right away.",
            "Protected pages cannot be opened again until you sign back in.",
        ],
        "future_ideas": [
            "Add a 'Sign out of all devices' option for extra safety if a laptop is shared or lost.",
            "Briefly confirm the sign-out with a small message ('You have been signed out').",
        ],
    },
    {
        "no": 7,
        "category": "Public & Sign-In Pages",
        "feature": "Forgot password",
        "route": "/forgot-password",
        "how_to_test": [
            ("Verify the screen.", [
                "An email field and a 'Send Reset Link' button should be visible.",
            ]),
            ("Try an invalid email format first, then submit a valid email.", []),
        ],
        "expected_result": [
            "Invalid formats are blocked by a clear message.",
            ("A valid submission switches the page into a success state.", [
                "It shows the email address you entered.",
                "It offers a 'Try a different email' link.",
                "It offers a link back to sign in.",
            ]),
        ],
        "future_ideas": [
            "Show the same success message for known and unknown emails so that no one can fish for registered emails.",
            "Add a 'Resend email' button that is enabled after 60 seconds.",
        ],
    },
    {
        "no": 8,
        "category": "Public & Sign-In Pages",
        "feature": "Reset password",
        "route": "/reset-password (opened from the reset email)",
        "how_to_test": [
            ("Open a real reset link from your email inbox.", []),
            ("Try a weak password and confirm the page blocks it.", []),
            ("Enter a strong password and a matching confirmation, then submit.", []),
            ("Also test opening /reset-password directly with no token, or an expired token.", []),
        ],
        "expected_result": [
            "The reset form loads correctly when the link is valid.",
            "A successful reset shows a success screen and then takes you to /login.",
            "An invalid or expired link shows a clear 'Invalid or expired link' screen with a way to request a new one.",
        ],
        "future_ideas": [
            "Show a password strength meter while typing.",
            "Warn the user that all other signed-in sessions will be logged out after reset.",
        ],
    },
    {
        "no": 9,
        "category": "Public & Sign-In Pages",
        "feature": "Email confirmation",
        "route": "/auth/confirm (opened from the confirmation email)",
        "how_to_test": [
            ("Open a valid confirmation link from your email inbox.", []),
            ("Watch the spinner and wait for the redirect.", []),
            ("Also open a broken / malformed confirmation URL.", []),
        ],
        "expected_result": [
            "A valid link signs you in and takes you to /onboarding (new user) or /dashboard (returning user).",
            "A malformed link shows a clear error screen.",
        ],
        "future_ideas": [
            "Show the confirmed email address on the success screen so the user can double-check the right account was confirmed.",
            "Offer a 'Resend confirmation email' button on the error screen.",
        ],
    },
    {
        "no": 10,
        "category": "Public & Sign-In Pages",
        "feature": "Google sign-in and sign-up",
        "route": "/login and /register",
        "how_to_test": [
            ("Click the Google button on the Login page.", []),
            ("Do the same on the Register page.", []),
            ("Approve Google's consent screen.", []),
        ],
        "expected_result": [
            "The browser goes to Google, then comes back through the app's callback page.",
            "You are signed in and taken to /onboarding (new user) or /dashboard (returning user).",
            "If Google sign-in is not configured in this environment, an error message should appear cleanly rather than breaking the page.",
        ],
        "future_ideas": [
            "Add Microsoft / Outlook and Apple sign-in buttons next to Google.",
            "Show the connected Google account email on the first screen after sign-in.",
        ],
    },
    {
        "no": 11,
        "category": "Public & Sign-In Pages",
        "feature": "Invite acceptance",
        "route": "/invite/<token> (opened from the invite email)",
        "how_to_test": [
            "Open the invite link from your email. Both '/invite/<token>' and '/invite/accept?token=<token>' should work.",
            "Confirm the page shows a 'You're invited!' headline, the role you are joining as (Agent, Transaction Coordinator, Team Lead, Attorney, Client, FSBO Customer, or Vendor), and the invited email.",
            "Fill in Full name (required), Create password (required, at least 8 characters with at least one number), and Phone (optional).",
            "Try an empty name or a weak password and confirm the page blocks Submit.",
            "Click 'Join Velvet Elves' to submit.",
            "Separately, paste a fake token in the browser (e.g. /invite/some-fake-token) and confirm an 'Invalid Invitation' screen appears with a 'Go to login' link.",
        ],
        "expected_result": [
            "A valid invite signs the new user in and takes them to /onboarding.",
            "On /onboarding, the wizard skips the Company / brokerage step and shows a read-only 'Joining: {brokerage name}' line instead.",
            "An invalid or expired token shows the 'Invalid Invitation' screen with a way back to login.",
        ],
        "future_ideas": [
            "Show a countdown for how long the invite is still valid (e.g. 'Expires in 6 hours').",
            "Add a password-strength meter that matches the one on the /register page so invitees see the same five rules.",
            "Show the inviter's name (for example 'Sam Closer invited you to Acme Realty') so the recipient knows it is legitimate.",
        ],
    },

    # ============================================================
    # SECTION 2 — FIRST-TIME USER EXPERIENCE
    # ============================================================
    {
        "no": 12,
        "category": "First-Time User Experience",
        "feature": "Onboarding wizard (role-aware flow)",
        "route": "/onboarding",
        "how_to_test": [
            "Sign up for a fresh account or accept an invite. The wizard should open automatically. An account that already finished onboarding is forwarded straight to /dashboard.",
            "Founder vs. invitee: if you self-registered on /register, Step 2 includes a Company / brokerage field and a Brand logo upload. If you joined via an invite, Step 2 hides those fields and shows a read-only 'Joining: {brokerage name}' line instead.",
            "Step 1 — Welcome: confirm the page greets you by first name and shows a short intro for your role. Internal roles see four value cards; external roles (Client, FSBO, Vendor) see two. Click 'Let's go' to advance.",
            "Step 2 — Your Profile: confirm Full name and Phone (auto-formats as you type) are pre-filled, and a Role dropdown lets you pick your role. Founders also see Company / brokerage and Brand logo (PNG, JPEG, WEBP, SVG, GIF; max 2 MB — upload a wrong type or oversize file and confirm a clear error). Switch the Role dropdown to an external role and confirm the email and e-sign steps disappear from the step list.",
            "Step 3 — Email Inbox (internal roles only): connect Gmail or Outlook by clicking Connect. Confirm a real OAuth popup opens; on success the row flips to Connected. You can also click 'Skip for now' to move on.",
            "Step 4 — E-signature (Agent / TC / Team Lead / Attorney only): connect DocuSign through its OAuth popup, or skip.",
            "Final step — All set: confirm you can click 'Create your first transaction' to open the New Transaction wizard, or 'Go to dashboard' to finish. After finishing, refresh — you should land on /dashboard, not back on /onboarding.",
            "Refresh mid-wizard and confirm your previously-saved fields (name, phone, role, company, logo) are remembered even though the wizard restarts at Step 1.",
        ],
        "expected_result": [
            "Each step shows only the fields that match your role.",
            "Internal roles see 4–5 steps; external roles see 3.",
            "Gmail / Outlook / DocuSign connections persist into Settings → Email & E-signature after onboarding.",
            "Logo files outside the allowed types or larger than 2 MB are rejected with a clear message.",
            "Either final-step button marks onboarding complete on the server and triggers the product tour the next time the dashboard loads.",
        ],
        "future_ideas": [
            "Add a 'Save and finish later' option that remembers the current step (today only the field values are remembered).",
            "Detect popup-blocked browsers up front and offer a fallback redirect-based OAuth path.",
            "Preview the uploaded logo at the exact size it will appear in the app's sidebar.",
            "Let a user preview which steps an external role will see before they switch the role dropdown.",
            "Add an invitee-only 'Welcome from {inviter name}' line so the invitee knows who brought them in.",
        ],
    },
    {
        "no": 13,
        "category": "First-Time User Experience",
        "feature": "Product Tour overlay (role-aware walkthrough)",
        "route": "/dashboard (and any other signed-in page once the tour is started)",
        "how_to_test": [
            "Finish onboarding with a fresh test account, then land on the dashboard. The product tour should start automatically.",
            "If you already finished the tour, replay it from Settings → Help & Tour → Start tour (internal roles); portal roles open it from the avatar menu → Settings → Help & tour.",
            "Step through the whole tour using Next / Back / Skip. Internal roles (Agent, Transaction Coordinator, Team Lead, Admin) see a 9-step tour covering sidebar KPIs, Active Transactions, My Task Queue, All Documents, AI Briefing, search, notifications, and the New Transaction button.",
            "Sign in as an Attorney and replay the tour — it should be a 5-step tour focused on the matter queue, documents, and AI briefing.",
            "Sign in as a Client, FSBO Customer, or Vendor and replay the tour — it should be a 5-step tour focused on My Properties, Documents, and Ask Velvet Elves AI.",
            "Use the keyboard: → or Enter to advance, ← to go back, Esc to skip. Confirm Cmd+K / Ctrl+K still opens global search mid-tour.",
            "Skip the tour mid-way and confirm it does not mark complete (Account → Help & tour → Start tour starts it again from the beginning).",
            "Finish the tour on the final step and confirm it does not auto-start the next time you log in.",
        ],
        "expected_result": [
            "The tour highlights the right element for each step and the tooltip stays on screen.",
            "Internal roles see a 9-step tour; Attorney and external roles see 5-step role-appropriate tours.",
            "Skipping does not lock the tour; only Finish marks it complete.",
            "Account → Help & tour always replays the tour for the role you are signed in as.",
        ],
        "future_ideas": [
            "Add per-feature mini-tours (e.g. 'tour just the Documents page') reachable from inline 'New here?' badges.",
            "Show a one-line tip about what to try after Finish (for example 'Try uploading your first contract').",
            "Add an option to slow down the spotlight animation for users with motion sensitivity.",
            "Track tour completion analytics (where users skip) so we can prioritise improvements.",
        ],
    },

    # ============================================================
    # SECTION 3 — DAILY AGENT / ELF WORKFLOW
    # ============================================================
    {
        "no": 14,
        "category": "Daily Agent / Elf Workflow",
        "feature": "Dashboard home (role-aware)",
        "route": "/dashboard",
        "how_to_test": [
            "Open /dashboard. You are automatically sent to the dashboard built for your role — a solo Agent / Transaction Coordinator, a Team Lead, an Admin, and an Attorney each see a different layout.",
            ("On the Agent / Team Lead dashboard, check the main areas.", [
                "A row of KPI tiles at the top (for example Pending commission, Pipeline volume, Closings this year, Active deals).",
                "An 'Action queue' card listing the most urgent things to do.",
                "A 'Priority transactions' area showing your most important deals as cards.",
                "A 'Portfolio health' and a 'Portfolio intelligence' (AI) card.",
                "A payments snapshot card.",
            ]),
            "Confirm every card either shows real numbers or a clean 'all clear' / empty message — nothing should be blank or broken.",
            "Click the '+ New Transaction' button (top bar and sidebar).",
            "Sign in as a Team Lead and confirm the dashboard shows team-wide numbers; sign in as an Admin and confirm the Admin console layout.",
        ],
        "expected_result": [
            "Each role lands on its own dashboard automatically — you never see another role's layout.",
            "Every card shows real data or a clean empty state.",
            "The New Transaction button opens the transaction wizard.",
        ],
        "future_ideas": [
            "Let the user reorder or hide cards to personalize their landing page.",
            "Add a single 'AI summary of my day' headline card at the top.",
            "Show a small 'this week vs last week' comparison on the KPI tiles.",
        ],
    },
    {
        "no": 15,
        "category": "Daily Agent / Elf Workflow",
        "feature": "Sidebar and top bar",
        "route": "Every signed-in page",
        "how_to_test": [
            ("Check the sidebar.", [
                "Navigation items change depending on the signed-in role (for example an Agent sees Deals, Workflow, Payments, Vendors, and Intelligence groups; a Team Lead and Admin also see a Team group with Team Overview and Teams, and an Admin sees an Oversight group with Communication Audit and Audit Log). All workspace configuration (Task Templates, Document Templates, Vendor Templates, Integrations, AI & Automation, Payment Access, Advertising, and the team playbook) now lives in the Settings hub, not the sidebar. Internal Velvet Elves staff also see a Platform group (Tenants, AI usage, Help center).",
                "KPI tiles in the sidebar show numbers such as overdue tasks, closings this week, active deals, and pipeline value.",
            ]),
            ("Check the top bar.", [
                "Click 'Today's AI Briefing' — a side panel should open (internal roles only).",
                "Click any status chip (Critical / Needs Attention / On Track) — it should filter the transactions list.",
                "Click the search box (or press Ctrl+K / Cmd+K) — a search panel should open and find deals, tasks, documents, and people as you type. Press Enter on a result to jump to it.",
                "Click the bell icon — a notifications panel should open. If you have overdue or upcoming task reminders, a small number badge shows the unread count.",
                "Open the avatar menu — confirm Settings, Help Center, and Log Out (Settings opens the Settings hub for internal roles; Help Center opens the public guides website in a new browser tab).",
                "On a narrow browser window, click the mobile menu icon.",
            ]),
        ],
        "expected_result": [
            "Sidebar navigation and KPIs adjust correctly to the user's role.",
            "The AI Briefing panel opens and closes cleanly.",
            "Status chips take you to the correct filtered transaction view.",
            "Search returns real results and the bell opens a real notifications list with an accurate unread count.",
            "The avatar menu opens the Settings hub (Settings), the public Help Center in a new tab (Help Center), or signs you out (Log Out). The mobile menu behaves correctly.",
        ],
        "future_ideas": [
            "Add a sidebar collapse toggle for users on smaller laptops.",
            "Let the user mark all notifications as read in one click from the panel.",
            "Add recent-searches and saved-search shortcuts to the search panel.",
        ],
    },
    {
        "no": 16,
        "category": "Daily Agent / Elf Workflow",
        "feature": "New Transaction wizard — upload, AI parsing, details, and confirm",
        "route": "Opens from the top bar, the sidebar, and the dashboard 'New Transaction' button",
        "how_to_test": [
            "Start the wizard. At the top, confirm a 'Step X of 5' progress strip with five phases: Upload, Review details, Timeline, Compliance, and Tasks & create.",
            ("Step 1 — Documents.", [
                "First pick 'Who are you representing?' (Buyer / Seller / Buyer & Seller). Until you pick, the upload area is dimmed and Continue stays off.",
                "Upload one PDF, then upload several at once. Remove a file and confirm the count updates.",
                "On a PDF with more than one page, confirm a 'Split' button (it is hidden for single-page PDFs); use it to pick page ranges.",
                "Or click 'Skip upload — enter details manually' to fill everything in by hand.",
            ]),
            ("Step 2 — AI Parsing.", [
                "If you uploaded a document, let AI parsing run; each document shows a confidence bar and the wizard moves on by itself.",
                "If parsing cannot finish, confirm a clear warning and a 'Continue Manually' button.",
            ]),
            ("Step 3 — Address & Contacts.", [
                "Fill in Street, City, State, ZIP (County is optional). The Street field suggests addresses you have used before.",
                "Add each party with 'Add Party'. Name, Email, and Phone are required on every party (marked with a red *).",
                "Confirm empty required fields are gently highlighted, and the highlight clears once you fill them in.",
            ]),
            ("Step 4 — Purchase Information.", [
                "Enter Purchase Price and Earnest Money (both show a $ and add commas as you type), Contract Acceptance Date (cannot be in the future), Closing Date (cannot be in the past), and Possession Date.",
                "Answer 'Is the buyer getting a mortgage?' — choosing Yes reveals a Mortgage Type dropdown.",
                "Answer the Yes/No questions for home inspection, home warranty, and HOA — choosing Yes reveals the matching days / ordered-by field.",
                "Pick who orders title (Buyer or Seller), add any custom contingencies, and add a note (you can pin it to the top of the transaction log).",
                "If you are a Transaction Coordinator, Team Lead, or Admin, confirm a 'Whose transaction is it?' owner picker (an Agent does not see this).",
            ]),
            ("Step 5 — Missing Information (only appears if a required field is still empty).", [
                "Leave one required field empty on purpose to trigger this step; confirm it shows the friendly field name, not a code.",
                "Fill it in by hand, or click 'AI Search' to have the AI look for the value.",
            ]),
            ("Step 6 — Confirm.", [
                "Review the summary: the property, price, key dates, parties, purchase info, and documents (each with an AI confidence pill).",
                "Click any Edit pencil to jump back, make a change, and confirm Continue brings you straight back to this summary.",
            ]),
            "Close the wizard mid-way and confirm a branded 'Discard this transaction?' warning appears (not a plain browser pop-up), with Keep editing / Discard.",
        ],
        "expected_result": [
            "You must choose who you are representing before you can upload.",
            "Continue is blocked until every required field (address, and each party's name + email + phone, and the core purchase fields) is filled.",
            "Required dates are limited sensibly (acceptance not in the future, closing not in the past), and money fields format as you type.",
            "After Confirm, the wizard continues to the Timeline, Compliance, and Review Tasks steps (covered next) before the transaction is created.",
        ],
        "future_ideas": [
            "Let the user start a new deal by duplicating a recent transaction.",
            "Remember a partly-filled wizard across devices, not just in this browser.",
            "Offer a short tips panel for first-time users.",
        ],
    },
    {
        "no": "16.1",
        "category": "Daily Agent / Elf Workflow",
        "feature": "New Transaction wizard — Review Your Timeline",
        "route": "New Transaction wizard, the Timeline phase (after Confirm)",
        "how_to_test": [
            "After Confirm, land on 'Review Your Timeline'. Confirm a card asking 'Does this Date of Acceptance look right?' with the date shown large.",
            "If the date came from a document, click 'View in document' and confirm the document viewer jumps to the page it was read from.",
            "If no acceptance date was found (or you are in manual mode), confirm the card says so honestly and offers a date picker — it never makes up a date.",
            "Confirm the deadline list below is dimmed until you confirm the date. Click 'Looks good' to unlock it.",
            "Confirm every deadline row shows a name, a date, and how it was worked out (for example '5 days after Date of Acceptance'). Core dates like Closing and Possession have no Remove button.",
            "Click 'Edit date' on the acceptance card and pick a different day. Confirm a summary appears ('N deadlines and M task dates moved') with a 'Review changes' list showing old → new dates.",
            "Click '+ Add deadline', give it a name (for example 'Septic inspection'), keep the suggested rule, and save. Confirm the new row appears with the calculated date. Remove it and use the Undo chip.",
            "If your document contained a deal-specific item, confirm an 'AI suggestions' group with cards — each showing the rule, a confidence chip, and a citation link that jumps to the page. Click 'Add' to keep one or 'Skip' to dismiss it.",
            "Click 'Confirm Timeline' to continue.",
        ],
        "expected_result": [
            "The deadline list stays locked until you confirm the acceptance date.",
            "Every deadline shows how its date was calculated; changing the acceptance date moves the others and shows you the change before it sticks.",
            "AI suggestions only appear when the AI can show where it read them, and none are added without your click.",
        ],
        "future_ideas": [
            "Let the user print or share the timeline before creating the deal.",
            "Offer common deadline presets (for example a standard inspection period) for one-click add.",
        ],
    },
    {
        "no": "16.2",
        "category": "Daily Agent / Elf Workflow",
        "feature": "New Transaction wizard — Confirm Your Compliance Checklist",
        "route": "New Transaction wizard, the Compliance phase (after Timeline)",
        "how_to_test": [
            "Confirm a checklist of documents that matches this deal — for example a Closing Disclosure on financed deals, Proof of Funds on cash deals, and HOA or Home Warranty rows only when those terms are true. Each row shows a due date and how it was set.",
            "If you uploaded exactly one document, confirm the Purchase Agreement row already shows green 'Uploaded · <your file>'. With several uploads, click 'Attach document' on a row and pick one of your files (or upload a new one) in the pop-up.",
            "After you attach or upload, confirm an 'AI is checking this document…' chip appears, then settles into a result. Attach a deliberately WRONG file to a row and confirm an amber warning ('AI read this as … — expected …') with 'Keep my type' and 'Remove file' — it warns you but never blocks you.",
            "Click '+ Add document' and confirm a pop-up opens: drop a file or browse (the file is optional), the name fills in from the file, and you can set a due date (none, a relative rule, or a specific date).",
            "Use the pencil to edit a row in place, and the trash to remove one (it shows an Undo chip; removed rows are kept as 'waived', not deleted).",
            "Try 'Use your own checklist' — paste two lines, click 'Read checklist', and confirm both appear checked but not yet added. Uncheck one and click 'Add 1 document' to add only the checked row.",
        ],
        "expected_result": [
            "The checklist matches the transaction type and the contract terms — no rows that do not apply.",
            "Attaching or uploading runs an AI check that warns on a wrong file but never stops you continuing.",
            "Items from your own pasted checklist are only added when you click Add.",
        ],
        "future_ideas": [
            "Let a brokerage save its own standard checklist as the default.",
            "Suggest the most likely document type for each uploaded file automatically.",
        ],
    },
    {
        "no": "16.3",
        "category": "Daily Agent / Elf Workflow",
        "feature": "New Transaction wizard — Review Tasks and create",
        "route": "New Transaction wizard, the Tasks & create phase (the final step)",
        "how_to_test": [
            "Confirm the list of tasks the wizard generated, each with a one-line description, and a 'Search tasks…' box that filters the list as you type.",
            "Click 'Rule…' next to a task's date, set something like '3 days before Closing Date', and save. Confirm the date updates to the worked-out value with a 'rule' chip. Typing a date directly instead clears the rule.",
            "Open a task's 'Related document' dropdown and confirm it lists the checklist rows from the previous step plus 'None'.",
            "Find a task aimed at someone whose email you captured (for example a Buyer task) and confirm an 'Auto-Email off' toggle. Turn it on and confirm the note: 'When this task comes due, I draft the email … for your review. Nothing sends without approval.' A task aimed at 'Agent' shows no toggle at all.",
            "Click 'Create Transaction'. Confirm a spinner and 'Creating…', then a success message with the number of tasks, and that you land on the transactions list with the new deal highlighted.",
        ],
        "expected_result": [
            "A task's date can be set by a relative rule or a specific date, but not both at once.",
            "The Auto-Email toggle only appears on tasks with a real captured email, and it only ever prepares a draft for your review — it never sends on its own.",
            "Creating the transaction generates its tasks and opens the list with the new deal highlighted.",
        ],
        "future_ideas": [
            "Let the user bulk-assign several tasks to a teammate at create time.",
            "Offer task bundles for common deal types.",
        ],
    },
    {
        "no": "16.4",
        "category": "Daily Agent / Elf Workflow",
        "feature": "New Transaction wizard — the AI plan: deal brief, Autopilot, and the command bar",
        "route": "New Transaction wizard (the Timeline, Compliance, Confirm, and Review Tasks steps)",
        "how_to_test": [
            "Above the timeline (and again on the final confirm), read the short deal brief: the price, the financing, the dates that matter, and any 'watch out' lines. Confirm every watch-out has a link into the document, and the facts match what is on the form.",
            "Autopilot — upload a clean, complete packet (all core fields present, every party with a name, email, and phone). After parsing, confirm a single 'Autopilot' banner that lists exactly what will be created and offers 'Looks good' and then 'Approve & Create' — about three clicks in total. Confirm 'Review step by step' is always offered next to it.",
            "Autopilot honesty — upload (or edit to create) a packet where a party has no phone number. Confirm Autopilot does NOT offer the shortcut: it says exactly what is missing and drops you on that step with the field highlighted. Filling it in brings the shortcut back.",
            "The command bar — on the Timeline, Compliance, Confirm, or Review Tasks step, find the 'Tell me what to change' bar and type 'add a septic inspection deadline 14 days after acceptance'. Confirm it shows a preview of what it understood with Apply / Cancel — nothing changes yet. Click Apply, then use Undo.",
            "Type something it cannot do (for example 'make it rain') and confirm it honestly says it cannot map that and lists the kinds of things it can do.",
        ],
        "expected_result": [
            "The deal brief's facts match the confirmed fields, and every watch-out links to the document it came from.",
            "Autopilot only offers the 3-click shortcut when the packet is genuinely complete; otherwise it tells you exactly what is missing.",
            "The command bar always previews a change before applying it, and every change it makes can be undone.",
        ],
        "future_ideas": [
            "Let the command bar handle more than one instruction at once.",
            "Add voice input to the command bar.",
            "Let the user set how complete a packet must be before Autopilot offers the shortcut.",
        ],
    },
    {
        "no": 17,
        "category": "Daily Agent / Elf Workflow",
        "feature": "Active Transactions page",
        "route": "/transactions",
        "how_to_test": [
            ("Check the page header.", [
                "The correct title and the total deal count are shown.",
                "Export CSV, Export Excel, and Print Report buttons are visible.",
            ]),
            ("Use the filter tabs.", [
                "All, Overdue, Due Today, Closing Soon, In Inspection, On Track, Unhealthy.",
                "The list updates as you switch tabs.",
            ]),
            ("Use the sort control.", []),
            ("From the sidebar, switch between Active Transactions, Pending, Closed, and All Transactions.", []),
            ("Click Export CSV and Export Excel — confirm the files download.", []),
            ("Click Print Report — confirm a printable window opens.", []),
            ("Click the floating Ask AI button.", []),
        ],
        "expected_result": [
            "Filter tabs and sidebar filters update the title and the list correctly.",
            "The exports download files named 'transactions.csv' and 'transactions.xls'.",
            "Ask AI opens the AI side panel.",
        ],
        "future_ideas": [
            "Let the user save custom filter views such as 'My hot deals this week'.",
            "Add multi-select and bulk actions on cards (bulk reassign, bulk status change).",
            "Let the user pick which columns to include in the CSV / Excel export.",
        ],
    },

    # ------------------------------------------------------------
    # Transaction Workspace — opening one deal with the AI agent.
    # The agent-centric detail page is the default for Agent / TC /
    # Team Lead / Admin at /transactions/<deal>. These sit right
    # after the Active Transactions list (item 17) because that is
    # where you open a single deal.
    # ------------------------------------------------------------
    {
        "no": "17.1",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Open one deal — the AI deal workspace and its layout",
        "route": "/transactions/<a deal> (click any deal on the Active Transactions list)",
        "how_to_test": [
            "From the Active Transactions list, open any deal. It opens as its own full page with its own web address — not just an expanded card.",
            ("On a normal computer screen, confirm the page has two sides.", [
                "On the left: the Velvet Elves AI assistant for this deal.",
                "On the right: a panel with tabs across the top — Timeline, Compliance, Documents, Tasks, People, Activity, and Email.",
            ]),
            ("Look at the top of the page.", [
                "A 'X% complete' progress bar (shown once the deal has tasks).",
                "A colored status chip — click it to change the deal's status (Active, Incomplete, Paused, Completed, Closed).",
                "A small button (computer screens) that hides or shows the AI assistant so the panel can fill the width.",
            ]),
            "Do something that saves — for example change the status or tick a task — and watch for a short 'Saving…' note at the top while it records.",
            "Open the same deal on a phone or a narrow window. Confirm there is no left/right split: the deal opens straight into the AI assistant, and the tabs are one tap away.",
            "Refresh the page and confirm you stay on the same deal and the same tab.",
        ],
        "expected_result": [
            "Every deal has its own page; a refresh keeps you on the same deal and tab.",
            "The top of the page shows progress, the status (which you can change), and the show/hide assistant button.",
            "A 'Saving…' note confirms the app is recording whatever you clicked.",
            "On a narrow screen the deal opens on the assistant, with the tabs available.",
        ],
        "future_ideas": [
            "Remember which tab you last used on each deal.",
            "Add a quick way to jump to the next or previous deal without going back to the list.",
            "Let the user choose whether the assistant starts open or closed by default.",
        ],
    },
    {
        "no": "17.2",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Talk to the deal's AI assistant (a saved conversation)",
        "route": "The AI assistant (left side on a computer; the 'Agent' tab on a phone)",
        "how_to_test": [
            "In the message box at the bottom, type 'When is the closing date?' and press Enter.",
            "Confirm your message appears immediately as a dark bubble on the right, a short 'thinking' animation shows while the assistant works, then a tidy, well-formatted answer appears.",
            "Reload the page and confirm the whole conversation is still there — it is saved with the deal, not a throwaway chat.",
            "Click 'Clear chat' (the eraser button at the top of the assistant). Confirm a warning explains it removes the conversation but keeps any changes the assistant already made.",
            "Confirm the clear, then reload: the conversation stays empty, but any tasks, dates, or drafts the assistant already applied are still in place on their tabs.",
        ],
        "expected_result": [
            "Your message shows instantly; the reply is correct and easy to read.",
            "The conversation is still there after a page reload.",
            "Clear chat empties the conversation for good but never undoes real work.",
        ],
        "future_ideas": [
            "Let the user search past messages in a long conversation.",
            "Add a voice option to dictate a question (a microphone button is shown but not active yet).",
            "Show suggested follow-up questions after an answer.",
        ],
    },
    {
        "no": "17.3",
        "category": "Daily Agent / Elf Workflow",
        "feature": "AI suggestions — fix a document mismatch in one click",
        "route": "The 'AI suggestions' block at the top of the assistant",
        "how_to_test": [
            "At the top of the assistant, read the 'AI suggestions' block: a one-line summary of what needs you (for example blockers, drafts waiting, items due this week) and a 'Scan' button. Click 'Show all' if there are more than three.",
            "Click 'Scan' to re-check the deal and confirm the suggestion list refreshes.",
            "Set up the test: on the Compliance tab, attach a file to a requirement that is really a different document. The row will show as Uploaded but with a 'mismatch' warning — it looks satisfied, which is the danger.",
            "Back in the assistant, find the suggestion card about the document-type mismatch and click its recommended fix (for example 'Detach and draft a request for the correct document').",
            "Confirm a 'Proposed action' card appears, showing exactly what it will do, with Approve and Dismiss buttons — and that nothing has changed yet.",
            "Click Approve. Confirm the card flips to 'Applied' with a result line and an Undo option.",
            "Open the Compliance tab and confirm the requirement is back to Missing. Open the Email tab and confirm a draft was created as 'Pending review' — nothing was sent.",
            "Back in the assistant, click Undo on the applied card and confirm the change reverses.",
        ],
        "expected_result": [
            "The suggestions are about this specific deal — not generic tips.",
            "The assistant never changes anything without first showing a preview and waiting for your Approve.",
            "Approving applies the change (you can see it on the matching tab); Undo reverses it.",
            "Email actions only ever create a draft for your review — the assistant can never send on its own.",
        ],
        "future_ideas": [
            "Let the user snooze a suggestion for a day.",
            "Add a short 'why this matters' note under each suggestion.",
            "Let the user approve a batch of low-risk suggestions at once.",
        ],
    },
    {
        "no": "17.4",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Point the assistant at any item, and quick commands",
        "route": "Any tab row + the assistant's message box",
        "how_to_test": [
            "On any tab (Tasks, Compliance, Documents, Timeline, People), open a row's three-dot menu (or hover the row) and choose 'Ask AI about this'.",
            "Confirm the item appears as a small tag in the message box and the assistant comes into focus. Type a question and send it — the answer should be about that item.",
            "Click the '+' in the message box to open a picker with tabs (Documents, Tasks, Deadlines, Requirements, People, Emails). Pick one to add it as a tag.",
            "Type '/' in the message box to open the command menu (for example /scan, /readiness, /summarize, /draft-email, /request-document, /add-deadline, /move-date).",
            "Click '/readiness' and confirm the assistant reports whether the deal is ready to close and what is still in the way.",
        ],
        "expected_result": [
            "Any row can be handed to the assistant with the mouse; the tag's name matches the row you picked.",
            "Typing '/' lists the commands, and each one produces a visible result.",
        ],
        "future_ideas": [
            "Let the user tag more than one item at once from the picker.",
            "Add a command to draft a full status update for the client.",
        ],
    },
    {
        "no": "17.5",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Move a key date safely (see what else moves first)",
        "route": "The AI assistant message box",
        "how_to_test": [
            "In the message box, type 'Move the closing date to' and a date about a week later, then send.",
            "Confirm the assistant proposes the change and lists exactly which deadlines and tasks will move — and which will not, with the reason — before anything happens.",
            "Click Approve, open the Timeline tab, and confirm the dates moved.",
            "Back in the assistant, click Undo and confirm the dates go back to what they were.",
        ],
        "expected_result": [
            "A core date never changes without first showing the full list of what else will move.",
            "Undo restores the original dates.",
        ],
        "future_ideas": [
            "Offer a one-click 'Tell the client about the new date' after a move.",
            "Let the user move several dates in one proposal.",
        ],
    },
    {
        "no": "17.6",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Upload a document for instant AI analysis",
        "route": "The paperclip in the assistant message box, or the Upload button on the Documents / Compliance tab",
        "how_to_test": [
            "Click the paperclip in the message box and pick a PDF. Confirm it uploads (a spinner shows), then the assistant posts a short summary of the document and flags anything worth noting.",
            "Try the same from the Documents tab and from the Compliance tab — uploading there should also make the assistant post its analysis in the conversation.",
            "Upload a file that is the wrong type for a requirement and confirm the assistant raises it as a mismatch you can resolve (see item 17.3).",
        ],
        "expected_result": [
            "Uploading a document anywhere on the deal makes the assistant read it and post a plain-language summary in the conversation.",
            "A wrong-type document is caught and offered as something to fix, not silently accepted.",
        ],
        "future_ideas": [
            "Show the document side-by-side with the assistant's summary.",
            "Let the user ask a follow-up question about the document right after it is analyzed.",
        ],
    },
    {
        "no": "17.7",
        "category": "Daily Agent / Elf Workflow",
        "feature": "The deal tabs (Timeline, Compliance, Documents, Tasks, People, Activity)",
        "route": "The tab bar across the top of the right-hand panel",
        "how_to_test": [
            ("Click each tab in turn and confirm it shows this deal's information.", [
                "Timeline — the key dates and deadlines, with 'Add deadline' and a 'Sync' control to push dates to your calendar.",
                "Compliance — the document checklist; attach or upload a document to a requirement, and the AI flags mismatches.",
                "Documents — every document on the deal, with Upload, Manage, Print, and download.",
                "Tasks — the deal's tasks; each task's status is a colored pill (Pending, In progress, Completed, Skipped), with 'Add Task' and a three-dot menu on each task.",
                "People — the buyers, sellers, and other contacts, with 'Manage client access' where your role allows it.",
                "Activity — the deal's history (date changes, status updates, checklist edits) and a Communications button.",
            ]),
            "On the Tasks tab, change a task's status and confirm the pill changes color right away and a 'Saving…' note appears at the top while it records.",
            "Confirm the main button on each tab (Add Task, Add deadline, Upload, Manage) is a colored (orange) button so it is easy to find.",
        ],
        "expected_result": [
            "Every tab shows real information for the open deal — nothing blank or made up.",
            "Changing a status or other field gives instant on-screen feedback, backed by the 'Saving…' note.",
        ],
        "future_ideas": [
            "Show a small count on each tab (for example how many tasks are open).",
            "Let the user reorder the tabs to match how they work.",
        ],
    },
    {
        "no": "17.8",
        "category": "Daily Agent / Elf Workflow",
        "feature": "The Email tab (Outbox / Inbox for this deal)",
        "route": "The Email tab on the right-hand panel",
        "how_to_test": [
            "Open the Email tab and confirm an Outbox / Inbox switch at the top, each showing a count.",
            "On Outbox, confirm any AI drafts for this deal are listed as 'Pending review', with a banner saying nothing sends without your approval. Sent and discarded emails appear below them.",
            "Click Inbox and confirm incoming emails matched to this deal are listed separately.",
            "Click a draft row and confirm it opens the full AI Email Review screen, where you can review, edit, approve, or discard it.",
            "Click Compose to start a new email for this deal.",
        ],
        "expected_result": [
            "Outbox and Inbox are kept separate, each with its own count.",
            "Drafts are clearly marked 'Pending review' — the assistant never sends on its own.",
            "Opening a draft takes you to the AI Email Review screen to finish it.",
        ],
        "future_ideas": [
            "Let the user re-file an email that was matched to the wrong deal.",
            "Show the most recent email at a glance without leaving the tab.",
        ],
    },

    {
        "no": 18,
        "category": "Daily Agent / Elf Workflow",
        "feature": "Transaction card — Tasks area",
        "route": "Expanded card on /transactions",
        "how_to_test": [
            ("Expand a card.", []),
            ("In the Tasks area, click the checkbox on a task to mark it complete, then click again to mark it incomplete.", []),
            ("Open the task status dropdown and change the status.", []),
            ("Click '+ Add' or '+ Add Task'.", []),
        ],
        "expected_result": [
            "The task state updates on screen right away.",
            "The Add Task button opens the Add Task window.",
        ],
        "future_ideas": [
            "Let the user rename a task inline without opening a modal.",
            "Allow drag-and-drop reordering of tasks.",
            "Add a 'Snooze until' option for tasks the user wants to hide temporarily.",
        ],
    },
    {
        "no": 19,
        "category": "Daily Agent / Elf Workflow",
        "feature": "Transaction card — Key Dates area",
        "route": "Expanded card on /transactions",
        "how_to_test": [
            ("Click any key date row.", []),
            ("Change the date in the popover and click Save.", []),
        ],
        "expected_result": [
            "The new date shows up on the card right away.",
            "Overdue dates are visually flagged.",
        ],
        "future_ideas": [
            "When a key date changes, preview which downstream tasks will move.",
            "Offer a one-click 'Notify client of new date' option.",
        ],
    },
    {
        "no": 20,
        "category": "Daily Agent / Elf Workflow",
        "feature": "Transaction card — Contacts area",
        "route": "Expanded card on /transactions",
        "how_to_test": [
            ("Expand a contact row and try the phone and email actions if those fields are filled in.", []),
            ("Click the plus button to add a new contact.", []),
            ("On an empty contact group, click the empty-state add area.", []),
        ],
        "expected_result": [
            "The contact creation window opens with the role already chosen, matching the group you clicked.",
        ],
        "future_ideas": [
            "Show 'Last contacted' date next to each person.",
            "Add a 'Log a call' one-click shortcut that records when you called them.",
        ],
    },
    {
        "no": 21,
        "category": "Daily Agent / Elf Workflow",
        "feature": "Add Task window",
        "route": "Opens from an expanded transaction card",
        "how_to_test": [
            ("Open the Add Task window.", [
                "Leave the task name empty and try to save — it should block you.",
                "Enter a name, pick a Completion Method, pick a Due Date, and pick an Assignee (Myself or AI Agent).",
            ]),
            ("Click 'Get AI Suggestions on How to Complete'.", []),
            ("Apply one of the AI suggestions, then submit the task.", []),
        ],
        "expected_result": [
            "The task is created and the window closes.",
            "Applying an AI suggestion fills in the Completion Method when the suggestion provides one.",
        ],
        "future_ideas": [
            "Add a dropdown of common task templates for one-click creation.",
            "Add a 'Recurring task' option (weekly, monthly).",
            "Allow attaching a document while creating the task.",
        ],
    },
    {
        "no": 22,
        "category": "Daily Agent / Elf Workflow",
        "feature": "Add Contact window",
        "route": "Opens from an expanded transaction card",
        "how_to_test": [
            ("Open the Add Contact window from different groups (Buyer, Lender, Title, etc.).", [
                "Confirm the role label matches the group you opened it from.",
                "Confirm Company Name appears for Lender and Title groups.",
            ]),
            ("Enter First Name, Last Name, Phone, Email, and optional Company.", []),
            ("Submit the form.", []),
        ],
        "expected_result": [
            "The new contact is added and the window closes.",
        ],
        "future_ideas": [
            "Offer 'Pick from existing contacts' so vendors used on other deals can be reused without retyping.",
            "Auto-format phone numbers while typing.",
        ],
    },
    {
        "no": 23,
        "category": "Daily Agent / Elf Workflow",
        "feature": "Documents window on a transaction card",
        "route": "Opens from the 'View / Add Documents' button on an expanded transaction card",
        "how_to_test": [
            "Expand a transaction card on /transactions/active and click View / Add Documents.",
            "Confirm the documents panel opens and lists every document on that transaction.",
            "Confirm each row shows the document name, a status badge (Signed / Awaiting / Flagged / Uploaded), the version, the upload date, and the size.",
            "Click Download on a row. Confirm the file opens in a new browser tab without being blocked by your pop-up blocker, even if the panel has been open for a while.",
            "Click the document name. Confirm it also opens the download in a new tab.",
            "Click Add Document and upload a new file. Confirm the new file appears in the list with its version and upload date.",
            "Click the Email icon on a row (internal users only). Confirm the Email Document window opens above the panel and is clickable.",
            "Click the Version history icon. Confirm the Version history panel opens above the panel and lists every version.",
            "Open the three-dot menu on any row. As an internal user you should see Rename / Classify, Upload new version, and Archive. As a Client / FSBO / Vendor you should see Flag for deletion (or Flagged if you already flagged it).",
            "Click Archive (internal) and confirm a confirmation window opens. Click Confirm and confirm the document leaves the list.",
            "Click Flag for deletion (Client / FSBO / Vendor) and submit a reason. Confirm the row updates to show Flagged.",
            "Press the Escape key — the panel should close.",
        ],
        "expected_result": [
            "Every action (Download, Email, Version history, Rename, Archive, Flag) opens its window or finishes its action visibly above the panel — no clicks land on something hidden behind the backdrop.",
            "Download is not blocked by the browser's pop-up blocker.",
            "Internal users can rename, classify, archive, and email a document from this panel. External users can flag a document for deletion.",
            "Escape closes the panel.",
        ],
        "future_ideas": [
            "Add a Preview window so PDFs render inline instead of opening in a new tab.",
            "Add Send for Signature, Refresh status, and Void envelope buttons here so internal users can finish the signature flow without leaving the transaction card.",
            "Add a status filter (All / Uploaded / Awaiting / Signed / Flagged) and a sort menu to make long document lists easier to scan.",
            "Add an internal Mark for follow-up flag so the agent can mark a document they need to come back to.",
            "Offer a 'Download all as zip' button for the full packet.",
        ],
    },
    {
        "no": 24,
        "category": "Daily Agent / Elf Workflow",
        "feature": "Transaction history panel",
        "route": "Opens from an expanded transaction card",
        "how_to_test": [
            ("Open the History panel from a card.", [
                "Check that it slides in from the right.",
            ]),
            ("Type in the search box and confirm the list filters.", []),
            ("Confirm events are grouped by date (Today, Yesterday, etc.).", []),
        ],
        "expected_result": [
            "Matching events stay visible and the rest are filtered out.",
        ],
        "future_ideas": [
            "Add filter chips by event type (emails, tasks, documents, AI flags).",
            "Add an 'Export this timeline as PDF' button for client hand-offs.",
        ],
    },
    {
        "no": 25,
        "category": "Daily Agent / Elf Workflow",
        "feature": "Print and AI actions on a transaction card",
        "route": "Expanded card on /transactions",
        "how_to_test": [
            ("Click Print on a card.", []),
            ("Click the next-step suggestion or any AI chip on a card.", []),
        ],
        "expected_result": [
            "Print opens a printable closing-checklist window.",
            "AI chips open the AI side panel.",
        ],
        "future_ideas": [
            "Let the user pick a print template (full checklist, one-pager, client summary).",
            "Let the user decide which sections to include in the printout.",
        ],
    },
    {
        "no": 26,
        "category": "Daily Agent / Elf Workflow",
        "feature": "Velvet Elves AI chat panel",
        "route": "Top bar 'Today's AI Briefing', floating Ask AI button, and AI chips on cards",
        "how_to_test": [
            ("Open and close the panel from each of the three entry points.", []),
            ("Click several of the suggested-action chips at the bottom.", []),
            ("Type your own message such as 'Summarize the Smith deal' and click Send.", []),
            ("Try a deal-specific question with a transaction card expanded so the AI picks up that transaction as context.", []),
            ("Try the message actions on a returned reply — Copy, Edit, Delete, and Regenerate.", []),
        ],
        "expected_result": [
            "Each Send produces a real AI reply (the panel is connected to a live AI service).",
            "Suggested chips fill the input with a useful starting prompt.",
            "Copy, Edit, Delete, and Regenerate all work on existing messages.",
        ],
        "future_ideas": [
            "Save past conversations so users can come back later.",
            "Let users 'pin' a transaction so the AI always uses it as context.",
            "Stream answers word-by-word instead of waiting for the full response.",
        ],
    },
    {
        "no": "26.1",
        "category": "Daily Agent / Elf Workflow",
        "feature": "My Task Queue",
        "route": "/tasks/queue (sidebar → Workflow → My Task Queue)",
        "how_to_test": [
            "Open My Task Queue from the sidebar. Confirm a 'Today's AI briefing' strip at the top and a 'Today's progress' summary.",
            "Use the type tabs (All / Documents / Communication / Milestones / Admin) and the Sort menu (Priority / Due date / Transaction / Task type). Confirm the list updates.",
            "Type in the search box ('Search tasks, deals, contacts…') and confirm the list filters.",
            "Tick a task's checkbox to complete it and confirm it moves to 'Completed today'; untick to bring it back.",
            "Click Add task, fill in the task, and confirm it appears in the list.",
            "Click the floating Ask AI button and confirm the AI panel opens.",
        ],
        "expected_result": [
            "The queue gathers your tasks across every deal in one place.",
            "Tabs, sort, and search all narrow or reorder the list correctly.",
            "Completing, adding, and AI assist all work.",
        ],
        "future_ideas": [
            "Add a 'snooze until' option to hide a task for a day.",
            "Allow drag-to-reorder within a priority group.",
            "Add bulk-complete for several tasks at once.",
        ],
    },
    {
        "no": "26.2",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Closing Calendar",
        "route": "/calendar (sidebar → Workflow → Closing Calendar)",
        "how_to_test": [
            "Open Closing Calendar. Confirm a month view of your transaction key dates and closings, with a way to switch to an agenda (list) view.",
            "Move between months and confirm events land on the right dates.",
            "Click an event and confirm it takes you to the matching transaction with that card opened.",
            "Use the 'Connect calendar' / sync controls to connect Google or Outlook (a sign-in popup opens), then use Sync to push your closings. Disconnect and confirm it stops syncing.",
        ],
        "expected_result": [
            "The calendar is built from your real transaction dates — nothing made up.",
            "Events deep-link into the right transaction.",
            "Google / Outlook calendar connect, sync, and disconnect all work.",
        ],
        "future_ideas": [
            "Add a week view alongside month and agenda.",
            "Color-code events by deal health.",
            "Add an .ics download for a single closing.",
        ],
    },
    {
        "no": "26.3",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Clients hub",
        "route": "/clients (sidebar → Deals → Clients)",
        "how_to_test": [
            "Open Clients from the sidebar. Confirm one row per represented client showing their deals and two 'needs me' signals: an unanswered client question and uploads awaiting review.",
            "Click a client's action (for example the unanswered-question signal) and confirm it takes you to the right transaction with the client Q&A drawer or client-access modal open.",
            "Use the phone / email shortcuts on a client row.",
        ],
        "expected_result": [
            "Every represented client is listed once with their deals and what needs your attention.",
            "Actions deep-link to the matching transaction and open the right drawer.",
        ],
        "future_ideas": [
            "Add a search box and a filter for 'only clients who need me'.",
            "Show the last time you contacted each client.",
        ],
    },
    {
        "no": "26.4",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Contacts directory",
        "route": "/contacts (sidebar / global search)",
        "how_to_test": [
            "Open Contacts. Confirm a searchable list of people (co-agents, loan officers, title reps, and other contacts) with their type, email, and phone.",
            "Type in the search box and confirm the list filters by name, email, or company.",
            "Use the type filter to narrow to one kind of contact.",
            "Confirm preferred contacts are marked (star) and vendor contacts are indicated.",
        ],
        "expected_result": [
            "Contacts are searchable and filterable by type.",
            "Email and phone shortcuts work; preferred and vendor contacts are clearly marked.",
        ],
        "future_ideas": [
            "Add an 'add contact' button directly on this page.",
            "Show which transactions each contact is attached to.",
        ],
    },
    {
        "no": "26.5",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Notifications center",
        "route": "/notifications (and the top-bar bell)",
        "how_to_test": [
            "Open the bell in the top bar, then open the full Notifications page.",
            "Confirm filter tabs (All / Overdue / Today / Tomorrow / Upcoming) and that each notification shows its urgency.",
            "Click a notification and confirm it takes you to the related task or transaction.",
            "Confirm reading notifications updates the unread count on the bell.",
        ],
        "expected_result": [
            "Notifications are real task reminders grouped by urgency.",
            "Opening one navigates to the right place and the unread count stays accurate.",
        ],
        "future_ideas": [
            "Add a 'mark all as read' button.",
            "Add notification types beyond task reminders (new client message, signature completed).",
        ],
    },
    {
        "no": "26.6",
        "category": "Daily Agent / Elf Workflow",
        "feature": "AI Suggestions (intelligence inbox)",
        "route": "/ai-suggestions (sidebar → Intelligence → AI Suggestions)",
        "how_to_test": [
            "Open AI Suggestions. Confirm a briefing hero, activity stats, and a row of category filter pills.",
            "Expand a suggestion card and confirm it shows the reason, the AI recommendation, an editable draft where relevant, and an action row.",
            "Apply an action on a card (for example send a draft, add a task, or dismiss) and confirm it does something real and the card updates.",
            "Use the category pills to filter the suggestions.",
        ],
        "expected_result": [
            "Every suggestion is based on real data about your deals — no generic tips.",
            "Each action runs a real, role-appropriate operation.",
        ],
        "future_ideas": [
            "Add a 'snooze this suggestion' option.",
            "Let the user thumbs-down a suggestion to tune what appears.",
        ],
    },
    {
        "no": "26.7",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Analytics (Reports)",
        "route": "/reports (sidebar → Intelligence → Analytics)",
        "how_to_test": [
            "Open Analytics. Confirm a row of KPI tiles (for example commission, transaction volume, on-time close rate) followed by charts (GCI by month, transaction volume, active pipeline, days-to-close, task completion, and so on).",
            "Confirm every chart shows your real numbers or a clean 'not enough data yet' message.",
            "Open the goals editor, set a goal, save it, and confirm the relevant tile reflects progress toward it.",
            "Click Export / Download and confirm a file downloads.",
        ],
        "expected_result": [
            "All tiles and charts are driven by your real transactions.",
            "Goals save and show progress; export downloads a file.",
        ],
        "future_ideas": [
            "Let the user pick the date range / compare two periods.",
            "Add a team-vs-me toggle for Team Leads.",
        ],
    },
    {
        "no": "26.8",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Invoices & Payments",
        "route": "/payments (sidebar → Payments → Invoices & Payments)",
        "how_to_test": [
            "Open Invoices & Payments. Confirm tabs: All Invoices, Open, Paid, Drafts, Void, and Payments.",
            "Switch tabs and confirm the list narrows; use the search box to find an invoice.",
            "Click New invoice, fill it in (transaction, amount, line items), and save a draft, then open it.",
            "Open an invoice to view its detail window; if you have permission, send it and confirm the client receives a pay link.",
            "Open the Payments tab and confirm recorded payments are listed.",
            "Sign in as a role without invoice permission and confirm the create / refund actions are not offered (read-only history is still visible).",
        ],
        "expected_result": [
            "Invoices and payments are real and filterable by status.",
            "Creating, sending, and viewing invoices works; permission gating hides actions a role cannot perform.",
        ],
        "future_ideas": [
            "Add recurring invoices.",
            "Add a one-click reminder for overdue invoices.",
        ],
    },
    {
        "no": "26.9",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Commission Payouts",
        "route": "/payments/payouts (sidebar → Payments → Commission Payouts)",
        "how_to_test": [
            "Open Commission Payouts (shown when your role / plan allows payouts).",
            "Confirm a list of payouts with amount and status; use the search box to find one.",
            "Click New payout, pick a transaction with the typeahead, enter the amount, and create it.",
            "Open a payout to view its detail window.",
        ],
        "expected_result": [
            "Payouts are listed with accurate amounts and statuses.",
            "Creating and viewing a payout works; the action is gated by payment permission.",
        ],
        "future_ideas": [
            "Add a split-commission helper for co-brokered deals.",
            "Export payouts to CSV for accounting.",
        ],
    },
    {
        "no": 27,
        "category": "Daily Agent / Elf Workflow",
        "feature": "All Documents page — open and access",
        "route": "/documents (also reachable as /documents/all)",
        "how_to_test": [
            "Open /documents while signed in as Agent, Transaction Coordinator, Team Lead, Attorney, or Admin.",
            "Try opening /documents while signed in as a Client, For-Sale-By-Owner Customer, or Vendor.",
            "From any page, press Cmd+K (Mac) or Ctrl+K (Windows), search for a document by name, and press Enter on the result.",
        ],
        "expected_result": [
            "Internal roles see the full All Documents page.",
            "Non-internal roles are sent back to their dashboard instead.",
            "The Cmd+K result opens /documents with the right document highlighted and scrolled into view.",
        ],
        "future_ideas": [
            "Show non-internal users a one-line note explaining where they can find documents inside their assigned transactions.",
            "Remember the last filter and sort the user picked between visits.",
        ],
    },
    {
        "no": "27.1",
        "category": "Daily Agent / Elf Workflow",
        "feature": "All Documents — AI Priority queue and Today's Priority card",
        "route": "/documents (default AI Priority tab)",
        "how_to_test": [
            "Open /documents and stay on the AI Priority tab.",
            "Read the Today's Priority card at the top — it should name the most urgent item with a short reason and one or two recommended actions.",
            "Read the AI Briefing strip just below it — a one-sentence summary plus a button that jumps to whichever tab is most relevant right now.",
            "Scroll down the priority list. Each row names the document or missing item, the transaction it belongs to, and the action buttons available.",
        ],
        "expected_result": [
            "The hero and list highlight whatever is most likely to delay a deal.",
            "Clicking the AI Briefing button switches to the matching tab (for example Missing or Sent for Signature).",
        ],
        "future_ideas": [
            "Show a short reason next to each priority row ('Closing in 3 days', 'Awaiting buyer signature for 5 days').",
            "Allow snoozing an item for a day so it drops out of the list without being marked complete.",
        ],
    },
    {
        "no": "27.2",
        "category": "Daily Agent / Elf Workflow",
        "feature": "All Documents — filter tabs and sort",
        "route": "/documents (tabs under the page title)",
        "how_to_test": [
            "Click each filter tab in turn: AI Priority, All Docs, Missing, Pending Review, Sent for Signature, Signed.",
            "Confirm the list narrows to match the chosen tab.",
            "Open the Sort menu and pick each option in turn: AI Impact, Close Date, Document Name, Recently Updated, Last Touched.",
            "Refresh the page after changing tab or sort.",
        ],
        "expected_result": [
            "Every tab shows only the documents that match its name.",
            "Sort instantly reorders the visible list.",
            "Tab and sort choices stay in place after a refresh (they are saved in the URL).",
        ],
        "future_ideas": [
            "Let the user save a favourite tab + sort combination as a one-click shortcut.",
            "Add a 'Group by document type' option alongside the existing sort options.",
        ],
    },
    {
        "no": "27.3",
        "category": "Daily Agent / Elf Workflow",
        "feature": "All Documents — actions on a priority row",
        "route": "/documents (any row in the AI Priority list)",
        "how_to_test": [
            "Click Request on a missing-document row. Confirm the email modal opens with the right recipient and template, and that the row stays in the queue with a 'Requested today' note after Send.",
            "Click Nudge on a row that already had a request. Confirm a shorter follow-up email is drafted and the row updates with a 'Nudged' note.",
            "Click Upload on a missing-document row. Confirm the Upload Document modal opens with the transaction and document type already filled in, and that uploading clears the row.",
            "Find a missing item that does not have a template ready (for example Appraisal Report). Confirm the Generate button is NOT shown — only Upload / Request / Call appear.",
            "Find a missing item that does have a template ready (for example Lead-Based Paint Disclosure). Click Generate.",
            "If the transaction has all the information the template needs, confirm a draft document is created and the preview opens.",
            "If the transaction is missing information (for example the seller name), confirm an 'Almost ready' window appears listing each missing piece in plain English and a Fix button next to each one.",
            "Click Fix on one of the missing pieces. Confirm you are taken straight to the transaction page with the right field highlighted. Fill it in.",
            "Come back to /documents. Confirm the Generate action retries on its own and either opens the draft or shows the next missing piece. You should not have to click Generate a second time.",
            "Trigger Generate on a requirement that has no template at all (rare). Confirm the window now offers real next-step buttons — Upload draft manually or Request from counter-party — instead of a dead 'Got it' button.",
            "Click Approve on a Pending Review row. Confirm the row leaves the queue and appears in the Cleared Today strip.",
            "Click Mark N/A on any row. Confirm a confirmation appears, the row disappears, and an Undo option is available on the Cleared Today card.",
            "Click Flag on any row. Confirm a flag icon appears next to the item and is still there after refresh.",
            "Click Call on a row with a phone number — confirm your phone app opens and the activity is logged on the transaction.",
            "Click Forward on a signed document. Confirm a forward-style email modal opens with the document attached.",
        ],
        "expected_result": [
            "Every action button does something real (sends an email, opens the right modal, or updates the queue).",
            "Request, Nudge, Call, and Forward keep the row visible with a 'last touched' note until the requirement is actually resolved.",
            "Mark N/A, Approve, Upload, and Generate clear the row and add it to the Cleared Today strip at the bottom of the page.",
            "Generate is only offered when the system actually has a template for that requirement.",
            "The missing-fields window never shows raw machine names — every line is plain English with a Fix button.",
        ],
        "future_ideas": [
            "Show 'Next follow-up due in N hours' on touched rows so the user knows when to nudge again.",
            "Add a 'Skip this field' option on the missing-fields window for rare cases where a value is genuinely unavailable.",
        ],
    },
    {
        "no": "27.4",
        "category": "Daily Agent / Elf Workflow",
        "feature": "All Documents — Cleared Today strip",
        "route": "/documents (strip at the bottom of the page)",
        "how_to_test": [
            "Read the Cleared Today header. There should be a one-line description telling you what the strip is — items resolved in the last 24 hours.",
            "Click the information icon next to the header. A short pop-up should explain what each kind of card means (Signed, Approved, Marked N/A, Generated, Replaced, Voided, Uploaded, Reassigned).",
            "Confirm the strip is visible on every tab (AI Priority, Missing, All Docs, Pending Review, Sent for Signature, Signed) — not only on AI Priority.",
            "Resolve a few items different ways (Approve a review, Mark something N/A, Upload a missing document, sign an envelope, generate a draft). Confirm each one shows up as a card on the strip with the matching label.",
            "Send a Request or Nudge from a row. Confirm those touches do NOT show up here — only true resolutions appear on the strip.",
            "On any card, click View Details. Confirm a small read-only window opens with the document name, the transaction, the badge, who cleared it, and when. (No more single whole-card click — every card has clear buttons.)",
            "On a card that points to a real document, click Open. Confirm the document preview opens.",
            "On a Mark N/A card (which does not have a document attached), click Open. Confirm the transaction's Documents view opens — not a dead window.",
            "On an Approved card, click Undo. Confirm the document goes back to Pending Review.",
            "On a Mark N/A card, click Undo. Confirm the row returns to the Missing list.",
            "On a Generated card, click Undo and confirm the draft is removed and the Missing row comes back.",
            "On a Signed card, confirm Undo is not offered, and a short explanation tells you to void the envelope from the document row instead.",
            "Use the filter buttons at the top of the strip: All, Me, Team. Confirm Me shows only items you cleared, Team shows your teammates' clears, and All shows both.",
            "Click 'View all cleared (last 7 days)'. Confirm a panel slides out listing every clear from the past week with the same filter buttons. Scroll down — more rows should load as you reach the bottom.",
            "If you are signed in as a solo agent with no teammates, confirm the Team filter button is not shown.",
        ],
        "expected_result": [
            "The strip explains itself — a new user can tell what it is and what each card means without help.",
            "Cards never use a single 'click the whole card' target. Every card has a View Details button, an Open button, and an Undo where it makes sense.",
            "Only true resolutions appear on the strip — requests, nudges, calls, and forwards do not.",
            "Undo works for Approve, Mark N/A, Generated, and Uploaded clears. Signed, Replaced, and Voided clears explain why they cannot be undone here.",
            "Switching to Me filters the strip to your own clears; Team shows your teammates'.",
            "The 7-day panel shows the same items beyond the last 24 hours, with the same filter buttons.",
        ],
        "future_ideas": [
            "Show a small badge on the Cleared Today header when new items land while you're on a different tab.",
            "Group the 7-day panel by day so it reads like a timeline.",
            "Let the user export the 7-day list to CSV for end-of-week reporting.",
        ],
    },
    {
        "no": "27.5",
        "category": "Daily Agent / Elf Workflow",
        "feature": "All Documents — Upload Document modal",
        "route": "Opens from the 'Upload Document' header button or from 'Upload' on a missing-document row",
        "how_to_test": [
            "Open the Upload modal from the page header button.",
            "Open it again from the Upload action on a missing-document row and confirm the transaction and document type are already filled in.",
            "Drag a PDF onto the drop zone, or browse and pick a file.",
            "Try a file over 20 MB or an unsupported file type (.zip, .xlsx) and confirm the modal shows a clear error.",
            "Click Upload Document to save.",
        ],
        "expected_result": [
            "Allowed file types: PDF, DOC, DOCX, JPG, PNG, WEBP, GIF, TXT, up to 20 MB.",
            "After upload, the new file appears in the document list and resolves the matching missing item if any.",
        ],
        "future_ideas": [
            "Auto-suggest the document type by reading the file contents on upload.",
            "Let the user drag and drop several files at once.",
        ],
    },
    {
        "no": "27.6",
        "category": "Daily Agent / Elf Workflow",
        "feature": "All Documents — Preview and Download",
        "route": "Preview (eye) and Download icons on any document row",
        "how_to_test": [
            "Click the Preview (eye) icon on a PDF or image — the document should open inside the window.",
            "Click Preview on a non-previewable file (for example a .docx) — the window should offer a Download button instead.",
            "Click the Download icon on a row. The document should open in a new browser tab.",
            "Confirm Download is not blocked by your browser's pop-up blocker, even after the page has been open for a while.",
            "From inside the Preview window, click Send for Signature.",
        ],
        "expected_result": [
            "Preview opens the document inside the page so you can read it without leaving.",
            "Download opens the document in a new tab and is never blocked by the pop-up blocker.",
            "Send for Signature from Preview opens the signature modal with the document already selected.",
        ],
        "future_ideas": [
            "Allow side-by-side comparison of two versions in the preview window.",
            "Remember the last zoom level between previews.",
        ],
    },
    {
        "no": "27.7",
        "category": "Daily Agent / Elf Workflow",
        "feature": "All Documents — Send for Signature (DocuSign)",
        "route": "'Send for Signature' header button, the row's signature icon, or the Preview footer",
        "how_to_test": [
            "Open the Send for Signature modal from each of the three entry points (header button, row icon, preview footer).",
            "If your DocuSign account is not yet connected, click the Connect DocuSign button on the banner and complete the popup sign-in.",
            "Add signers using the suggested-party chips (buyer, listing agent, etc.) and / or the 'Add signer' button. Confirm you cannot send until each row has a name and a valid email.",
            "Edit the subject and message that will go to the signers.",
            "Click Send for Signature.",
        ],
        "expected_result": [
            "After Send, a toast confirms the envelope was sent and the row updates to 'Sent for Sig.' with an 'Awaiting: {names}' note.",
            "Recipients receive a DocuSign email immediately.",
            "If sending fails, the error appears both as a toast and inside the modal so the user does not need to reopen it.",
        ],
        "future_ideas": [
            "Let the sender choose sequential vs. parallel signing order.",
            "Show live signature progress inside the modal so the user does not have to refresh the row.",
        ],
    },
    {
        "no": "27.8",
        "category": "Daily Agent / Elf Workflow",
        "feature": "All Documents — manage envelopes already sent for signature (Refresh, Void, Resend)",
        "route": "Rows whose status is Sent for Sig., Declined, or Voided",
        "how_to_test": [
            "Send a document for signature but do not let recipients sign yet. Confirm the row shows the 'Awaiting: …' line listing the recipients.",
            "Click Refresh on the row and confirm a toast reports the current signature status.",
            "Open the row's three-dot menu and click Void Envelope. Confirm a toast appears and the row switches to a 'Voided' state with a Resend option.",
            "Simulate a declined signer in DocuSign (Decline) and refresh — the row should switch to 'Declined' and offer Resend.",
            "Click Resend for Signature from a voided or declined row and confirm the Send for Signature modal reopens with the same document.",
        ],
        "expected_result": [
            "The agent can see who has not signed yet without opening DocuSign.",
            "Refresh updates the row from DocuSign on demand.",
            "Voided and declined envelopes always offer a one-click Resend so a stuck file is never a dead end.",
        ],
        "future_ideas": [
            "When voiding, prompt the user for a short reason that gets passed to DocuSign.",
            "Add a single 'Force-sync all' button that refreshes every in-flight envelope at once.",
        ],
    },
    {
        "no": "27.9",
        "category": "Daily Agent / Elf Workflow",
        "feature": "All Documents — Email Document, Edit Document (Rename / Reclassify / Reassign), Version History",
        "route": "Three-dot menu on any document row (internal roles only). Edit Document is also reachable from the AI Priority queue rows and the priority detail window.",
        "how_to_test": [
            "Open the three-dot menu and click Email Document. Add at least one recipient and click Send Email. Confirm a toast reports the email is queued.",
            "Click Edit Document. The window should show two groups of fields: Identity (file name, label, document type) and Assignment (the transaction this document belongs to).",
            "Try to save with an empty file name and confirm it is blocked.",
            "Change the file name, label, or type and click Save. Confirm the row updates immediately.",
            "Open Edit Document again. Use the Transaction picker to move the document to a different transaction you have access to. A short note should warn you that moving the document reattaches its history and version chain.",
            "Save. Confirm the document leaves the original transaction's list and now appears under the new transaction.",
            "If the document was satisfying a missing requirement on the new transaction, confirm a 'Reassigned' card appears in Cleared Today.",
            "Try to reassign a document that has an envelope already sent for signature, or one that is already signed. Confirm the save is blocked with a clear message asking you to void the envelope first.",
            "Open Edit Document from a row in the AI Priority list (not just from the Three-dot menu on a document row). Confirm it works the same way.",
            "Click Version History. Confirm every version is listed with v1, v2, … and that downloading any historical version still works. Upload a replacement and confirm the latest version is marked Current.",
        ],
        "expected_result": [
            "Email Document queues the email and records it in the transaction's communication history.",
            "Edit Document handles renaming, reclassifying, and reassigning a document in one window. The label is the same everywhere — there is no separate 'Rename' or 'Reassign' button.",
            "Reassigning the document moves it between transactions and shows up in the audit trail of both transactions.",
            "An envelope-in-flight or signed document cannot be reassigned, and the system tells you exactly why.",
            "Uploading a new version moves the previous one to Legacy.",
        ],
        "future_ideas": [
            "Offer saved email templates ('Client intro', 'Title hand-off', etc.).",
            "Allow rolling a Legacy version back to Current.",
            "Use AI to suggest the right document type during Edit Document.",
            "Add a 'Move all related documents' option when reassigning a document — copy / addendum / disclosure as a group.",
        ],
    },
    {
        "no": "27.10",
        "category": "Daily Agent / Elf Workflow",
        "feature": "All Documents — Archive a document",
        "route": "Three-dot menu on any document row (internal roles only)",
        "how_to_test": [
            "Open the three-dot menu on a document and click Archive Document.",
            "Confirm a confirmation dialog appears explaining the document will be archived and can be restored by an authorized user.",
            "Click Archive and confirm the document leaves the list.",
            "Open the Restore Archived panel (test 27.13) and confirm the document appears there, ready to bring back.",
        ],
        "expected_result": [
            "Archiving requires a confirmation step.",
            "After archive, the document leaves the list and is no longer counted in the tabs.",
            "The Restore Archived panel always shows recently archived documents, so an accidental archive can be undone from there.",
        ],
        "future_ideas": [
            "Add a one-click Undo button on the toast right after archive so the user does not have to open the Restore Archived panel for a fresh mistake.",
            "Allow batch archive of several documents at once.",
            "Auto-clean documents that have been archived for over 90 days, with an admin warning two weeks before.",
        ],
    },
    {
        "no": "27.11",
        "category": "Daily Agent / Elf Workflow",
        "feature": "All Documents — Deletion Approval Queue",
        "route": "'Deletion Queue' header button on /documents (internal roles only)",
        "how_to_test": [
            "Sign in as Team Lead, Admin, or any other internal role.",
            "If no clients have flagged anything yet, confirm the Deletion Queue button is not shown at all (the header stays clean).",
            "When at least one document has been flagged for deletion (use 27.15 to flag one as a client / FSBO customer), confirm the Deletion Queue button now appears in the page header with a small number badge showing how many requests are waiting.",
            "Confirm you also see flagged documents inside the AI Priority list as 'Review deletion request' items, with higher importance if the document has already been signed.",
            "Click Deletion Queue. Review each pending request — confirm you can see the document name, the reason given by the requester, and a decision-notes field.",
            "Click Approve on one request. Confirm a clear confirmation window opens before anything is archived. If the document was signed, confirm the warning says so explicitly. Click Archive document to confirm.",
            "Click Reject on another request without typing a reason. Confirm you are blocked with a clear 'please enter a reason' message — Reject requires a reason so the requester knows why their request was turned down.",
            "Add a reason and click Reject. Confirm the document stays in place and the request is closed.",
            "Sign in as the original requester (the client / FSBO / vendor who flagged the document). Confirm they receive an in-app notification AND an email telling them the decision and the reviewer's reason.",
            "Try to approve deletion of a document that has an envelope already sent for signature. Confirm the system blocks it with a message telling you to void the envelope first.",
        ],
        "expected_result": [
            "Every flagged request from a client / FSBO / vendor appears here for an internal reviewer, and also as a priority item in the AI Priority list so it cannot be missed.",
            "The button only shows up when there is work to do — the badge tells you how many requests are waiting.",
            "Approve always requires an explicit confirmation step. Reject always requires a reason.",
            "Both decisions are recorded for audit, and the requester always hears back by notification and email.",
        ],
        "future_ideas": [
            "Add bulk approve / bulk reject for high-volume cleanup.",
            "Show the requester their previous flag history if they have flagged several documents in a short period.",
        ],
    },
    {
        "no": "27.12",
        "category": "Daily Agent / Elf Workflow",
        "feature": "All Documents — Connect DocuSign wizard",
        "route": "'Connect DocuSign' button inside Send for Signature (when no provider is connected yet)",
        "how_to_test": [
            "Make sure DocuSign is not yet connected. Open the Send for Signature modal and click Connect DocuSign on the banner.",
            "In the wizard, click Continue to DocuSign and complete sign-in in the popup that opens.",
            "After the popup closes, click Back to Send for Signature.",
            "Try the failure path: cancel the DocuSign popup before signing in. Confirm the wizard offers a Retry button.",
        ],
        "expected_result": [
            "The whole connection flow happens inside Velvet Elves; the user never leaves /documents.",
            "On success, a toast confirms the connection and the Send for Signature modal becomes usable.",
            "If the popup is cancelled, the wizard offers a retry without restarting from scratch.",
        ],
        "future_ideas": [
            "Send a tiny test envelope to the signed-in user's own email after a successful connection so they can verify the link works.",
            "Detect popup-blocked browsers up front and offer a redirect flow instead.",
        ],
    },
    {
        "no": "27.13",
        "category": "Daily Agent / Elf Workflow",
        "feature": "All Documents — Restore Archived",
        "route": "'Restore archived' button in the page header on /documents (internal roles only)",
        "how_to_test": [
            "Archive one or two documents from the three-dot menu on a row (see test 27.10).",
            "Click the Restore archived button in the page header (it sits next to Send for Sig and Upload).",
            "Confirm a panel opens listing recently archived documents.",
            "Confirm each row shows the document name and the date it was archived.",
            "Click Restore on a row. Confirm the document comes back to the active list and the row leaves the Restore Archived panel.",
        ],
        "expected_result": [
            "Recently archived documents can be restored from this panel without contacting an admin.",
            "Restoring a document brings it back exactly where it was — same transaction, same version history.",
        ],
        "future_ideas": [
            "Show who archived the document and the reason, so the reviewer can decide whether to restore.",
            "Allow searching the panel by document name or transaction address for tenants with a lot of archived documents.",
            "Add a 'Restored by' note to the document's history so the audit trail is complete.",
        ],
    },
    {
        "no": "27.14",
        "category": "Daily Agent / Elf Workflow",
        "feature": "All Documents — Bulk actions on the Missing tab",
        "route": "/documents (Missing tab)",
        "how_to_test": [
            "Open the Missing tab. Each row should have a small checkbox at the left edge.",
            "Tick the checkboxes on two or three rows. A bulk-action bar should appear above the list showing 'N items selected'.",
            "From the bulk-action bar click Mark N/A. Confirm all selected rows clear at once.",
            "Select two or three more rows and click Request. Confirm a request email window opens for the first selected row (the system asks one recipient at a time — sending one email to several different recipients at once is not yet supported).",
            "Select rows and click Upload / Assign. Confirm the Upload window opens with the first selected row's transaction and document type already filled in.",
            "Click Clear on the bulk-action bar. Confirm the bar disappears and no rows stay selected.",
            "Confirm the bulk-action bar is only shown when at least one row is selected — no checked rows, no bar.",
        ],
        "expected_result": [
            "Multi-select is only offered on the Missing tab — other tabs do not show row checkboxes.",
            "Mark N/A runs on every selected row at once.",
            "Request and Upload / Assign use the first selected row to pre-fill their modal — they do not yet send to several recipients at once.",
            "There is no bulk Reassign on the Missing tab — there's no existing document to move yet.",
        ],
        "future_ideas": [
            "Allow Request to send to several different recipients in one pass.",
            "Add a Select-all checkbox at the top of the list.",
            "Add a 'Select rows by transaction' shortcut so the user can clear an entire deal in two clicks.",
            "Add bulk Reassign on the All Docs and Pending Review tabs once that surface is ready.",
        ],
    },
    {
        "no": "27.15",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Client / FSBO Documents portal — view and flag a document for deletion",
        "route": "/client/documents (Client) or /fsbo/documents (FSBO Customer)",
        "how_to_test": [
            "Sign in as a Client (a buyer or seller invited to a transaction) and open Documents from the sidebar. Confirm a real list of documents shared with you appears (not a stub or count-only summary).",
            "Sign in as an FSBO Customer and open Documents from the sidebar. Confirm the same kind of real list appears.",
            "Confirm each row shows the document name, the upload date, and the document type. If a document is already flagged, a small 'Flagged' badge appears next to its name.",
            "On a row that is not yet flagged, click the Flag for deletion button on the right edge of the row. Fill in the reason field and submit.",
            "Confirm the row now shows a 'Flagged' badge and the Flag for deletion button is no longer offered for that row.",
            "Sign in as the Agent or Team Lead for that transaction. Confirm the document now appears on the Deletion Queue (test 27.11) and as a 'Review deletion request' item in the AI Priority list.",
            "After the agent approves or rejects, sign back in as the original requester. Confirm a notification appears in the bell and an email is in your inbox with the decision and any reviewer note.",
        ],
        "expected_result": [
            "Client and FSBO users see a real document list, not a placeholder.",
            "Anyone outside the agent team can flag a document for deletion, and the agent always hears about it through the queue.",
            "Once flagged, the row shows a Flagged badge so the requester does not flag the same document twice.",
            "The requester always hears back about the decision — through both the in-app bell and email.",
        ],
        "future_ideas": [
            "Add an inline preview of each document so the requester can re-read it before flagging.",
            "Show which transaction each document belongs to (currently only the date and type are shown).",
            "Allow the requester to attach a short note when their reason needs more than one line.",
            "Show the same list to Vendors on the vendor portal once that surface is ready.",
        ],
    },
    {
        "no": 28,
        "category": "Daily Agent / Elf Workflow",
        "feature": "Settings hub — find any setting in one place",
        "route": "Avatar menu (top-right) → Settings (internal roles). Web address: /settings.",
        "how_to_test": [
            "Open the avatar menu in the top-right corner and click Settings. Confirm a Settings page with a search box on top and cards grouped under 'Personal Settings' and 'Workspace Settings' (internal Velvet Elves staff also see a 'Platform' group).",
            "Confirm the Personal cards include Profile, Notifications, Email & E-signature, Email Templates, My Playbook, and Help & Tour.",
            "As an Admin or workspace owner, confirm the Workspace cards include Company, Branding, Users & Invites, Task Templates, Document Templates, Vendor Templates, Team Playbook, Integrations & Webhooks, AI & Automation, Payment Access, Advertising, and Delete Organization. (A 'Billing & Credits' card also appears here once credit billing is switched on — it is off for now.)",
            "Sign in as an Agent or Transaction Coordinator and confirm you only see the Personal cards (no Workspace cards) — you should never see a card that leads to a 'not allowed' page. A Team Lead additionally sees Users & Invites, Task Templates, Vendor Templates, and Team Playbook.",
            "Type into the search box (for example 'docusign', 'logo', or 'credits') and confirm the cards filter as you type; clear it to see them all again.",
            "Click any card and confirm it opens that setting, with a breadcrumb or back link to return to the hub.",
        ],
        "expected_result": [
            "Every setting is reachable from this one hub — you no longer hunt through separate Account and Organization menus.",
            "You only see the cards your role can actually use.",
            "Search narrows the cards by title, description, or keyword.",
        ],
        "future_ideas": [
            "Show a 'recently used' settings row at the top.",
            "Add a keyboard shortcut to jump straight into settings search.",
            "Let an Admin pin the cards their team uses most.",
        ],
    },
    {
        "no": "28.1",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Settings — Profile (your identity)",
        "route": "Settings → Profile card (/settings/account). Portal roles (Client / FSBO / Vendor) open the same Profile from the avatar menu → Settings.",
        "how_to_test": [
            "Open the avatar menu → Settings, then open the Profile card.",
            "Confirm your photo, Full name, Email, Phone, a short Bio, and (internal roles) an Email signature box.",
            "Upload a photo: drag an image onto the photo box or click Upload photo (PNG or JPG, up to 5 MB). Confirm the preview updates, then remove it and confirm it clears.",
            "Edit your name, phone, or bio. Confirm 'Save changes' only becomes active once you change something.",
            "Change your Email to a different valid address. Confirm a badly-formed email is blocked with a clear message, and a valid one saves (this becomes your new sign-in address).",
            "Internal roles: type an Email signature. It signs the AI email replies sent from your mailbox; leave it blank to fall back to your name, company, and phone.",
            "Click Save changes and confirm a success toast, then reload the page and confirm your edits stuck.",
        ],
        "expected_result": [
            "Your photo, name, phone, bio, email, and signature all save.",
            "Your email address is now editable in the app (it used to be read-only) and is checked for a valid format before saving.",
            "Your photo and name update across the app — and across every workspace you belong to — after saving.",
        ],
        "future_ideas": [
            "Add an account-security area to change your password from inside the app (today use 'Forgot password').",
            "Confirm an email change with a verification link before it takes effect.",
            "List the workspaces you belong to on the Profile section.",
        ],
    },
    {
        "no": "28.2",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Settings — Notifications, My Playbook, and Help",
        "route": "Settings → the Personal cards: Notifications, My Playbook, and Help & Tour.",
        "how_to_test": [
            ("Open the Notifications card.", [
                "Confirm a list of notification categories with Email / Push / In-app switches. Flip a few and click Save — confirm a success toast, then reload and confirm they stuck.",
            ]),
            ("Open the My Playbook card.", [
                "Confirm one page with a row of buttons across the top: Checklist, Tagged notes, Preferred vendors, and Resources. Click each to switch the tool in place. Each is your own personal list — add, edit, and remove an entry, then save. These are the 'My …' personal copies; the team-wide versions live under the Settings hub's Team Playbook card (feature 32.14).",
                "Click 'Preview closing checklist' in the header and confirm a preview of how your whole playbook prints on a sample closing.",
            ]),
            ("Open the Help & Tour card.", [
                "Confirm two cards: a Help Center card with an 'Open Help Center' button (it opens the searchable guides website in a new browser tab), and a Help & tour card with a 'Start tour' button. Click Start tour and confirm the product tour starts for whatever role you are signed in as (see feature 13).",
            ]),
            "Portal roles get these as a lightweight Account window from the avatar menu → Settings: sign in as a Client or Vendor and confirm it shows Profile only; sign in as an FSBO Customer and confirm it shows Profile plus a Preferences section.",
        ],
        "expected_result": [
            "Every section saves its own changes with a clear success toast and the changes persist after reloading.",
            "My Playbook keeps all four personal tools on one page; the team-wide copies live on the Team Playbook card.",
            "Internal roles see all personal cards; Client and Vendor see Profile only; FSBO sees Profile plus Preferences.",
            "The Help Center button opens the public guides website in a new tab; Help & Tour replays the role-appropriate product tour at any time.",
        ],
        "future_ideas": [
            "Add a one-click 'copy my checklist template to the team' shortcut.",
            "Add a daily-digest email option in Notifications.",
            "Add a per-feature mini-tour launcher in Help & Tour.",
        ],
    },
    {
        "no": "28.3",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Settings — Email Templates",
        "route": "Settings → Email Templates card (/email-templates). Agent, Transaction Coordinator, Team Lead, and Admin (Attorney does not see this card).",
        "how_to_test": [
            "From the Settings hub, open the Email Templates card. Confirm a searchable library with filter buttons (All / Starter / Shared / Personal) and a search box on one line.",
            "Confirm three kinds of template, each with a small label: Starter (built-in, read-only), Shared (visible to your whole brokerage), and Personal (your own).",
            "Open a Starter template and confirm you cannot edit it directly, but you can click 'Save a copy' to make an editable Personal copy.",
            "Click New template, fill in the name, category, subject, and body, then save. Confirm it appears with a 'Personal' label. As an Admin or Team Lead, confirm you can also create a Shared template for the whole brokerage.",
            "Confirm the body and subject can use placeholders (for example the client's name or the property address) that fill in from the deal when the email is sent.",
            "Edit one of your templates, then delete it and confirm a confirmation step appears before it is removed.",
            "Type in the search box and confirm the list narrows by name, subject, or category.",
        ],
        "expected_result": [
            "Starter templates are read-only; you customise them with 'Save a copy'.",
            "Personal templates are yours alone; Shared templates are visible to the whole brokerage and only Admins / Team Leads can create them.",
            "These templates appear in the Compose email flow, with placeholders filled from the deal.",
        ],
        "future_ideas": [
            "Show a live preview of a template with a sample deal filled in.",
            "Let a user duplicate any template (not only Starter ones) as a starting point.",
            "Track how often each template is used so the team can retire low-value ones.",
        ],
    },
    {
        "no": "29.1",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Workspace settings — Company details",
        "route": "Settings → Company card (/organization?section=company). Admin or workspace owner only.",
        "how_to_test": [
            "From the Settings hub, open the Company card.",
            "As an Admin, confirm an editable Organization name field with a Save changes button, plus a read-only Plan badge and a seats summary.",
            "Type a new name and click Save changes — confirm a success toast.",
            "Sign in as another member of the same workspace and confirm they see the new name.",
            "Note: Company, Branding, Billing & Credits, and Delete Organization are the Workspace cards; per-user Email & E-signature connections now live under the Personal 'Email & E-signature' card (feature 29.2).",
        ],
        "expected_result": [
            "Admins / owners can rename the workspace and see the current plan and seat usage.",
            "The new name shows up for every member of the brokerage.",
            "Non-admins do not see the Company card at all (it is a Workspace setting).",
        ],
        "future_ideas": [
            "Show a preview of how the name appears in the sidebar, invitation email, and outbound transaction emails.",
            "Add a 'workspace web address' (subdomain) field.",
            "Show the workspace owner's name and email on this card.",
        ],
    },
    {
        "no": "29.2",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Settings — Email & E-signature connections (Gmail, Outlook, DocuSign)",
        "route": "Settings → Email & E-signature card (/settings/connections). These connections are per-user — every internal user connects their own accounts.",
        "how_to_test": [
            ("Email integrations.", [
                "Confirm a Gmail row and an Outlook row (iCloud is intentionally hidden for now).",
                "Click Connect on the Gmail row and complete sign-in in the Google popup. After approval the row switches to Connected with your email and the date.",
                "Repeat on the Outlook row using a Microsoft 365 account.",
                "Cancel a popup mid-way and confirm the row stays on 'Connect' without an error.",
                "Click Disconnect on a connected row, read the warning that inbound sync and AI email automation will stop, then cancel (row stays connected) and try again to confirm (row returns to Connect).",
                "Click Refresh to re-fetch the list.",
            ]),
            ("E-signature (DocuSign).", [
                "If DocuSign is not connected, click Connect and complete the wizard (Intro → DocuSign popup → Done). Confirm the row then shows your DocuSign account email and the date.",
                "Click Disconnect, read the warning that Send for Signature will be disabled across the app, and confirm.",
            ]),
        ],
        "expected_result": [
            "Every provider connects through its official sign-in popup — no password is typed into Velvet Elves.",
            "Disconnect always asks for confirmation first.",
            "At least one inbox must be connected for AI Email Review (features 29.7+) to send replies, and DocuSign must be connected for Send for Signature on the Documents center.",
        ],
        "future_ideas": [
            "Show a 'Last synced' time and a manual 'Sync now' button per inbox.",
            "Re-enable the iCloud row once the Apple app-specific-password flow is reviewed.",
            "Show the monthly DocuSign envelope count remaining so users do not hit their quota by surprise.",
        ],
    },
    {
        "no": "29.3",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Workspace settings — Document Templates (your own fillable forms)",
        "route": "Settings → Document Templates card (/settings/document-templates). Admin or workspace owner only.",
        "how_to_test": [
            "From the Settings hub, open the Document Templates card (Admins / owners only).",
            "Click Upload template and pick one of your brokerage's PDF forms. Confirm it uploads and an editor opens on the new draft.",
            "In the editor, set a Document type, then map each detected field to a deal value (for example a 'Buyer' field maps to the buyer's name). Click Save & activate.",
            "Back on the list, confirm the template now shows an Active badge and how many fields are mapped. A form that already has data typed into it shows a 'Contains data' warning.",
            "Click Edit to change a mapping, and the trash icon to delete a template (confirm it disappears).",
            "Note: once a template is Active, the app fills your own form in place with the deal's data (keeping its exact layout) and flattens it whenever you generate that document — your form takes precedence over the built-in one.",
        ],
        "expected_result": [
            "You can upload, map, activate, edit, and delete your own PDF forms.",
            "An active template is used to generate that document type, in your form's exact layout.",
            "Only Admins / owners see this card; the editor warns when an uploaded form already contains data.",
        ],
        "future_ideas": [
            "Map checkbox and signature fields, not just text fields.",
            "Let the AI suggest the field mapping automatically from the form's labels.",
            "Support Word (DOCX) templates alongside fillable PDFs.",
        ],
    },
    {
        "no": "29.4",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Workspace settings — Branding (white-label)",
        "route": "Settings → Branding card (/organization?section=branding). Admin or workspace owner only.",
        "how_to_test": [
            "From the Settings hub, open the Branding card as an Admin. Confirm a logo upload, a brand-color field, a display-name field, and a live preview.",
            "Upload a logo (PNG, JPEG, WEBP, SVG, or GIF, up to 2 MB). Try a wrong file type or an oversized file and confirm a clear error.",
            "Pick a brand color and a display name, watch the live preview update, then click Save branding.",
            "Refresh the page and confirm your logo, color, and display name are still there.",
            "Confirm the saved logo and color now show in the sidebar and on the sign-in page; check that they also appear on outbound/printed documents.",
            "Sign in as a non-Admin and confirm the Branding card does not appear in the Settings hub.",
        ],
        "expected_result": [
            "Branding now fully saves — logo, brand color, and display name persist and apply across the app immediately (this was a placeholder before and is now live).",
            "Invalid logo files are rejected with a clear message.",
            "Only Admins can change branding.",
        ],
        "future_ideas": [
            "Add a 'Reset to Velvet Elves defaults' button.",
            "Show a sample invitation email and a sample printed document in the preview.",
            "Let an Admin preview dark-mode branding before saving.",
        ],
    },
    {
        "no": "29.7",
        "category": "Daily Agent / Elf Workflow",
        "feature": "AI Email Review — overview, list, and filter tabs",
        "route": "/ai-emails",
        "how_to_test": [
            "Open Intelligence → AI Email Review from the sidebar (Agent, Transaction Coordinator, Team Lead, and Admin only; Attorney and external roles do not see this entry).",
            "Confirm the filter tabs: All, Needs Review, Ready to Send, Low Confidence, Escalated — each with a count.",
            "Click each tab in turn and confirm the list narrows to that subset.",
            "Click Refresh and confirm the list reloads.",
            "Leave the page open for at least 60 seconds. The list should silently re-fetch every minute, and any new draft should appear without a manual refresh.",
            "Open the top-bar bell on another page and confirm the unread count matches the All tab here.",
            "Force an empty state by switching to a tab with no matches and confirm a clear empty message appears.",
        ],
        "expected_result": [
            "The list only contains drafts your role is allowed to act on.",
            "Auto-refresh runs every 60 seconds; the list also refreshes immediately after Approve / Edit / Discard / Regenerate.",
            "Empty and error states show a clear next step.",
        ],
        "future_ideas": [
            "Add an inline search box and sort control inside the list.",
            "Add per-row checkboxes plus bulk Approve / Discard so a reviewer can clear a backlog quickly.",
            "Wire the deep-link /ai-emails/:logId so notifications open the exact draft in question.",
            "Stream live updates so the list does not have to poll every minute.",
        ],
    },
    {
        "no": "29.8",
        "category": "Daily Agent / Elf Workflow",
        "feature": "AI Email Review — what each draft row shows",
        "route": "/ai-emails (list of drafts)",
        "how_to_test": [
            "Look at any draft row. Confirm it shows: the subject, the recipient(s), the kind of email (Factual question / Document request / Vendor reply / Uncertain / Other), a confidence percent, an Escalated marker if the draft is past its deadline, and a 'how long ago' timestamp.",
            "Click a row and confirm the detail pane on the right loads the full draft (see 29.9).",
            "Note: there is no per-row search, sort, or bulk action yet — the only filtering surface is the tab row at the top. Please flag if this is missing for the way you want to work.",
        ],
        "expected_result": [
            "Every row tells you at a glance how confident the AI is, who the email is to, and whether it is overdue.",
            "Selection survives the 60-second auto-refresh as long as the draft is still in the same filter.",
        ],
        "future_ideas": [
            "Show the transaction address on each row so a reviewer can scan deals without opening each draft.",
            "Show an attachment indicator (planned alongside attachment support).",
            "Offer a compact / comfortable density toggle for high-volume reviewers.",
        ],
    },
    {
        "no": "29.9",
        "category": "Daily Agent / Elf Workflow",
        "feature": "AI Email Review — what the draft detail shows",
        "route": "/ai-emails (open any draft)",
        "how_to_test": [
            "Open any draft and confirm the top of the detail shows the email kind, the confidence percent, and an Escalated marker if applicable. The subject and the To / Cc lines are also visible (by default the file owner is automatically CC'd).",
            "Read the AI-drafted reply. Any phrases the AI is unsure about should be highlighted in the body, and if it made any explicit assumptions, those are listed below the body as 'Flagged assumptions'.",
            "On the right side, confirm an 'AI Verified From' card lists every piece of source data the AI used (address, closing date, status, document names, etc.). If the list is empty, a warning explains that no source data was cited and that the agent should verify each fact manually.",
            "Below it, confirm an 'Original Inbound' card shows the email that triggered this draft (sender, time, subject, body). If the original can't be loaded, the card says so cleanly.",
        ],
        "expected_result": [
            "The confidence percent, kind, and escalation status match whatever was shown on the list row.",
            "Highlighted phrases in the body match the explicit 'Flagged assumptions' list.",
            "AI Verified From only shows facts the AI actually used — it does not guess.",
            "Original Inbound shows the email that triggered this reply, not the entire thread (full-thread view is a future improvement).",
        ],
        "future_ideas": [
            "Show the entire thread, not just the immediate inbound, so a reviewer can see prior context.",
            "Make each AI Verified From row clickable so it deep-links to the source field on the transaction record.",
            "Add an inline 'Flag this fact as wrong' button that pushes a correction back into the AI's training data.",
        ],
    },
    {
        "no": "29.10",
        "category": "Daily Agent / Elf Workflow",
        "feature": "AI Email Review — Approve & Send, Edit & Send, Regenerate, Discard",
        "route": "/ai-emails (action buttons on the open draft)",
        "how_to_test": [
            "Open a draft and click Approve & Send. Confirm the email is sent and a toast confirms.",
            "Open another draft and click Edit. The subject and body should become editable. Make a change and click Send Edit — confirm the edited version is sent.",
            "Click Regenerate on a draft. Confirm the AI redraws a fresh reply from the original inbound email.",
            "Click Discard. Confirm a warning explains the draft will be removed but the original inbound message stays in the communication log. Confirm Discard removes the draft.",
            "Disconnect your email provider in Settings → Email & E-signature, then click Approve & Send on a draft. Confirm a clear error explains that no email provider is connected.",
            "Confirm the actions that are NOT here yet: no Reassign, no 'Mark Reviewed', no attachment uploader, no scheduled-send.",
        ],
        "expected_result": [
            "Approve, Edit, Regenerate, and Discard all complete with a toast and the list refreshes.",
            "Editing a draft also clears its flagged assumptions, since the human reviewer rewrote the content.",
            "Discard preserves the original inbound message in the communication log — only the draft disappears.",
            "If your role is not allowed to act, an error toast explains it on click rather than the button being hidden.",
        ],
        "future_ideas": [
            "Add a Reassign button so a Team Lead can hand a draft to a colleague without opening it.",
            "Add a Mark Reviewed (no-send) status for drafts the human read but does not want to send.",
            "Add support for attachments when replying so contracts and addenda can ride out with the AI reply.",
            "Add a per-tenant 'Auto-send when confidence is over X%' threshold (backend already supports it; no UI yet).",
        ],
    },
    {
        "no": "29.11",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Vendor directory",
        "route": "/vendors",
        "how_to_test": [
            "Open the sidebar and click Vendors, or go to /vendors directly.",
            "Confirm a card for every vendor company shows the name, category, contact email, and phone.",
            "Use the search box to find a vendor by company name, category, or email.",
            "Use the category dropdown and the 'Preferred only' toggle to narrow the list.",
            "Click New vendor and add a vendor company. Confirm it appears in the directory.",
            "Click a vendor card to open the vendor detail page.",
        ],
        "expected_result": [
            "Every vendor stored for your brokerage shows up in the directory and stays reachable from any transaction.",
            "Search, category, and preferred filters all narrow the visible list in real time.",
            "New vendors created here can be assigned to transactions afterwards.",
        ],
        "future_ideas": [
            "Show the number of open transactions each vendor is currently assigned to.",
            "Allow tagging vendors with custom labels (for example 'fast turnaround', 'cash only').",
        ],
    },
    {
        "no": "29.12",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Vendor detail page (contacts, portfolio, background refresh)",
        "route": "/vendors/{vendorId}",
        "how_to_test": [
            "Open a vendor from the directory.",
            "Confirm the page shows the company name, category, address, phone, email, and preferred status.",
            "Under Contacts, confirm each person on file shows a name, email, and phone, with the primary contact clearly labelled.",
            "Click the Email button on a contact and confirm the Send Vendor Request modal opens with that contact pre-filled.",
            "Hover the Call button — until SMS / voice integration is enabled it should show a 'Coming soon' note.",
            "Under Portfolio, confirm a list of transactions where this vendor is on the team.",
            "Click Refresh info — the Background refresh drawer opens (see entry 29.16).",
            "Click Add colleague (public link) — confirm a one-time link is created and copied to the clipboard.",
        ],
        "expected_result": [
            "The page summarises everything you need to know about the vendor in one screen.",
            "Email opens the request flow with the right contact selected.",
            "Add colleague creates a public link that the vendor can use to add a teammate without logging in.",
            "Refresh info opens the suggestions drawer; nothing applies without an explicit click.",
        ],
        "future_ideas": [
            "Show the latest vendor reply per transaction directly on the portfolio list.",
            "Allow the agent to set vendor-specific reply expectations (for example 'always reply within 24h').",
        ],
    },
    {
        "no": "29.13",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Send Vendor Request (template-based email)",
        "route": "Email button on a vendor contact card, or from a task that has a vendor assignment",
        "how_to_test": [
            "Open a vendor detail page and click Email on one of the contacts.",
            "Pick a template from the list (for example 'Inspection — Schedule Visit'). Confirm the right pane previews the subject and body with the transaction address, task name, and primary contact name filled in.",
            "Confirm the body ends with the reply footer 'Reply with: Scheduled: YYYY-MM-DD' (or the equivalent footer for the chosen template).",
            "Edit the subject or body inline if you want.",
            "Click Send request.",
            "Open the Communications panel for the same transaction and confirm a new outbound row appears with the right vendor and timestamp.",
            "Check the vendor's inbox in another window — the email should be there.",
        ],
        "expected_result": [
            "The request is sent through the agent's connected Gmail or Outlook (not from a Velvet Elves address).",
            "Both the email and a record of it appear in the transaction's communication log.",
            "If no email provider is connected, the modal explains how to fix it and does not silently fail.",
        ],
        "future_ideas": [
            "Allow attaching the most recent contract or inspection report directly from the modal.",
            "Schedule the send for later (overnight or a specific weekday).",
        ],
    },
    {
        "no": "29.14",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Vendor proposals queue",
        "route": "/vendor-proposals (sidebar → Intelligence → Vendor Proposals)",
        "how_to_test": [
            "Open /vendor-proposals from the sidebar.",
            "Confirm three tabs: 'Awaiting decision', 'Awaiting vendor', and 'All open'.",
            "Each proposal card should show the task name, the vendor company, the original task date, and the date the vendor is proposing.",
            "Click Accept & update task on a pending proposal — confirm a toast reports the task date was updated, and check the matching transaction's task tab to verify the new date.",
            "Click Ask vendor to clarify on a vague proposal (one with no clear date) — confirm the proposal moves to the 'Awaiting vendor' tab.",
            "Click Reject on another proposal — confirm a toast reports the task date is unchanged.",
        ],
        "expected_result": [
            "Every vendor reply that proposes a new date lands here for the agent to approve.",
            "Accept is the only action that changes a task's due date — nothing happens automatically.",
            "Decisions are recorded in the audit log for compliance.",
        ],
        "future_ideas": [
            "Show the original vendor email inline so the agent does not have to bounce between this page and /ai-emails.",
            "Suggest an alternative date when Reject is clicked, so the agent can immediately reply with one.",
        ],
    },
    {
        "no": "29.15",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Add a vendor colleague (public link)",
        "route": "'Add colleague (public link)' button on a vendor contact card; the public page is /v/{token}",
        "how_to_test": [
            "On a vendor detail page, click Add colleague (public link). Confirm a one-time link is copied to your clipboard.",
            "Open the link in a private / incognito browser window. The page should show only your brokerage name and the vendor company name (no transaction details).",
            "Submit the form with first name, last name, email, optional phone, and optional title.",
            "Confirm the success screen reads 'You're on the thread.'",
            "Back in the authenticated app, reload the vendor detail page and confirm the new contact appears in the Contacts list (not marked primary).",
            "Try opening the same link again — the page should report the link is no longer valid (single-use).",
        ],
        "expected_result": [
            "Vendors can attach a colleague without creating an account.",
            "The public page never leaks any transaction information.",
            "Each link can be used exactly once and expires after the default window (7 days).",
        ],
        "future_ideas": [
            "Allow the vendor to set their own preferred contact channel (email or SMS) right on the public form.",
            "Let the agent customise the welcome message that appears on the public page.",
        ],
    },
    {
        "no": "29.16",
        "category": "Daily Agent / Elf Workflow",
        "feature": "Vendor background refresh (suggested updates)",
        "route": "'Refresh info' button on a vendor detail page",
        "how_to_test": [
            "Open a vendor detail page that already has one or more contacts and click Refresh info.",
            "In the drawer, click Run refresh. Confirm suggestions appear as 'Current value' vs. 'Suggested value' cards with a confidence score and source label.",
            "Tick one or two suggestions and click Apply selected.",
            "Confirm a toast reports the vendor record was updated, and the At-a-glance section on the vendor detail page reflects the new values.",
            "Reopen the drawer and confirm the suggestions you accepted are gone but the others remain.",
        ],
        "expected_result": [
            "Suggestions are based on existing tenant data (other transactions, other contacts) — no public web search.",
            "Nothing changes on the vendor record without an explicit click.",
            "Every accepted change is recorded in the audit log per field.",
        ],
        "future_ideas": [
            "Pull from approved public sources (state licensing boards, business directories) once that is reviewed.",
            "Run the refresh on a schedule for preferred vendors so suggestions are always fresh.",
        ],
    },

    # ============================================================
    # SECTION 4 — ADMIN / TEAM LEAD EXTRAS
    # ============================================================
    {
        "no": 30,
        "category": "Admin / Team Lead Extras",
        "feature": "Task Templates list",
        "route": "Settings → Task Templates card (/admin/task-templates). Team Lead and Admin.",
        "how_to_test": [
            ("From the Settings hub, open the Task Templates card.", []),
            ("Type in the search box to filter templates.", []),
            ("Confirm templates are grouped by category.", []),
            ("Click New Template and create a template in the pop-up window.", [
                "Fill in Name, Description, Automation Level, and Category.",
                "Submit.",
            ]),
            ("Click an existing template row to open its details in a pop-up window (see feature 31).", []),
        ],
        "expected_result": [
            "The newly created template appears in the list right away.",
            "Creating and opening a template both happen in a pop-up window — you stay on the list the whole time.",
        ],
        "future_ideas": [
            "Allow CSV import and export of templates.",
            "Add a 'Duplicate this template' shortcut.",
        ],
    },
    {
        "no": 31,
        "category": "Admin / Team Lead Extras",
        "feature": "Task Template detail",
        "route": "Settings → Task Templates card → click any template row (opens a pop-up window).",
        "how_to_test": [
            ("Open a template by clicking its row. A details pop-up opens with the read-only view.", []),
            ("Click Edit Template and update several fields in the edit pop-up.", [
                "Name, description, target, milestone label, dependency, float days, automation level, category, sort order, active toggle.",
            ]),
            ("Try the dependency rule builder section.", []),
            ("Click Save.", []),
        ],
        "expected_result": [
            "All changes persist and the pop-up returns to the read-only view.",
        ],
        "future_ideas": [
            "Add a visual dependency graph so admins can see how tasks connect.",
            "Show a 'where used' indicator listing which transaction types use this template.",
        ],
    },
    {
        "no": 32,
        "category": "Admin / Team Lead Extras",
        "feature": "Member details pop-up",
        "route": "Settings → Users & Invites card → click any member row (opens a pop-up window). Also opens from the Teams page roster.",
        "how_to_test": [
            "Open the Users & Invites card (Settings hub) and click any member row in the table. A details pop-up should open.",
            "Confirm the pop-up shows the member's name, role, email, phone, team, status (Active / Deactivated), joined date, last sign-in, and bio (if any). The workspace owner is clearly marked, and a member invited from another brokerage shows a 'Guest' badge.",
            "Confirm the owner-only Transfer ownership button and the Admin-only Deactivate button appear in the pop-up footer only when they apply (see features 32.5 and 32.6); they never appear on your own row or on the owner's row.",
            "Close the pop-up and confirm you return to the table.",
        ],
        "expected_result": [
            "Every member's full profile opens in a pop-up window — there is no separate page to navigate to.",
            "The owner and any guest members are clearly marked, and the management actions only appear for users who can use them.",
        ],
        "future_ideas": [
            "Add an audit trail showing the user's recent actions (logins, role changes, transactions worked).",
            "Add an inline Edit form so an Admin can update name, phone, or role from the pop-up.",
            "Show recent transactions assigned to this user so a Team Lead has full context.",
        ],
    },
    {
        "no": "32.1",
        "category": "Admin / Team Lead Extras",
        "feature": "Team Overview page",
        "route": "/team — sidebar → Team → Team Overview (Team Lead and Admin only)",
        "how_to_test": [
            "Sign in as Team Lead or Admin and open Team Overview from the sidebar's Team group.",
            "Try the same as Agent, Transaction Coordinator, Attorney, Client, FSBO, or Vendor — the page should redirect to Unauthorized.",
            "Confirm the four numbers across the top: Active Members, Active Tx (open transactions), Pipeline (open volume), and Pending Invites.",
            "Confirm the main 'Agents & production' table lists each member with their role, team, active-transaction count, and pipeline value, sorted with the biggest producers first. The workspace owner is marked.",
            "Confirm the supporting cards: By team (each team with its lead and counts), Role coverage (how many people in each role), Awaiting acceptance (pending invites with time left), and Last seen (recent sign-ins).",
            "Click a member's name and confirm it opens the Users & Invites list; click Manage teams in the header (or a team in the 'By team' card) and confirm it opens the Teams page (feature 32.13).",
            "Sign in as a Team Lead and confirm the page is scoped to only your own team (your team's name in the title, only your team's members, transactions, and invites) — you never see other teams or workspace-wide totals.",
        ],
        "expected_result": [
            "Only Team Lead and Admin can reach the page; an Admin sees the whole brokerage, a Team Lead sees only their own team.",
            "Every panel reflects real data, and the production table never disagrees with the member count.",
            "The header and team cards navigate to the Teams page.",
        ],
        "future_ideas": [
            "Let the user filter the table by clicking a role on the Role coverage card.",
            "Show a 7-day activity sparkline on the top numbers.",
            "Allow pinning a favourite member for quick contact from the side rail.",
        ],
    },
    {
        "no": "32.2",
        "category": "Admin / Team Lead Extras",
        "feature": "Users & Invites — Active members tab",
        "route": "Settings → Users & Invites card (/admin/users). Team Lead and Admin.",
        "how_to_test": [
            "Open the Users & Invites card from the Settings hub (or by clicking a member's name on Team Overview).",
            "Sign in as Agent or another lower role and confirm the card is not shown and the URL redirects to Unauthorized.",
            "Confirm two tabs: 'Active members' and 'Pending invitations', plus an 'Invite user' button in the header.",
            "On the Active members tab, confirm the members appear in a table (Member, Role, Team, Last sign-in). Type in the search box and confirm the list filters by name and email as you type.",
            "Open the role dropdown, pick a role, and confirm only members with that role remain. As an Admin, also use the team dropdown to show only one team's members. Combine the filters and confirm they all apply.",
            "Click any member row and confirm the member-details pop-up opens (feature 32), where the Transfer ownership (owner only) and Deactivate (Admin only) actions live.",
            "Apply a search term with no matches and confirm a clear 'No members match this filter' message appears.",
        ],
        "expected_result": [
            "Members are shown in a clean table; clicking a row opens the details pop-up.",
            "Search, role filter, and (Admin) team filter narrow the list in real time.",
            "The owner and any guest members are clearly marked in the table.",
        ],
        "future_ideas": [
            "Show a 'last active' time more prominently in the table.",
            "Add bulk selection so an Admin can deactivate or change role for multiple members at once.",
            "Add an inline 'Resend welcome email' action for members who never signed in.",
        ],
    },
    {
        "no": "32.3",
        "category": "Admin / Team Lead Extras",
        "feature": "Users & Invites — Pending invitations tab",
        "route": "Settings → Users & Invites card → 'Pending invitations' tab (/admin/users)",
        "how_to_test": [
            "Switch to the Pending invitations tab on the Users & Invites page.",
            "If there are pending invites, confirm each row shows the invited email, the invited role, and how much time is left before the invite expires.",
            "Click Copy link on a row and confirm the invite URL is copied to your clipboard.",
            "In the row's three-dot menu, try Resend email, Extend by 72h, and Revoke invitation. Each should show a confirming toast and update the row.",
            "If there are no pending invites, confirm a clear empty state pointing at the Invite user button.",
            "Block your browser clipboard and click Copy link — confirm a clear error appears rather than a silent failure.",
        ],
        "expected_result": [
            "All four row actions (Copy link, Resend, Extend, Revoke) work on pending invitations.",
            "Accepted invitations do not appear on this tab.",
            "Revoking an invitation removes the row and prevents the invitee from accepting.",
        ],
        "future_ideas": [
            "Add a Bulk revoke expired button so cleanup is one click instead of many.",
            "Show who originally sent each invitation so a Team Lead can see which Admin invited a particular email.",
            "Show a preview of the full invite URL before copying.",
        ],
    },
    {
        "no": "32.4",
        "category": "Admin / Team Lead Extras",
        "feature": "Invite user modal",
        "route": "Settings → Users & Invites card → 'Invite user' button. (The Teams page has a matching 'Invite to team' button that pre-locks the team — see feature 32.13.)",
        "how_to_test": [
            "Click Invite user on the Users & Invites page.",
            "Confirm two required fields: Email address and Role. A short hint under the dropdown explains what each role can do.",
            "Sign in as different inviter roles (Agent, Team Lead, Admin) and confirm the dropdown only offers roles you are allowed to grant — for example only Admins can invite new Admins.",
            "Try to send with an empty email and confirm it is blocked.",
            "Try to invite an email that already belongs to a member and confirm an inline error explains the email is already in use.",
            "If your plan has a seat cap, try inviting beyond the cap and confirm a seat-limit error appears.",
            "Send a valid invitation and confirm a success toast appears, the modal closes, and the new row appears on the Pending invitations tab.",
        ],
        "expected_result": [
            "Roles you cannot grant are not offered in the dropdown.",
            "Every failure mode shows a clear inline error rather than breaking the modal.",
            "A successful invite appears in the Pending invitations list immediately.",
        ],
        "future_ideas": [
            "Allow inviting multiple emails at once (comma-separated) so an admin can bulk-onboard a team.",
            "Show a 'view what the invitee will see' preview link before sending.",
            "Let the admin attach a team or a transaction at invite time (planned for Phase 5).",
        ],
    },
    {
        "no": "32.5",
        "category": "Admin / Team Lead Extras",
        "feature": "Transfer workspace ownership",
        "route": "Settings → Users & Invites → click a member → Transfer ownership button in the pop-up (workspace owner only)",
        "how_to_test": [
            "Sign in as the workspace owner and confirm you are clearly marked as the owner in the Users & Invites table.",
            "Click a different member's row to open their details pop-up and confirm a Transfer ownership button is offered in the footer. It should not appear on your own row or on a member who is already the owner.",
            "Click Transfer ownership. Confirm the warning dialog explains: the target will be promoted to Admin if they are not already one; you will stay as Admin but lose owner-only abilities (schedule deletion, transfer ownership); and the action is logged.",
            "Confirm the transfer. After success, the owner mark moves from your row to the new owner's row.",
            "Refresh the page and confirm the new ownership state persists.",
            "Sign in as a regular Admin (not the owner) and confirm Transfer ownership is hidden from every member's pop-up.",
        ],
        "expected_result": [
            "Only the current owner sees the Transfer ownership control.",
            "The confirmation dialog clearly describes the side effects before you commit.",
            "After transfer, the previous owner can no longer schedule deletion or transfer ownership — those controls disappear from Settings.",
        ],
        "future_ideas": [
            "Send an automatic email to the new owner immediately so they know responsibility has moved.",
            "Add an undo window (for example 30 minutes) on a recently-transferred ownership.",
            "Show the full audit trail of past ownership transfers on the Team Overview page.",
        ],
    },
    {
        "no": "32.6",
        "category": "Admin / Team Lead Extras",
        "feature": "Deactivate a member",
        "route": "Settings → Users & Invites → click a member → Deactivate button in the pop-up (Admin only)",
        "how_to_test": [
            "Sign in as an Admin who is not the target. Open another member's details pop-up and confirm a Deactivate button is offered in the footer.",
            "Confirm Deactivate is hidden on the workspace owner's row and on your own row.",
            "Click Deactivate. A confirmation dialog should explain the member will no longer be able to sign in and can be re-activated later.",
            "Cancel and confirm the member is still listed. Try again and confirm — the member disappears from Active members.",
            "Switch to Pending invitations and confirm any invitations the deactivated user created are still listed there.",
        ],
        "expected_result": [
            "Deactivation requires confirmation and updates the list immediately.",
            "The owner cannot be deactivated — the option is hidden on their row.",
            "You cannot deactivate yourself — the option is hidden on your own row.",
        ],
        "future_ideas": [
            "Add a paired Re-activate control that surfaces deactivated members in a separate sub-list.",
            "Capture a deactivation reason at confirmation time so the audit log records why.",
            "Auto-revoke any active browser sessions for the deactivated member.",
        ],
    },
    {
        "no": "32.7",
        "category": "Admin / Team Lead Extras",
        "feature": "Workspace settings — Delete Organization (Danger Zone)",
        "route": "Settings → Delete Organization card (/organization?section=danger). Workspace owner only.",
        "how_to_test": [
            "Sign in as the workspace owner, open the Settings hub, and confirm a red 'Delete Organization' card appears under Workspace Settings.",
            "Sign in as any other role on the same workspace and confirm the Delete Organization card is not shown at all.",
            "As the owner, open Delete organization. Confirm the page asks you to type the workspace name exactly before Schedule deletion becomes available.",
            "Schedule deletion and confirm a clear message reports the exact date and time it will run, plus a note that audit logs and a full snapshot are archived under the 2-year retention policy.",
            "Click Cancel deletion and confirm the workspace returns to normal.",
            "If your workspace is on legal hold, confirm the Delete button is disabled with a clear explanation pointing the owner at platform support.",
        ],
        "expected_result": [
            "Only the workspace owner can see and use the Danger Zone.",
            "Scheduling deletion requires typing the workspace name exactly, and the action can be reversed during the grace period.",
            "Members can still sign in while deletion is scheduled, so the owner can cancel it.",
            "Legal hold blocks the action with a clear plain-language reason.",
        ],
        "future_ideas": [
            "Show a top-of-page banner with the deletion date and days remaining while deletion is scheduled.",
            "Send a daily reminder email to the owner while deletion is scheduled.",
            "Let the owner customise the grace-period length (default is 30 days).",
            "Add a 'Download a full export before I leave' button next to the schedule-deletion CTA.",
        ],
    },
    {
        "no": "32.9",
        "category": "Admin / Team Lead Extras",
        "feature": "Platform Admin — Tenants list (internal Velvet Elves staff only)",
        "route": "/platform/tenants (platform admins only)",
        "how_to_test": [
            "Sign in as a platform admin and open /platform/tenants.",
            "Sign in as any non-platform user and try /platform/tenants directly — confirm you get a 404 (not a 403), so the route's existence is not leaked.",
            "Use the filter dropdown (All / Active / Suspended / On legal hold) and confirm the table narrows correctly.",
            "Each tenant row should show the tenant name, slug, current status, plan, and an Actions menu.",
            "Click Details on a row and confirm the tenant detail page opens (feature 32.10).",
            "Click Suspend on an active tenant. Confirm a warning explains members will not be able to sign in and that you can reactivate later. Confirm and check the row updates.",
            "Confirm the Suspend button is disabled for a tenant on legal hold.",
            "Click Reactivate on a suspended tenant and confirm it returns to Active.",
        ],
        "expected_result": [
            "Non-platform users get a 404 — the route's existence is not leaked.",
            "Filters narrow the list correctly.",
            "Suspend / Reactivate updates the row immediately and is recorded in the platform audit log.",
        ],
        "future_ideas": [
            "Add a search box that matches name / slug / owner email.",
            "Add a Schedule deletion action on the row alongside Suspend (currently only available inside the tenant's own Settings → Delete Organization).",
            "Show owner email and member count as extra columns.",
            "Add a 'Force re-verify domain' action for tenants on a custom domain.",
        ],
    },
    {
        "no": "32.10",
        "category": "Admin / Team Lead Extras",
        "feature": "Platform Admin — Tenant detail (internal Velvet Elves staff only)",
        "route": "/platform/tenants/<tenantId> (platform admins only)",
        "how_to_test": [
            "Open a tenant detail by clicking Details on /platform/tenants or by pasting the URL.",
            "Confirm the page shows the tenant name, slug, created date, status (Active, Suspended, Legal hold, or Deletion scheduled), and an All tenants back link.",
            "Confirm the Identity panel lists ID, slug, custom domain (if any), domain verification status and timestamp, owner user id, and the invite base URL.",
            "Confirm the Plan & lifecycle panel lists plan name, seat limit, staff seats used, trial-ends date (if any), scheduled-deletion timestamp (if any), and the legal-hold reason (if any).",
            "Sign in as any non-platform user and try the URL directly — confirm you get a 404.",
            "Open the URL with an invalid tenant id and confirm the page shows a clear 'Tenant not found' message with a back link.",
        ],
        "expected_result": [
            "All read-only data renders without errors.",
            "Non-platform users cannot reach this page even with a direct URL.",
            "An invalid or deleted tenant id shows a clean fallback, not a broken page.",
        ],
        "future_ideas": [
            "Add inline edit for tenant name and plan from this detail page.",
            "Show the tenant's recent audit-log events (last 50) directly on this page.",
            "Add a Reset domain verification control for tenants stuck on unverified.",
            "Add a Set / clear legal hold control on this page (currently the field is read-only).",
        ],
    },
    {
        "no": "32.11",
        "category": "Admin / Team Lead Extras",
        "feature": "Vendor email templates (Admin / Team Lead)",
        "route": "Settings → Vendor Templates card (/admin/vendor-templates). Team Lead and Admin.",
        "how_to_test": [
            "Sign in as Admin or Team Lead and open the Vendor Templates card from the Settings hub.",
            "Confirm five system templates ship with every tenant: Inspection — Schedule Visit, Inspection — Reschedule, Appraisal — Schedule Visit, Title — Document Request, Generic Vendor — Scheduling.",
            "Open any system template. Confirm its body includes the required reply footer 'Reply with: Scheduled: YYYY-MM-DD'. Tune the subject or body and click Save.",
            "Click New template, fill in the name, category, subject, and body, then save. Confirm the new template appears with a 'Custom' label.",
            "Sign out and back in as an Agent, then open the Send Vendor Request modal from a vendor's contact card — confirm the new custom template appears in the template picker.",
            "Deactivate a custom template and confirm it no longer appears for agents.",
        ],
        "expected_result": [
            "Every tenant has the five system templates available out of the box.",
            "Admins and Team Leads can add, edit, and deactivate custom templates; agents can use them but cannot edit them.",
            "System template names and categories are locked, but their subject and body can still be tuned.",
        ],
        "future_ideas": [
            "Show a live preview of the rendered email with a sample transaction filled in.",
            "Allow importing and exporting templates as a single file.",
            "Track how often each template is used so the team can retire low-value ones.",
        ],
    },
    {
        "no": "32.12",
        "category": "Admin / Team Lead Extras",
        "feature": "Communication audit page",
        "route": "/admin/communications (the legacy /communications URL still redirects here)",
        "how_to_test": [
            "Sign in as Team Lead or Admin and open Oversight → Communication Audit from the sidebar.",
            "Confirm two tabs: 'Audit log' (visible to Team Lead and Admin) and 'Export requests' (visible to Admin only).",
            "On the Audit log tab, use the filter row: search, channel, direction, date range, AI-only, vendor-only, and party email.",
            "Filter to a single transaction and click Download CSV — a single-transaction CSV downloads.",
            "Open Export requests (Admin only) and submit a multi-transaction export request. Confirm it appears in the list and can be downloaded once it finishes.",
            "Try to open /admin/communications as a regular Agent — the page should redirect to Unauthorized.",
        ],
        "expected_result": [
            "Team Lead and Admin can audit every communication across every transaction.",
            "Single-transaction CSV downloads work in the foreground; multi-transaction exports go through the request queue.",
            "Agents and lower roles cannot reach the page.",
        ],
        "future_ideas": [
            "Add a saved-search feature so admins can revisit the same filter quickly.",
            "Email the requester automatically when a multi-transaction export is ready.",
        ],
    },
    {
        "no": "32.13",
        "category": "Admin / Team Lead Extras",
        "feature": "Teams",
        "route": "/admin/teams (sidebar → Team → Teams). Admin and Team Lead.",
        "how_to_test": [
            "Sign in as an Admin and open Teams from the sidebar. Confirm a two-part layout: a list of teams on the left, and the selected team's roster and actions on the right. The first team is selected automatically.",
            "In the left list, confirm each team shows its lead and member count. Click a different team and confirm the right side updates.",
            "For the selected team, confirm its member table appears right there (no jump to another page), with these action buttons: Invite to team, Manage members, and (Admin only) Rename and Delete.",
            "Click New team (top right, Admin only), fill in the setup dialog (name and lead), and save — confirm the new team appears and is selected.",
            "Click Manage members and confirm you can add members to, or remove members from, the team. Click Invite to team and confirm the invite window opens with the team already filled in.",
            "Click Rename, change the name or lead, and save. Click Delete and confirm a warning explains the members will be unassigned (their accounts stay active) before it is removed.",
            "Sign in as a Team Lead and confirm Teams shows only your own team, with Invite to team and Manage members available but no New team / Rename / Delete.",
        ],
        "expected_result": [
            "Teams is a single workspace: pick a team on the left, manage it on the right.",
            "Admins can create, rename, and delete teams and manage any team's members; a Team Lead manages only their own team's members.",
            "Every destructive change asks for confirmation and updates the list immediately.",
        ],
        "future_ideas": [
            "Show each team's active-deal count and pipeline value in the left list.",
            "Allow moving several members between teams at once.",
        ],
    },
    {
        "no": "32.14",
        "category": "Admin / Team Lead Extras",
        "feature": "Team Playbook (team-wide)",
        "route": "Settings → Team Playbook card (/admin/team-settings). Team Lead and Admin.",
        "how_to_test": [
            "From the Settings hub, open the Team Playbook card. Confirm one page with a row of buttons across the top — Checklist, Tagged notes, Preferred vendors, and Resources — that switch the team tool in place (the same shape as My Playbook, but team-wide).",
            "As an Admin, use the team picker in the header ('Editing for …') to choose which team you are editing. A Team Lead is automatically locked to their own team and sees an 'Editing your team's playbook' note instead of the picker.",
            "In each tool, add or edit an entry and save. Confirm it persists.",
            "Click 'Preview closing checklist' and confirm a preview of how this team's playbook prints on a sample closing.",
            "Confirm these are the team-wide versions; the personal 'My …' copies live in Settings → My Playbook (feature 28.2).",
        ],
        "expected_result": [
            "One card edits a team's whole playbook (checklists, notes, vendors, resources) from a single page.",
            "Admins choose a team; Team Leads are scoped to their own team automatically.",
            "Saved changes apply to everyone on that team.",
        ],
        "future_ideas": [
            "Add a 'copy from another team' shortcut to clone a playbook.",
            "Show which agents are using each checklist template.",
        ],
    },
    {
        "no": "32.15",
        "category": "Admin / Team Lead Extras",
        "feature": "Admin — Integrations (CRM / webhooks)",
        "route": "Settings → Integrations & Webhooks card (/admin/integrations). Admin or workspace owner.",
        "how_to_test": [
            "Sign in as an Admin and open the Integrations & Webhooks card from the Settings hub.",
            "Click to register a new webhook endpoint, pick the events it should receive, and save.",
            "Copy the signing secret with the copy button.",
            "Fire a test event and confirm a delivery attempt is recorded with its result.",
            "Review the delivery history for an endpoint.",
        ],
        "expected_result": [
            "An Admin can set up a CRM / integration webhook end-to-end without an engineer.",
            "Test events and delivery history both work.",
        ],
        "future_ideas": [
            "Add ready-made templates for popular CRMs.",
            "Add automatic retry with backoff for failed deliveries.",
        ],
    },
    {
        "no": "32.16",
        "category": "Admin / Team Lead Extras",
        "feature": "Admin — AI & Automation",
        "route": "Settings → AI & Automation card (/admin/confidence). Admin or workspace owner.",
        "how_to_test": [
            "Sign in as an Admin and open the AI & Automation card from the Settings hub.",
            "Confirm plain-English 'what AI can do' and 'what AI cannot do' lists.",
            "Adjust the confidence thresholds (for example the review threshold and the auto-send threshold) and save.",
            "Confirm the change affects how AI Email Review classifies and auto-sends drafts (feature 29.7+).",
        ],
        "expected_result": [
            "Admins control the AI confidence thresholds for the whole workspace.",
            "AI never sends below the review threshold without a human; the page states this clearly.",
        ],
        "future_ideas": [
            "Show a recent history of auto-sent vs held drafts at each threshold.",
            "Allow per-team thresholds, not just per-workspace.",
        ],
    },
    {
        "no": "32.17",
        "category": "Admin / Team Lead Extras",
        "feature": "Admin — Payment Access",
        "route": "Settings → Payment Access card (/admin/payment-access). Admin or workspace owner.",
        "how_to_test": [
            "Sign in as an Admin and open the Payment Access card from the Settings hub.",
            "Confirm a grid of roles (Agent, Elf / Transaction Coordinator, Team Lead) against capabilities (create invoice, refund, trigger payout).",
            "Turn a capability on or off for a role and save.",
            "Sign in as that role and confirm the matching action appears or disappears on the Payments pages (features 26.8–26.9).",
        ],
        "expected_result": [
            "Admins decide which roles can create invoices, issue refunds, and trigger payouts.",
            "Changes take effect on the Payments pages for that role.",
        ],
        "future_ideas": [
            "Allow per-person overrides, not only per-role.",
            "Show an audit trail of who changed which capability.",
        ],
    },
    {
        "no": "32.18",
        "category": "Admin / Team Lead Extras",
        "feature": "Admin — Advertising",
        "route": "Settings → Advertising card (/admin/advertising). Admin or workspace owner.",
        "how_to_test": [
            "Sign in as an Admin and open the Advertising card from the Settings hub. Confirm three cards: Workspace ads (a single on/off toggle, OFF by default), Your house ads, and Performance.",
            "Turn Workspace ads on and confirm sponsored placements may then appear in the workspace; turn it off and confirm they stop.",
            "Create a house ad: add the details, upload an image, and approve it. Confirm each ad shows a plain-English 'why it is / isn't showing' chip.",
            "Pause a house ad and confirm it stops showing.",
            "Check the Performance card for impressions, clicks, and click-through rate.",
        ],
        "expected_result": [
            "Ads are OFF until an Admin explicitly turns them on.",
            "House ads can be created, approved, and paused, each with a clear status reason.",
            "Performance numbers reflect this workspace's ads.",
        ],
        "future_ideas": [
            "Add scheduling (start / end dates) for a house ad.",
            "Add simple A/B testing for two creatives.",
        ],
    },
    {
        "no": "32.19",
        "category": "Admin / Team Lead Extras",
        "feature": "Admin — Audit Log",
        "route": "/admin/audit-logs (Admin only; sidebar → Oversight → Audit Log)",
        "how_to_test": [
            "Sign in as an Admin and open Oversight → Audit Log from the sidebar.",
            "Confirm a list of recorded actions across the workspace (documents, transactions, tasks, users, invitations, vendors, AI emails, and so on).",
            "Filter by entity type and confirm the list narrows.",
            "Expand an entry and confirm it shows who did what and when.",
            "Scroll to load more entries.",
        ],
        "expected_result": [
            "Every meaningful action is recorded and filterable.",
            "Each entry clearly shows the actor, the action, and the time.",
        ],
        "future_ideas": [
            "Add a date-range filter and free-text search.",
            "Add CSV export for a filtered view.",
        ],
    },
    {
        "no": "32.20",
        "category": "Admin / Team Lead Extras",
        "feature": "Platform Admin — Advertising (internal Velvet Elves staff only)",
        "route": "/platform/advertising (platform admins only)",
        "how_to_test": [
            "Sign in as a platform admin and open the Platform Advertising card from the Settings hub (Platform group).",
            "Sign in as any non-platform user and try /platform/advertising directly — confirm you get a 404 (the route's existence is not leaked).",
            "Review and manage the platform-wide / partner ad inventory shown here.",
        ],
        "expected_result": [
            "Only platform admins can reach the page; everyone else gets a 404.",
            "Platform-level ad management renders without errors.",
        ],
        "future_ideas": [
            "Add per-tenant targeting controls for partner ads.",
            "Show platform-wide ad performance roll-ups.",
        ],
    },
    {
        "no": "32.21",
        "category": "Admin / Team Lead Extras",
        "feature": "Platform Admin — AI Usage & cost (internal Velvet Elves staff only)",
        "route": "/platform/ai-usage (platform admins only)",
        "how_to_test": [
            "Sign in as a platform admin and open Platform → AI Usage.",
            "Sign in as any non-platform user and try /platform/ai-usage directly — confirm you get a 404 (the route's existence is not leaked).",
            "Confirm summary totals (total AI cost, calls, and tokens) and read-only breakdown tables by tenant, by transaction, and by feature.",
        ],
        "expected_result": [
            "Only platform admins can reach the page; everyone else gets a 404.",
            "The cost and usage figures render as read-only totals and tables — this is a measurement tool, with no pricing or checkout.",
        ],
        "future_ideas": [
            "Add a date-range filter to compare periods.",
            "Let staff export the usage breakdown as a CSV.",
            "Add a per-tenant cost-per-deal average column.",
        ],
    },
    {
        "no": "32.22",
        "category": "Admin / Team Lead Extras",
        "feature": "Platform Admin — Help Center authoring (internal Velvet Elves staff only)",
        "route": "/platform/help — sidebar → Platform → Help center (platform admins only)",
        "how_to_test": [
            "Sign in as a platform admin and open Platform → Help center from the sidebar. This is where the public Help Center website's content is written and published.",
            "Sign in as any non-platform user and try /platform/help directly — confirm you get a 404 (the route's existence is not leaked).",
            "Confirm a knowledge-base layout: a list of collections (categories) on the left, and an editing pane on the right. Use the search box to find a collection.",
            "Click New collection, give it a name, and save — confirm it appears in the tree. Open a collection and confirm you can add an article inside it.",
            "Open an article. Confirm a writing column with a simple formatting toolbar and an Edit / Preview toggle (Preview shows the reader's view). Make a change and click Save.",
            "Confirm the article status (Draft / Published) is shown as a read-only badge, and a separate Publish (or Unpublish) button asks you to confirm before it changes what the public sees. Confirm Archive and Delete live in a 'More' menu.",
            "Use the Feedback and Settings buttons in the header and confirm they open the reader-feedback list and the Help Center settings.",
        ],
        "expected_result": [
            "Only platform admins can reach the page; everyone else gets a 404.",
            "Collections and articles can be created, edited, previewed, and published, with publishing always behind a clear confirmation.",
            "Changes here drive the public Help Center website that customers reach from the in-app 'Help Center' link.",
        ],
        "future_ideas": [
            "Show where each article links from so staff can spot orphaned guides.",
            "Add a side-by-side edit-and-preview mode for longer articles.",
            "Track which articles get the most 'not helpful' votes so staff know what to rewrite first.",
        ],
    },

    # ============================================================
    # SECTION 5 — ATTORNEY WORKSPACE
    # ============================================================
    {
        "no": 33,
        "category": "Attorney Workspace",
        "feature": "Attorney workspace — Matters list",
        "route": "/transactions (as Attorney); the sidebar entry is 'Matters'",
        "how_to_test": [
            "Sign in as an Attorney and open Matters from the sidebar (or /transactions). The page should automatically load the attorney layout, not the agent transactions list.",
            "Confirm the attorney KPI row at the top (for example Hard Stops, Release Ready, Active Matters, Needs Review).",
            "Use the filter tabs: All, Needs Review, Missing Docs, Ready To Release, Clean Files. Confirm the matter cards narrow as you switch.",
            "On a matter card, click Review to open that matter's full workspace (feature 33.1).",
            "Click the floating Ask AI button and confirm the AI panel opens.",
            "Click the '+ Upload Legal Packet' button (sidebar footer / top bar) and confirm the legal-packet upload flow opens.",
        ],
        "expected_result": [
            "The attorney layout loads with the correct KPI row, tabs, and matter cards.",
            "Review opens the matter workspace; Upload Legal Packet opens the intake flow; Ask AI opens the AI panel.",
            "Releases, State Rules, and Recording Calendar are now their own sidebar pages (features 33.2–33.4) — they are no longer header buttons.",
        ],
        "future_ideas": [
            "Add a 'Sign off all in this packet' bulk action from the card.",
            "Let the attorney save a custom matter filter (for example 'closing this week, needs review').",
            "Show the responsible agent on each matter card.",
        ],
    },
    {
        "no": "33.1",
        "category": "Attorney Workspace",
        "feature": "Attorney matter workspace (one matter in depth)",
        "route": "/transactions/<id> (as Attorney) — opens from Review on a matter card",
        "how_to_test": [
            "Open a matter from the Matters list. Confirm a full-screen workspace with a header (property address, status, and a matter switcher to jump between matters) rather than a simple scrolling page.",
            "Use the left section rail to move between Overview, Review, Brief, Timeline, People, Activity, and Releases. Confirm each section loads its own content.",
            "On the Review section, work through the document review items. On the Releases section, confirm you can start a packet release.",
            "Use the matter switcher in the header to jump to a different matter without going back to the list.",
        ],
        "expected_result": [
            "The matter opens as a focused workspace; each rail section shows real data for that matter.",
            "The matter switcher moves you between matters in place.",
        ],
        "future_ideas": [
            "Add keyboard shortcuts to move between rail sections.",
            "Add a one-click 'export this matter file as PDF' for the closing binder.",
        ],
    },
    {
        "no": "33.2",
        "category": "Attorney Workspace",
        "feature": "Attorney Releases Queue",
        "route": "/attorney/releases (sidebar → Releases Queue)",
        "how_to_test": [
            "Open Releases Queue from the sidebar. Confirm a list of matters that are ready to release and a history of recently released packets.",
            "On a ready matter, click the release action and confirm the Send Packet window opens with the recipients and documents pre-filled.",
            "Send a packet and confirm it moves into the released history with the date and recipients.",
        ],
        "expected_result": [
            "Only matters that are actually ready to release appear in the ready list.",
            "Sending a packet records who it went to and when.",
        ],
        "future_ideas": [
            "Add a one-click 'release all ready matters' for a quiet end-of-day sweep.",
            "Let the attorney attach a short cover note to the released packet.",
        ],
    },
    {
        "no": "33.3",
        "category": "Attorney Workspace",
        "feature": "Attorney State Rules",
        "route": "/attorney/state-rules (sidebar → State Rules)",
        "how_to_test": [
            "Open State Rules from the sidebar.",
            "Confirm a clean reference document listing, per state, the closing type (attorney / title / escrow), the recording window, and whether same-day disbursement is allowed.",
            "Scroll through and confirm it reads as a reference page, not a dashboard.",
        ],
        "expected_result": [
            "The page is a read-only reference of state recording / closing rules.",
            "Every state your matters touch is listed.",
        ],
        "future_ideas": [
            "Add a search box to jump to a state quickly.",
            "Link each state to the matters currently in that state.",
        ],
    },
    {
        "no": "33.4",
        "category": "Attorney Workspace",
        "feature": "Attorney Recording Calendar",
        "route": "/attorney/recording-calendar (sidebar → Recording Calendar)",
        "how_to_test": [
            "Open Recording Calendar from the sidebar.",
            "Confirm a month grid with recording deadlines / closings marked on their dates.",
            "Move between months with the arrows.",
            "Click Print and confirm a printable calendar opens.",
        ],
        "expected_result": [
            "The calendar shows the matters' recording deadlines on the right dates.",
            "Month navigation and Print both work.",
        ],
        "future_ideas": [
            "Let the attorney click a day to see every matter due that day.",
            "Add an agenda (list) view alongside the month grid.",
        ],
    },

    # ============================================================
    # SECTION 6 — CLIENT, FSBO & VENDOR PORTALS
    # ============================================================
    {
        "no": 34,
        "category": "Client, FSBO & Vendor Portals",
        "feature": "Client workspace — Home (closing concierge)",
        "route": "/client/home (Client sign-in lands here)",
        "how_to_test": [
            "Sign in as a Client (a buyer or seller invited to a transaction). You should land on a warm 'closing concierge' Home with its own navy sidebar — not the staff app layout.",
            "Confirm the Home shows where your deal stands, what is coming next, your key dates, and a short list of documents and your agent.",
            "Use the 'Ask Velvet' / 'Ask your agent' box to send a message, and confirm it posts to a real two-way thread.",
            "Confirm the left sidebar shows: Home, Next Steps, Timeline, Documents, Updates. (Your Payments and Agent Info pages are reachable from links on the Home cards.)",
            "Open Next Steps (your action items and key dates) and Updates (recent updates plus the message thread with your agent) and confirm each loads with real content.",
            "If you are a brand-new client with no transaction yet, confirm a friendly empty Home appears instead of fake sample data.",
        ],
        "expected_result": [
            "Clients get their own concierge workspace, not the internal staff layout.",
            "Every number and date is real; an empty account shows an honest empty state.",
            "The Ask box reaches the agent and shows replies.",
        ],
        "future_ideas": [
            "Add a single 'what should I do next' banner at the very top.",
            "Let the client switch between several of their transactions from the Home header.",
        ],
    },
    {
        "no": "34.1",
        "category": "Client, FSBO & Vendor Portals",
        "feature": "Client workspace — Timeline",
        "route": "/client/milestones (sidebar → Timeline)",
        "how_to_test": [
            "Open Timeline from the client sidebar. Confirm one card per transaction summarising where the deal stands and the closing date.",
            "Tap a transaction card and confirm it opens that deal's full step-by-step timeline.",
        ],
        "expected_result": [
            "Clients with more than one deal see a calm card per deal; each opens its own detailed timeline.",
            "Every step and date is driven by the real transaction, not a template.",
        ],
        "future_ideas": [
            "Add an estimated date next to upcoming steps.",
            "Let the client turn on email alerts when a step completes.",
        ],
    },
    {
        "no": "34.2",
        "category": "Client, FSBO & Vendor Portals",
        "feature": "Client workspace — Documents",
        "route": "/client/documents (sidebar → Documents)",
        "how_to_test": [
            "Open Documents from the client sidebar. Confirm it leads with anything waiting on you, then a real status summary, then your document list.",
            "Click the upload action, pick the transaction and document type, and upload a file. Confirm it appears in the list.",
            "On a document, use Flag for deletion, give a reason, and confirm a 'Flagged' badge appears (this routes to the agent's Deletion Queue — feature 27.11).",
        ],
        "expected_result": [
            "The client sees a real document list and status summary, never a hardcoded zero board.",
            "Upload and flag-for-deletion both work and the agent is notified of flags.",
        ],
        "future_ideas": [
            "Add an inline preview so the client can re-read a document before flagging.",
            "Show which transaction each document belongs to when the client has several deals.",
        ],
    },
    {
        "no": "34.3",
        "category": "Client, FSBO & Vendor Portals",
        "feature": "Client / FSBO workspace — Payments (invoices)",
        "route": "/client/invoices (Client and FSBO sidebar → Payments)",
        "how_to_test": [
            "Sign in as a Client or FSBO Customer and open Payments from the sidebar.",
            "Confirm a list of invoices on your transactions with their amount, status, and due date.",
            "Open an invoice and confirm you can pay it securely (Stripe). After paying in test mode, confirm the status updates to Paid.",
        ],
        "expected_result": [
            "Clients and FSBO customers see only their own invoices.",
            "Paying an invoice updates its status and records the payment.",
        ],
        "future_ideas": [
            "Email a receipt automatically after payment.",
            "Let the payer save a card for future invoices.",
        ],
    },
    {
        "no": "34.4",
        "category": "Client, FSBO & Vendor Portals",
        "feature": "Client workspace — Agent Info (your team)",
        "route": "/client/agent (sidebar → Agent Info)",
        "how_to_test": [
            "Open Agent Info from the client sidebar.",
            "Confirm your agent's details and the deal's key contacts (loan officer, title, etc.) with one-tap call and email.",
            "Confirm the agent's short bio appears (this is the bio the agent set in their Account → Profile).",
        ],
        "expected_result": [
            "The page shows the real agent and key contacts for the deal, read-only.",
            "Call and email shortcuts work on a phone.",
        ],
        "future_ideas": [
            "Add the agent's photo and office hours.",
            "Add a one-tap 'message my agent' that opens the Ask thread.",
        ],
    },
    {
        "no": 35,
        "category": "Client, FSBO & Vendor Portals",
        "feature": "FSBO workspace — Overview & sidebar",
        "route": "/fsbo (FSBO Customer sign-in lands here)",
        "how_to_test": [
            "Sign in as an FSBO Customer. Confirm you land on an Overview dashboard with KPI tiles (for example My Properties, Missing Docs, Share Links Live, Days To Close) and a floating Ask AI button.",
            "Confirm the sidebar shows: Dashboard, My Properties, Documents, Payments, Messages.",
            "Confirm the sidebar footer has a 'Share milestones' button.",
            "Click through each sidebar item and confirm every page loads (none should be a 'Coming Soon' placeholder).",
        ],
        "expected_result": [
            "The FSBO Overview shows real numbers or clean empty states.",
            "Every FSBO sidebar destination is a working page.",
            "Share milestones opens the share-link manager.",
        ],
        "future_ideas": [
            "Simplify the home screen further for less technical sellers.",
            "Add a guided 'first week as a FSBO seller' checklist.",
        ],
    },
    {
        "no": "35.1",
        "category": "Client, FSBO & Vendor Portals",
        "feature": "FSBO workspace — My Properties (and property detail)",
        "route": "/fsbo/properties (sidebar → My Properties)",
        "how_to_test": [
            "Open My Properties. Confirm one card per home with its status, closing date, outstanding-document count, and unread-message count.",
            "Open a property card and confirm its workspace opens (milestones, documents, and messages for that one home).",
            "From a property, use the 'Manage' / share action and confirm the share-link manager opens.",
        ],
        "expected_result": [
            "Every property the seller owns appears as a scannable card.",
            "Opening a property shows that home's full workspace.",
        ],
        "future_ideas": [
            "Add a filter strip (active / closing soon / closed) above the property cards.",
            "Show the next action needed on each property card.",
        ],
    },
    {
        "no": "35.2",
        "category": "Client, FSBO & Vendor Portals",
        "feature": "FSBO workspace — Documents",
        "route": "/fsbo/documents (sidebar → Documents)",
        "how_to_test": [
            "Open Documents. Confirm a count badge and a filter-tab strip (All / Missing / In progress / Uploaded / Verified / Complete) over one combined list across all your properties.",
            "Each row should show the document, its status, and a tag for which property it belongs to.",
            "On a missing requirement, click Upload — confirm the upload window opens with that property pre-selected.",
            "On a document, use Flag for deletion and confirm a reason is required and the row updates.",
        ],
        "expected_result": [
            "All documents across every property show in one place, filterable by status.",
            "Upload and flag-for-deletion both work.",
        ],
        "future_ideas": [
            "Add a 'download everything for this property' button.",
            "Auto-detect the document type from the uploaded file.",
        ],
    },
    {
        "no": "35.3",
        "category": "Client, FSBO & Vendor Portals",
        "feature": "FSBO workspace — Messages",
        "route": "/fsbo/milestones (sidebar → Messages)",
        "how_to_test": [
            "Open Messages. Confirm a single inbox of everything your coordinator has sent you, across all properties.",
            "Confirm each message is tagged with the property it belongs to.",
            "Use the filter-tab strip to narrow the list.",
        ],
        "expected_result": [
            "One unified message inbox across every property.",
            "Each message clearly shows which property it relates to.",
        ],
        "future_ideas": [
            "Let the seller reply to a coordinator message directly from this inbox.",
            "Add unread / read filters.",
        ],
    },
    {
        "no": 36,
        "category": "Client, FSBO & Vendor Portals",
        "feature": "Vendor portal — Document requests and uploads",
        "route": "/portal/vendor (Vendor sign-in lands here)",
        "how_to_test": [
            "Sign in as a Vendor. Confirm a focused portal showing the documents requested from you and an upload area — not the staff app.",
            "Upload a requested document and confirm it is accepted and shows in your uploads.",
            "Switch to the My Uploads view (sidebar) and confirm your previously uploaded files are listed.",
            "Confirm the portal only ever shows your own requests and uploads — no transaction details you should not see.",
        ],
        "expected_result": [
            "Vendors see only their own document requests and uploads.",
            "Upload works and the file reaches the right transaction for the agent.",
        ],
        "future_ideas": [
            "Let a vendor reply to a request with a short note alongside the file.",
            "Show the due date for each requested document.",
        ],
    },

    # ============================================================
    # SECTION 7 — PUBLIC LINKS (NO SIGN-IN)
    # ============================================================
    {
        "no": 37,
        "category": "Public Links (No Sign-In)",
        "feature": "Public milestone viewer",
        "route": "/milestones/<share token> (opened from a share link an agent or seller sent)",
        "how_to_test": [
            "Open a milestone share link in a browser where you are signed out (or a private window).",
            "Confirm a clean, read-only progress page showing the property address and the milestone steps with their status — branded with the sharing brokerage's name and color.",
            "Confirm there is no sign-in prompt and no private contact or financial detail.",
            "Open a made-up / expired token and confirm a clear 'link not available' message rather than a broken page.",
        ],
        "expected_result": [
            "Anyone with the link sees the milestone progress without signing in.",
            "The page is read-only and never leaks private details.",
            "An invalid or expired link shows a clean message.",
        ],
        "future_ideas": [
            "Show an estimated closing date on the viewer.",
            "Let the viewer subscribe to email updates when a step completes.",
        ],
    },
    {
        "no": 38,
        "category": "Public Links (No Sign-In)",
        "feature": "Public invoice payment link",
        "route": "/pay/invoices/<invoice id> (opened from an invoice email)",
        "how_to_test": [
            "Open a public invoice pay link while signed out.",
            "Confirm the page shows the amount, who it is for, the property, and the due date, with a secure 'Pay' button (Stripe).",
            "Pay in test mode and confirm you are taken to a 'payment complete' confirmation page.",
            "Open the link again after paying and confirm it shows the invoice is already paid rather than charging twice.",
        ],
        "expected_result": [
            "A payer with no account can pay an invoice securely from the link.",
            "After payment a clear confirmation page appears and the invoice is marked paid.",
        ],
        "future_ideas": [
            "Offer a downloadable PDF receipt on the confirmation page.",
            "Support partial / installment payments.",
        ],
    },
    {
        "no": 39,
        "category": "Public Links (No Sign-In)",
        "feature": "Advertise storefront",
        "route": "/advertise (public marketing + checkout)",
        "how_to_test": [
            "Open /advertise while signed out. Confirm a marketing landing page explaining advertising on Velvet Elves with a clear call to action.",
            "Start the checkout flow and confirm it walks you through choosing a placement and paying.",
            "Complete a test checkout and confirm a completion / confirmation page appears.",
        ],
        "expected_result": [
            "The storefront is reachable with no sign-in and explains the offer clearly.",
            "Checkout and completion both work end to end in test mode.",
        ],
        "future_ideas": [
            "Show live example placements so advertisers see what they are buying.",
            "Add package tiers (week / month / quarter) with clear pricing.",
        ],
    },

    # ============================================================
    # SECTION 6 — DIRECT LINKS AND ERROR PAGES
    # ============================================================
    {
        "no": 40,
        "category": "Direct Links and Error Pages",
        "feature": "Unauthorized page",
        "route": "/unauthorized (or any blocked page)",
        "how_to_test": [
            ("Sign in as a basic Agent and open one of these URLs directly:", [
                "/team",
                "/admin/task-templates",
                "/admin/users/some-id",
            ]),
        ],
        "expected_result": [
            "The app redirects you to /unauthorized.",
            "A clear 'Access denied' screen appears.",
        ],
        "future_ideas": [
            "Add a 'Request access from your admin' button that emails the tenant admin.",
            "Explain in plain language which role is needed to open the page.",
        ],
    },
    {
        "no": 41,
        "category": "Direct Links and Error Pages",
        "feature": "Not Found (404) page",
        "route": "Any unknown URL, for example /this-does-not-exist",
        "how_to_test": [
            ("Type an invalid URL in the browser address bar and press Enter.", []),
        ],
        "expected_result": [
            "A 404 page appears with a Back to Dashboard button.",
        ],
        "future_ideas": [
            "Add a small search bar on the 404 page that suggests the likely intended page based on the URL.",
        ],
    },
]


TARGET_BUILDERS = {
    "requirements": create_requirements_doc,
    "milestones": create_milestones_doc,
    "design-feedback": create_design_feedback_doc,
    "testing-review": create_testing_review_doc,
    "transaction-system": create_transaction_system_doc,
    "payment-system": create_payment_system_doc,
    "godaddy-route53-config": create_godaddy_route53_config_doc,
    "transaction-processing-method": create_transaction_processing_method_doc,
    "transaction-automation-guide": create_transaction_automation_guide_doc,
    "gmail-approval-plan": create_gmail_approval_plan_doc,
}

TARGET_ALIASES = {
    "requirements.txt": "requirements",
    "requirements.docx": "requirements",
    "milestones.txt": "milestones",
    "milestones.docx": "milestones",
    "design-feedback.md": "design-feedback",
    "design-feedback.docx": "design-feedback",
    "design_feedback": "design-feedback",
    "feedback": "design-feedback",
    "testing_review": "testing-review",
    "testing": "testing-review",
    "review": "testing-review",
    "frontend_client_testing_review.csv": "testing-review",
    "frontend_client_testing_review.docx": "testing-review",
    "transaction_system": "transaction-system",
    "transaction-system-guide": "transaction-system",
    "transaction_system_guide": "transaction-system",
    "transaction_system_guide.md": "transaction-system",
    "transaction_system_guide.docx": "transaction-system",
    "guide": "transaction-system",
    "payment_system": "payment-system",
    "payment-system-guide": "payment-system",
    "payment_system_guide": "payment-system",
    "payment_system_guide.md": "payment-system",
    "payment_system_guide.docx": "payment-system",
    "payments": "payment-system",
    "payment": "payment-system",
    "godaddy_route53_configuration_guide": "godaddy-route53-config",
    "godaddy_route53_configuration_guide.md": "godaddy-route53-config",
    "godaddy_route53_configuration_guide.docx": "godaddy-route53-config",
    "godaddy-route53": "godaddy-route53-config",
    "route53-godaddy": "godaddy-route53-config",
    "transaction_processing_method": "transaction-processing-method",
    "transaction-processing": "transaction-processing-method",
    "transaction_processing": "transaction-processing-method",
    "transaction_processing_method.md": "transaction-processing-method",
    "transaction_processing_method.docx": "transaction-processing-method",
    "processing-method": "transaction-processing-method",
    "method": "transaction-processing-method",
    "gmail_google_approval_plan": "gmail-approval-plan",
    "gmail_google_approval_plan.md": "gmail-approval-plan",
    "gmail_google_approval_plan.docx": "gmail-approval-plan",
    "gmail-approval": "gmail-approval-plan",
    "approval-plan": "gmail-approval-plan",
    "all": "all",
}


def normalize_targets(raw_targets):
    """Normalize CLI target names into supported document keys."""
    normalized_targets = []

    for raw_target in raw_targets:
        target = TARGET_ALIASES.get(raw_target.lower(), raw_target.lower())
        if target == "all":
            for name in TARGET_BUILDERS:
                if name not in normalized_targets:
                    normalized_targets.append(name)
            continue

        if target not in TARGET_BUILDERS:
            valid = ", ".join(list(TARGET_BUILDERS.keys()) + ["all"])
            raise SystemExit(f"Unknown document target '{raw_target}'. Use one of: {valid}")

        if target not in normalized_targets:
            normalized_targets.append(target)

    return normalized_targets


def main():
    """Parse CLI args and generate the requested documents."""
    parser = argparse.ArgumentParser(description="Generate Word documents from project source files.")
    parser.add_argument(
        "documents",
        nargs="*",
        help="Documents to generate: requirements, milestones, design-feedback, or all.",
    )
    args = parser.parse_args()

    targets = normalize_targets(args.documents or ["requirements", "milestones"])
    for target in targets:
        TARGET_BUILDERS[target]()

    print(f"Done! Created {len(targets)} document(s) in {BASE_DIR}")


if __name__ == "__main__":
    main()
